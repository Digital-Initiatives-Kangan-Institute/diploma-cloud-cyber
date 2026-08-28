#!/usr/bin/env python3
"""Build the supplied Accounting System Baseline Solution Design (.docx) — in-world document.

The PRACTICE analogue of the YAT-LMS Baseline Solution Design: the supplied cloud architecture
students implement in the AT2 *practice* build (the Accounting System / Ledgerline engagement),
paralleling the real AT2's LMS baseline. Parallel-but-different: PostgreSQL on Amazon Linux (not
MySQL on Windows), 10.20.0.0/16 (not 10.0.0.0/16), business-hours 99.5% (not 24/7 99.9%), commercial
licensing + 7-year financial-records retention. Single-AZ baseline (non-HA); HA sections
marked "Not applicable" per the convention. In-world artefact (no UoC tags).

Output to the website documents folder for printing to PDF, then wiring into the AT2/AT3 intranet states.

Usage:  python scripts/scenario/build_accounting_solution_design.py [output.docx]
Default: ../diploma-cloud-cyber-website-s1/public/documents/YAT-Accounting-Baseline-Solution-Design.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # content-repo scripts/ (brand + registry)  # noqa: E402
sys.path.insert(0, str(next(d / "scripts" for d in Path(__file__).resolve().parents if (d / "scripts" / "helpers" / "__init__.py").exists())))  # umbrella scripts/ (engine)  # noqa: E402
from helpers.docx_body_text import add_body_paragraph, add_bullet_list  # noqa: E402
from helpers.docx_tables import add_data_table  # noqa: E402
from helpers.docx_styling import add_field, paragraph_bottom_rule, set_cell_borders, shade_cell  # noqa: E402
from brand import ADDRESS, CREAM, GREY, TEAL, TERRACOTTA  # noqa: E402
from helpers.scenario_document import build_header_footer, configure_styles, wordmark  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def na(doc, reason):
    p = doc.add_paragraph()
    r = p.add_run(f"Not applicable — {reason}")
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    p.paragraph_format.space_after = Pt(6)
    return p


def lab_note(doc, text):
    """Where the design and the build environment differ, say so in one plain sentence.

    Same convention as the LMS baseline design: the design is the real design; some of it cannot be
    built in an AWS Academy Learner Lab, and in one place it is deliberately simplified so a student
    can see their own work running. Each point carries a note under the thing it qualifies.
    """
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    lead = p.add_run("Building this in a lab ")
    lead.bold = True; lead.font.size = Pt(9.5)
    lead.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    r = p.add_run(text)
    r.italic = True; r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    return p


# Published diagrams (not for-documents-not-website): this topology is also served as an intranet
# page, so the .drawio/.svg/.png live at the top level and one asset serves both surfaces.
DIAGRAM_DIR = (Path(__file__).resolve().parents[3] / "diploma-cloud-cyber-website-s1"
               / "public" / "diagrams")


def diagram_figure(doc, caption, image_name, width_cm=16.0):
    """Place a diagram exported alongside its .drawio, captioned.

    The picture is placed by the generator, not pasted in afterwards, so it survives a rebuild.
    The .drawio is generated from a committed spec under scripts/scenario/diagrams/; the
    .drawio.png beside it is exported from draw.io, so the document shows the same picture as the
    published SVG. Fails loudly if the image is absent rather than leaving a silent hole.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    image = DIAGRAM_DIR / image_name
    if not image.exists():
        raise FileNotFoundError(
            f"Diagram not found: {image}\n"
            f"Render the .drawio from its spec, then export the .drawio.png from draw.io.")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image), width=Cm(width_cm))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.italic = True; cr.font.size = Pt(9); cr.font.color.rgb = RGBColor.from_string(GREY)
    return p


def diagram_placeholder(doc, caption, source):
    """A bordered, shaded drop-zone for a network diagram — replace with the exported image."""
    from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    set_cell_borders(cell)
    shade_cell(cell, CREAM)
    cell.width = Cm(16.6)
    t.rows[0].height = Cm(9.5)
    t.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("[ NETWORK TOPOLOGY DIAGRAM — PASTE HERE ]")
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    p2 = cell.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(source)
    r2.italic = True; r2.font.size = Pt(9); r2.font.color.rgb = RGBColor.from_string(GREY)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.italic = True; cr.font.size = Pt(9); cr.font.color.rgb = RGBColor.from_string(GREY)
    return t


def build(path):
    doc = Document()
    configure_styles(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    build_header_footer(sec)

    # ---- COVER ----
    wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(GREY)
    paragraph_bottom_rule(doc.add_paragraph(), TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Solution Design")
    sub = doc.add_paragraph().add_run("YAT Accounting System Cloud Architecture — Baseline Design")
    sub.font.size = Pt(15); sub.bold = True; sub.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    doc.add_paragraph()
    cover = [
        ("Engagement", "YAT Accounting System Cloud Migration — Foundation Build"),
        ("Document type", "Technical design (Solution Design)"),
        ("Version", "v1.0 — Approved for implementation"),
        ("Authored by", "MTS Senior Architecture Team, in consultation with YAT ICT"),
        ("Approved by", "Pat Lin (MTS Senior Consultant) · Sam Walker (YAT ICT Manager)"),
        ("Implemented by", "the Accounting System foundation-build Deployment Report"),
        ("Classification", "Internal — YAT ICT, Finance, and MTS personnel on signed MSA"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    for k, v in cover:
        cells = ct.add_row().cells
        set_cell_borders(cells[0]); set_cell_borders(cells[1]); shade_cell(cells[0], CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        cells[1].paragraphs[0].add_run(v).font.size = Pt(10)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    doc.add_paragraph("Contents", style="Heading 1")
    add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                 placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h3 = lambda t: doc.add_paragraph(t, style="Heading 3")

    h1("1. Purpose and Scope")
    add_body_paragraph(doc, "This document specifies the baseline AWS architecture MTS will implement as the foundation "
                 "build phase of the YAT Accounting System cloud migration. It translates the approved "
                 "direction into a concrete, implementable design for the Ledgerline finance and office-"
                 "administration system. The design stops at “infrastructure ready for application "
                 "deployment”: the EC2 instance is provisioned with the OS, the RDS instance with an empty "
                 "PostgreSQL engine, and the load balancer with placeholder health checks — no "
                 "application binaries or financial data are placed by the MTS build.")
    h3("In scope of this design")
    add_bullet_list(doc, [
        "The production cloud foundation for the Ledgerline accounting application.",
        "All compute, networking, identity, storage, database, autoscaling and baseline monitoring needed to run Ledgerline as a staff-facing, business-hours workload in AWS.",
        "Single-region, single-Availability-Zone deployment in ap-southeast-2 (Sydney).",
    ])
    lab_note(doc, "the design calls for ap-southeast-2 (Sydney). If you are building this in an AWS "
                  "Academy Learner Lab, deploy to whichever region that environment gives you — us-east-1 — "
                  "and treat it as standing in for Sydney throughout. That is acceptable.")
    h3("Out of scope — deferred to the follow-on HA design phase")
    add_bullet_list(doc, [
        "High-availability hardening (Multi-AZ database, cross-AZ compute resilience, failure-simulation testing).",
        "Disaster recovery to a second AWS region; DR runbook and tabletop testing.",
        "Application re-platforming (Ledgerline remains Amazon Linux + PostgreSQL).",
    ])
    h3("Out of MTS scope entirely — YAT ICT responsibility")
    add_bullet_list(doc, [
        "Ledgerline application installation onto the EC2 instance(s) after handover.",
        "Database migration (extract from the on-premises database, load into RDS for PostgreSQL).",
        "Cutover — DNS switch, parallel running, decommissioning, user redirection (avoiding month-end and EOFY).",
        "Organisational change management — CAB approvals, communications, training, post-cutover support.",
    ])

    h1("2. Design Inputs and Requirements")
    h3("2.1 Inputs")
    add_bullet_list(doc, [
        "Accounting System Application Specification — functions, user load, data, integrations, SLAs, data residency.",
        "Accounting System Migration Requirements — platform preservation, availability, recovery, licensing, retention.",
        "Accounting System Server Specifications — current server, measured availability, utilisation, growth.",
        "Engagement Role Brief and Consultation Notes — engagement scope; OS, application and database preservation.",
    ])
    h3("2.2 Requirements the design must meet")
    add_data_table(doc, ["Requirement", "Target / note"],
              [["Region / data residency", "ap-southeast-2 (Sydney); financial records + staff/debtor PII within Australia (Privacy Act APP 8; financial-records retention)"],
               ["Application stack", "Preserved — Amazon Linux 2023 · PostgreSQL · Ledgerline"],
               ["Concurrent users", "15–25 typical; 45–55 peak at month-end close and EOFY; idle out of hours"],
               ["Availability", "≥ 99.5% business-hours (Mon–Fri ~07:30–18:00); no 24/7 requirement"],
               ["Recovery", "RPO ≤ 1 hour (no loss of financial transactions); RTO ≤ 1 business day (≤ 8 business hours)"],
               ["Data footprint", "~22 GB database, growing ~5 GB/year, plus scanned-document attachments"],
               ["Retention", "Financial records and audit logs retained ≥ 7 years"],
               ["Licensing", "Commercial — Ledgerline per-user licences; the database is open-source (no database licence) is material"],
               ["Integrations", "AD authentication; O365 SMTP; LMS fee-status; payroll-bureau SFTP; banking/payment-gateway file exchange"],
               ["High availability", "Out of baseline scope — Multi-AZ resilience deferred to the follow-on HA design"]],
              widths=[4.5, 11.5])

    h1("3. Review of Existing Architecture")
    na(doc, "this is a greenfield cloud foundation; there is no existing cloud baseline to review. The "
            "on-premises current state is documented separately in the Accounting System Server "
            "Specifications and the ICT Environment Overview.")

    h1("4. Architecture Design")
    h3("4.1 Assumptions and constraints")
    add_data_table(doc, ["#", "Assumption / constraint", "Source"],
              [["A1", "Region must be ap-southeast-2 (Sydney) for financial-records data residency", "Migration Requirements; Privacy Policy"],
               ["A2", "Application stack preserved: Amazon Linux 2023, PostgreSQL, Ledgerline", "Migration Requirements; Role Brief"],
               ["A3", "Data footprint ~22 GB database, +~5 GB/year, plus scanned attachments", "Application Spec; Server Specs"],
               ["A4", "Concurrent users 15–25 typical, 45–55 at month-end / EOFY; idle out of hours", "Application Spec"],
               ["A5", "Financial records + staff/debtor PII must remain within Australia; retained ≥ 7 years", "Privacy Act 1988 APP 8; financial-records obligations"],
               ["A6", "Preserve AD authentication over a private network to campus AD; staff-only access", "Application Spec; Migration Requirements"],
               ["A7", "Baseline is NOT high-availability; Multi-AZ resilience deferred to the HA design", "Scoping decision"]],
              widths=[1.0, 9.5, 5.5])
    h3("4.2 AWS account and region")
    add_body_paragraph(doc, "An AWS account scoped to the migration engagement, provisioned by YAT and shared with MTS "
                 "for the build. Region ap-southeast-2 (Sydney); no deployment outside Australian regions, to "
                 "satisfy financial-records residency. Availability Zone ap-southeast-2a for the baseline; the "
                 "follow-on HA design introduces a second AZ.")
    h3("4.3 Identity and Access Management (IAM)")
    add_data_table(doc, ["Group", "Purpose", "Indicative permissions"],
              [["YAT-ICT-Admins", "YAT ICT day-to-day ops post-handover", "Read-only on infra; full CloudWatch/RDS/EC2 console; no IAM changes"],
               ["MTS-Consultants", "MTS during build + support", "Full admin during build; reduced post-handover"],
                              ["Finance-Auditors", "Finance / external auditors", "Read-only on logs, metrics, configs (financial-audit support)"]],
              widths=[3.5, 5.0, 7.5])
    add_bullet_list(doc, [
        "No long-lived access keys for humans; programmatic access via IAM roles only.",
        "The application servers reach the database and the logging service through an EC2 instance role, so no credentials are stored on a server.",
        "Configuration decision left to the implementer: the MTS-Consultants permission boundary during build vs after handover.",
    ])
    h3("4.4 Network topology")
    add_body_paragraph(doc, "VPC ledgerline-vpc, 10.20.0.0/16, with DNS hostnames and DNS resolution enabled. "
                 "This is an internal, staff-only service — there is no public internet ingress to the "
                 "application. Single-AZ subnets in ap-southeast-2a:")
    add_data_table(doc, ["Subnet", "CIDR", "Tier", "Internet-facing?"],
              [["ledgerline-public-a", "10.20.1.0/24", "Load balancer, NAT gateway", "Yes"],
               ["ledgerline-public-b", "10.20.2.0/24", "Load balancer second zone — empty", "Yes"],
               ["ledgerline-app-a", "10.20.11.0/24", "Application / Ledgerline EC2", "No"],
               ["ledgerline-data-a", "10.20.21.0/24", "Database (RDS for PostgreSQL)", "No"],
               ["ledgerline-data-b", "10.20.22.0/24", "Database subnet group second zone — empty", "No"]],
              widths=[4.0, 3.5, 6.0, 2.5])
    add_bullet_list(doc, [
        "Internet Gateway used only for NAT egress (Windows Update, vendor patches); no inbound internet path to the app.",
        "Route tables: public-egress → IGW; private-app → NAT; private-data → no internet route.",
        "Staff reach the service over the campus Site-to-Site VPN; AD authentication runs over the same private link.",
        "The follow-on HA design adds the corresponding -b subnets in ap-southeast-2b.",
    ])
    diagram_figure(doc,
                   "Figure 4.4 — Ledgerline baseline network topology. The workload runs in "
                   "ap-southeast-2a; the two subnets in ap-southeast-2b carry nothing.",
                   "network-accounting-baseline-singleaz.drawio.png")
    h3("4.5 Compute (EC2 + Auto Scaling)")
    add_bullet_list(doc, [
        "EC2: general-purpose burstable — t3.micro or t3.small, a C1 implementer decision; Amazon Linux 2023 AMI; placed in private-app-a (no public IP).",
        "EBS: gp3 root 80 GB + a gp3 data volume sized by the implementer (footprint + 12-month growth + headroom).",
        "Auto Scaling Group: min 1 / desired 1 / max 2 (baseline); target-tracking on CPU at 70%; ELB+EC2 health checks; 300 s cooldown.",
        "The workload is business-hours and idle overnight; the follow-on HA design adds cross-AZ capacity for resilience (not for load).",
    ])
    lab_note(doc, "Ledgerline's real finance workload — 15–25 typical and 45–55 concurrent users at "
                  "month-end close — would warrant an instance several sizes larger than either option "
                  "above. An AWS Academy Learner Lab caps what will launch, so the two offered are ones "
                  "that actually run there. Size between them on the reasoning you would use at full "
                  "scale; the reasoning is what matters, not the vCPU count.")
    h3("4.6 Load balancing (ALB)")
    add_bullet_list(doc, [
        "Internet-facing ALB, ledgerline-alb, spanning ledgerline-public-a and ledgerline-public-b — it will not create with only one zone; HTTP:80 listener forwarding to ledgerline-tg.",
        "Target group = the ASG instances; HTTP health check on the application health endpoint (30 s; 2 unhealthy → out of service).",
    ])
    lab_note(doc, "the listener is deliberately plain HTTP on port 80 rather than HTTPS. That is a "
                  "simplification, and a production front door for a finance system would terminate TLS "
                  "here. It is done this way so that you can paste the load balancer's address into a "
                  "browser and see the platform you built actually serving a page.")
    h3("4.7 Database (RDS for PostgreSQL)")
    add_bullet_list(doc, [
        "Amazon RDS for PostgreSQL (preserves the existing engine and data); version confirmed against Ledgerline at build time.",
        "General-purpose burstable instance class — db.t3.micro or db.t3.small, a C2 implementer decision; gp3 storage, 20 GB — sized to ~22 GB + ~5 GB/year growth.",
        "Multi-AZ DISABLED for the baseline (enabled in the HA design); storage encryption enabled (KMS).",
        "Placed via ledgerline-db-subnet-group, spanning ledgerline-data-a and ledgerline-data-b — a subnet group requires two zones; not publicly accessible. Automated backups sized to meet RPO ≤ 1 hour, 7-day retention.",
        "Schema and data migration are YAT ICT's responsibility, not MTS's — MTS provisions an empty instance.",
    ])
    lab_note(doc, "the same applies here as to the application tier: the ~22 GB financial data footprint "
                  "and the month-end close profile would in reality call for a database class and storage "
                  "allocation beyond either option above. Both are sized to what will deploy in a Learner Lab.")
    h3("4.8 Storage")
    add_data_table(doc, ["Resource", "Type", "Purpose"],
              [["EC2 root volume", "EBS gp3 8 GB", "OS and Ledgerline application install"],
               ["Database storage", "RDS gp3 20 GB, encrypted", "The PostgreSQL data files"]],
              widths=[4.0, 4.5, 7.5])
    add_bullet_list(doc, ["Object storage for scanned invoices, purchase orders and supporting documents is out of scope for this baseline and is a candidate for the follow-on phase, where the 7-year financial-records hold can be met with versioning and an archive lifecycle."])
    h3("4.9 Security")
    add_data_table(doc, ["Security group", "Inbound", "Outbound"],
              [["ledgerline-alb-sg", "HTTP:80 from 0.0.0.0/0", "default allow-all"],
               ["ledgerline-app-sg", "HTTP:80 from ledgerline-alb-sg", "default allow-all"],
               ["ledgerline-db-sg", "PostgreSQL:5432 from ledgerline-app-sg only", "default allow-all"]],
              widths=[3.0, 7.0, 6.0])
    add_bullet_list(doc, [
        "Each rule names a security group as its source rather than an address range, so it follows the instances however many the Auto Scaling group launches.",
        "Administrative access to instances is by Session Manager: no key pair, no open management port, no public IP address and no bastion host.",
        "Encryption at rest: EBS and RDS both enabled (KMS where customer-managed keys are warranted for financial data).",
        "Operates under the AWS Shared Responsibility Model — AWS secures the cloud; YAT/MTS secure the OS, application, IAM, data and access in the cloud.",
    ])
    h3("4.10 Monitoring (baseline)")
    add_body_paragraph(doc, "Standard CloudWatch metrics for EC2, RDS, ALB and Auto Scaling. Baseline alarms (HA-tuned "
                 "alarms come in the follow-on HA design):")
    add_data_table(doc, ["Alarm", "Metric", "Threshold"],
              [["ledgerline-unhealthy-hosts", "UnHealthyHostCount, per load balancer per target group",
                "Maximum ≥ 1 over 1 minute — any target has failed its health check"],
               ["ledgerline-db-storage-low", "FreeStorageSpace on the database instance",
                "Minimum below 15% of allocated storage, over two 5-minute periods"]],
              widths=[4.6, 5.4, 6.0])
    add_bullet_list(doc, ["Both alarms notify an SNS topic. The Auto Scaling policy tracks CPU and creates its own alarms, so no separate CPU alarm is required."])
    h3("4.11 Naming and tagging conventions")
    add_body_paragraph(doc, "Naming pattern ledgerline-<resource-type> — ledgerline-vpc, ledgerline-alb, ledgerline-tg, ledgerline-lt, ledgerline-db. Subnets are named for their tier and zone: ledgerline-public-a/-b, ledgerline-app-a, ledgerline-data-a/-b. Mandatory tags:")
    add_data_table(doc, ["Tag", "Value"],
              [["Project", "YAT-Accounting-Migration"], ["Environment", "Production"], ["Owner", "YAT-ICT"],
               ["ManagedBy", "MTS-Migration during build → YAT-ICT post-handover"],
               ["CostCentre", "YAT-Accounting"], ["DataClassification", "Financial / Confidential (PII)"]],
              widths=[5.0, 11.0])
    h3("4.12 Backup")
    add_data_table(doc, ["Resource", "Mechanism", "Retention"],
              [["RDS database", "Automated daily backups + transaction-log backups (point-in-time recovery)", "Sized to RPO ≤ 1 h; long-term financial copies retained ≥ 7 years"],
               ["EC2 EBS volumes", "None — the servers hold no state that is not in the database. They are rebuilt from the launch template, not restored", "n/a"]],
              widths=[4.0, 7.0, 5.0])
    add_bullet_list(doc, ["Cross-Region backup copies are out of scope for the baseline — addressed in the follow-on HA design."])
    h3("4.13 Recovery objectives — baseline state")
    add_body_paragraph(doc, "The baseline meets RPO ≤ 1 hour through RDS automated + transaction-log backups (point-in-"
                 "time recovery). It does NOT meet the RTO of 2 hours: recovery is by restore from backup, which does not reliably complete inside that window. The single AZ and "
                 "the single RDS instance remain known single points of failure: tolerable for a business-"
                 "hours service in the short term, but resilience against an AZ failure is the objective of the "
                 "follow-on HA design.")
    h3("4.14 Components requiring vertical scaling")
    add_body_paragraph(doc, "RDS instance-class changes require a modify-and-apply (a brief interruption, taken in the "
                 "maintenance window, outside business hours); EBS volumes support online resize with no downtime.")
    h3("4.15 Single points of failure removed")
    na(doc, "this baseline is single-AZ by design; the single RDS instance and the single AZ are known "
            "single points of failure, deliberately deferred to the follow-on HA design.")
    h3("4.16 Configuration decisions left to the implementer")
    add_body_paragraph(doc, "The design is opinionated where it matters and silent where the implementer must show "
                 "judgement. Each decision below is to be made and evidenced in the Deployment Report, "
                 "justified against the Ledgerline workload.")
    add_data_table(doc, ["#", "Decision", "Why left open"],
              [["C1", "Application-tier instance type (general-purpose burstable)",
                "Size against the Ledgerline workload (15–25 typical / 45–55 month-end peak)"],
               ["C2", "Database instance class and storage size (general-purpose burstable)",
                "Size against the database workload and its ~22 GB footprint"]],
              widths=[1.0, 7.0, 8.0])

    h1("5. Implementation Sequencing")
    na(doc, "this is a greenfield build in a new account, not a change to a running system; build order is "
            "at the implementer's discretion and there is no live service to sequence around or roll back.")

    h1("6. Simulation and Verification Plan")
    na(doc, "failure/resize simulation and availability verification are the subject of the follow-on HA "
            "design and its deployment report; the baseline build is verified functionally per the "
            "Deployment Report's testing section.")

    h1("7. Out of Scope")
    add_body_paragraph(doc, "Stated explicitly so the implementer knows what not to build (these are the deliberate "
                 "inputs to the follow-on HA design):")
    add_bullet_list(doc, [
        "Multi-AZ database; the ledgerline-app-b subnet; ASG capacity ≥ 2 across AZs.",
        "HA-tuned monitoring (cross-AZ latency, replica lag); cross-Region backup copies and DR runbook.",
        "Failure-simulation testing; automated availability reporting against the recovery objectives.",
        "And, out of MTS scope entirely: Ledgerline install, database migration, cutover, and change management (YAT ICT).",
    ])

    h1("8. References")
    add_bullet_list(doc, [
        "Accounting System Application Specification; Migration Requirements; Server Specifications; Operational Costing.",
        "Engagement Role Brief; Consultation Notes; ICT Environment Overview.",
        "Privacy Policy; User Access Policy; Security and Incident Response; Industry Standards Reference (AWS Well-Architected, ACSC Essential Eight).",
    ])

    h1("Document control")
    add_data_table(doc, ["Field", "Value"],
              [["Document version", "v1.0 — Approved for implementation"],
               ["Authored by", "MTS Senior Architecture Team, in consultation with YAT ICT"],
               ["Approved by", "Pat Lin (MTS Senior Consultant) · Sam Walker (YAT ICT Manager)"],
               ["Implemented by", "the Accounting System foundation-build Deployment Report"],
               ["Successor document", "YAT Accounting System Cloud Architecture — HA Design (follow-on phase)"]],
              widths=[5.0, 11.0])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "../diploma-cloud-cyber-website-s1/public/documents/YAT-Accounting-Baseline-Solution-Design.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
