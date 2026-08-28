#!/usr/bin/env python3
"""Build the supplied AT2 Baseline Solution Design (.docx) — in-world document.

The cloud architecture design a second MTS consultant produced while the student was on
leave, which the student implements in AT2. An in-world supplied artefact (NOT an
exemplar — no UoC tags): the branded Solution Design document, rendered from the existing
baseline-design content. Being a greenfield single-AZ design, the HA sections (§3 review,
§5 sequencing, §6 simulation, §4.15) are marked "Not applicable" — demonstrating the
convention from the design side, and modelling the Solution Design genre for AT3 (where
the student authors their own HA Solution Design).

Output to the website documents folder for printing to PDF, then wiring into the AT2+
intranet state in place of the current markdown page.

Usage:  python scripts/scenario/build_at2_baseline_solution_design.py [output.docx]
Default: ../diploma-cloud-cyber-website-s1/public/documents/YAT-LMS-Baseline-Solution-Design.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # content-repo scripts/ (brand + registry)  # noqa: E402
sys.path.insert(0, str(next(d / "scripts" for d in Path(__file__).resolve().parents if (d / "scripts" / "helpers" / "__init__.py").exists())))  # umbrella scripts/ (engine)  # noqa: E402
from helpers.docx_body_text import add_body_paragraph, add_bullet_list  # noqa: E402
from helpers.docx_tables import add_data_table  # noqa: E402
from helpers.docx_styling import add_field, paragraph_bottom_rule, set_cell_borders, shade_cell  # noqa: E402
from brand import ADDRESS, CREAM, GREY, TEAL, TERRACOTTA  # noqa: E402
from helpers.scenario_document import build_header_footer, configure_styles, wordmark  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def na(doc, reason):
    p = doc.add_paragraph()
    r = p.add_run(f"Not applicable — {reason}")
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    p.paragraph_format.space_after = Pt(6)
    return p


def lab_note(doc, text):
    """Where the design and the build environment differ, say so in one plain sentence.

    The design is the real design. Some of it cannot be built in an AWS Academy Learner Lab,
    and in one place (the HTTP front door) it is deliberately simplified so a student can see
    their own work running. Each of those points carries one of these notes, immediately under
    the thing it qualifies, rather than being quietly reworded to match the lab.
    """
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    lead = p.add_run("Building this in a lab ")
    lead.bold = True
    lead.font.size = Pt(9.5)
    lead.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    return p


DIAGRAM_DIR = Path(__file__).resolve().parents[3] / "diploma-cloud-cyber-website-s1" / "public" / "diagrams"


def diagram_figure(doc, caption, image_name, width_cm=16.0):
    """Place a diagram from the website's public/diagrams/ folder, captioned.

    The picture is placed by the generator, not pasted in afterwards, so it survives a rebuild.

    Which file: the `.drawio.png` exported from draw.io alongside the `.drawio.svg` the intranet
    embeds — so the document and the website show the same picture, including any tidying done in
    draw.io. (python-docx cannot embed SVG, which is why the document needs the raster copy.) The
    `.drawio` itself is generated from a committed spec under scripts/scenario/diagrams/.

    Fails loudly if the image is absent rather than emitting a document with a silent hole in it.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    image = DIAGRAM_DIR / image_name
    if not image.exists():
        raise FileNotFoundError(
            f"Diagram not found: {image}\n"
            f"Render it first:\n"
            f"  .claude/skills/draw-diagram/.venv/bin/python "
            f".claude/skills/draw-diagram/draw_diagram.py \\\n"
            f"      --spec diploma-cloud-cyber-content-s1/scripts/scenario/diagrams/"
            f"{Path(image_name).stem}.json \\\n"
            f"      --out {DIAGRAM_DIR / (Path(image_name).stem + '.drawio')} --png {image}")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image), width=Cm(width_cm))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.italic = True; cr.font.size = Pt(9); cr.font.color.rgb = RGBColor.from_string(GREY)
    return p


def build(path):
    doc = Document()
    configure_styles(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    build_header_footer(sec)

    # ---- COVER ----
    wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(GREY)
    paragraph_bottom_rule(doc.add_paragraph(), TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Solution Design")
    sub = doc.add_paragraph().add_run("YAT LMS Cloud Architecture — Baseline Design")
    sub.font.size = Pt(15); sub.bold = True; sub.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    doc.add_paragraph()
    cover = [
        ("Engagement", "YAT LMS Cloud Migration — Foundation Build"),
        ("Document type", "Technical design (Solution Design)"),
        ("Version", "v1.0 — Approved for implementation"),
        ("Authored by", "MTS Senior Architecture Team, in consultation with YAT ICT"),
        ("Approved by", "Pat Lin (MTS Senior Consultant) · Sam Walker (YAT ICT Manager)"),
        ("Implemented by", "the AT2 build run sheet (foundation build)"),
        ("Classification", "Internal — YAT ICT, and MTS personnel on signed MSA"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    for k, v in cover:
        cells = ct.add_row().cells
        set_cell_borders(cells[0]); set_cell_borders(cells[1]); shade_cell(cells[0], CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        cells[1].paragraphs[0].add_run(v).font.size = Pt(10)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    doc.add_paragraph("Contents", style="Heading 1")
    add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                 placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h3 = lambda t: doc.add_paragraph(t, style="Heading 3")

    h1("1. Purpose and Scope")
    add_body_paragraph(doc, "This document specifies the baseline AWS architecture MTS will implement as the first "
                 "build phase of the YAT LMS cloud migration. It translates the action plan approved at the "
                 "close of the Business Case engagement into a concrete, implementable design. The design "
                 "stops at “infrastructure ready for application deployment”: the EC2 instance is provisioned "
                 "with the OS, the RDS instance with an empty MySQL schema, and the ALB with placeholder "
                 "health checks — no application binaries or production data are placed by the MTS build.")
    h3("In scope of this design")
    add_bullet_list(doc, [
        "The production cloud foundation for the DOODLE LMS application.",
        "All compute, networking, identity, storage, database, autoscaling and monitoring needed to run the LMS as a multi-tier web workload in AWS.",
        "Single-region deployment in ap-southeast-2 (Sydney), with the workload in a single Availability Zone.",
    ])
    lab_note(doc, "the design calls for ap-southeast-2 (Sydney). If you are building this in an AWS "
                  "Academy Learner Lab, deploy to whichever region that environment gives you — "
                  "us-east-1 — and treat it as standing in for Sydney throughout. That is acceptable.")
    h3("Out of scope — deferred to the follow-on HA design phase")
    add_bullet_list(doc, [
        "High-availability hardening (Multi-AZ database, cross-AZ compute resilience, failure-simulation testing).",
        "Disaster recovery to a second AWS region; DR runbook and tabletop testing.",
        "Application re-platforming (the LMS remains Windows Server 2016 + DOODLE + MySQL).",
    ])
    h3("Out of MTS scope entirely — YAT ICT responsibility")
    add_bullet_list(doc, [
        "LMS application installation onto the EC2 instance(s) after handover.",
        "Database migration (extract from on-prem MySQL, load into RDS).",
        "Cutover — DNS switch, parallel running, decommissioning, user redirection.",
        "Organisational change management — CAB approvals, communications, training, post-cutover support.",
    ])

    h1("2. Design Inputs and Requirements")
    h3("2.1 Inputs")
    add_bullet_list(doc, [
        "LMS Application Specification — workload, SLAs, data footprint, integration points.",
        "LMS Cloud Migration Requirements — SLA, RPO/RTO targets.",
        "Engagement Role Brief — engagement scope, OS and application preservation.",
        "ICT Environment Overview and the On-Premises Network Diagram — the current state being migrated from.",
    ])
    h3("2.2 Requirements the design must meet")
    add_data_table(doc, ["Requirement", "Target / note"],
              [["Region / data residency", "ap-southeast-2 (Sydney); all student PII within Australia (Privacy Act, APP 8)"],
               ["Application stack", "Preserved — Windows Server · MySQL · DOODLE"],
               ["Concurrent users", "200–300 typical; 500–700 peak during assessment windows"],
               ["Data footprint", "~178 GB, growing ~25 GB/year"],
               ["Authentication", "Preserve AD-LDAP integration over a private network to campus AD"],
               ["High availability", "Out of baseline scope — the 99.9% target is deferred to the follow-on HA design"]],
              widths=[5.0, 11.0])

    h1("3. Review of Existing Architecture")
    na(doc, "this is a greenfield cloud foundation; there is no existing cloud baseline to review. The "
            "on-premises current state is documented separately in the ICT Environment Overview.")

    h1("4. Architecture Design")
    h3("4.1 Assumptions and constraints")
    add_data_table(doc, ["#", "Assumption / constraint", "Source"],
              [["A1", "Region must be ap-southeast-2 (Sydney) for data residency", "LMS App Spec; Privacy Policy"],
               ["A2", "Application stack preserved: Windows Server, MySQL, DOODLE", "Engagement Role Brief"],
               ["A3", "Data footprint ~178 GB, +~25 GB/year (attachments + submissions dominate)", "LMS App Spec"],
               ["A4", "Concurrent users 200–300 typical, 500–700 peak", "LMS App Spec"],
               ["A5", "All student PII must remain within Australia", "Privacy Act 1988, APP 8"],
               ["A6", "Preserve AD-LDAP integration over private network to campus AD", "LMS App Spec"],
               ["A7", "Baseline is NOT high-availability; 99.9% deferred to the HA design", "Scoping decision"]],
              widths=[1.0, 9.0, 6.0])
    h3("4.2 AWS account and region")
    add_body_paragraph(doc, "An AWS account scoped to the migration engagement, provisioned by YAT and shared with MTS "
                 "for the build. Region ap-southeast-2 (Sydney); no deployment outside Australian regions. "
                 "The workload sits in ap-southeast-2a; the follow-on HA design puts application capacity "
                 "into the second AZ.")
    h3("4.3 Identity and Access Management (IAM)")
    add_data_table(doc, ["Group", "Purpose", "Indicative permissions"],
              [["YAT-ICT-Admins", "YAT ICT day-to-day ops post-handover", "Read-only across the account; no IAM changes"],
               ["MTS-Consultants", "MTS during build + support", "Full admin during build; reduced post-handover"],
               ["Read-Only-Auditors", "Compliance / external auditors", "Read-only on logs, metrics, configs"]],
              widths=[3.5, 5.0, 7.5])
    add_bullet_list(doc, [
        "Individual users are placed in a group, and may also carry a policy of their own where their job needs "
        "something the group does not give everyone — Sam Walker (YAT ICT Manager) sits in YAT-ICT-Admins and "
        "carries CloudWatch access directly.",
        "No long-lived access keys for humans; programmatic access via IAM roles only.",
        "The application servers reach the database and the logging service through an EC2 instance role, so no "
        "credentials are stored on a server.",
        "Identity administration stays with whoever governs the account: the group that operates the platform can "
        "see everything and change nothing about who has access.",
    ])
    lab_note(doc, "an AWS Academy Learner Lab does not permit groups, users or roles to be created, and supplies a "
                  "ready-made role and instance profile instead. Build this section as far as the environment "
                  "allows, capture the refusal, and use the supplied role where this design calls for the "
                  "application's own.")
    h3("4.4 Network topology")
    add_body_paragraph(doc, "VPC 10.0.0.0/16 (room to expand), with DNS hostnames and DNS resolution enabled. The "
                 "workload runs in ap-southeast-2a. Two further subnets exist in ap-southeast-2b carrying "
                 "nothing: the load balancer and the database subnet group each refuse to be created "
                 "unless they are given subnets in two Availability Zones.")
    add_data_table(doc, ["Subnet", "CIDR", "Zone", "Tier", "Internet-facing?"],
              [["public-web-a", "10.0.1.0/24", "ap-southeast-2a", "Web / public load balancer", "Yes"],
               ["public-web-b", "10.0.2.0/24", "ap-southeast-2b", "Load balancer second zone — empty", "Yes"],
               ["private-app-a", "10.0.11.0/24", "ap-southeast-2a", "Application / LMS EC2", "No"],
               ["private-data-a", "10.0.21.0/24", "ap-southeast-2a", "Database (RDS)", "No"],
               ["private-data-b", "10.0.22.0/24", "ap-southeast-2b", "Database subnet group second zone — empty", "No"]],
              widths=[3.2, 2.6, 3.2, 4.6, 2.4])
    add_bullet_list(doc, [
        "Internet Gateway for public-subnet traffic; NAT Gateway in public-web-a for private-app outbound (Windows Update etc.).",
        "Route tables: public-rt carries both public subnets → IGW; private-app-rt carries private-app-a → NAT. "
        "The data subnets stay on the VPC's default route table, which has no internet route at all.",
        "Connectivity to campus AD: Site-to-Site VPN (baseline choice); Direct Connect deferred unless latency requires it.",
        "The follow-on HA design adds private-app-b, the one subnet the second zone is still missing.",
    ])
    diagram_figure(doc,
                   "Figure 4.4 — LMS baseline network topology. The workload runs in "
                   "ap-southeast-2a; the two subnets in ap-southeast-2b carry nothing.",
                   "network-at3-start-non-hardened.drawio.png")
    h3("4.5 Compute (EC2 + Auto Scaling)")
    add_bullet_list(doc, [
        "Every application server is built from a launch template, yat-lms-lt, so nothing is configured on a server by hand.",
        "Instance type: general-purpose burstable — t3.micro or t3.small, a C1 implementer decision sized against the concurrent-user load.",
        "Windows Server AMI; placed in private-app-a by the Auto Scaling group; no key pair and no public IP.",
        "EBS: gp3 root 30 GB, plus a gp3 8 GB data volume at device name xvdb for application data.",
        "The instance profile carries the EC2 instance role, which is how the servers reach the database and the logging service.",
        "Auto Scaling Group: min 1 / desired 1 / max 2, in private-app-a only; target-tracking on CPU at 70%; "
        "Elastic Load Balancing health checks with a 300 s grace period; 60 s warm-up.",
        "The follow-on HA design expands the ASG (min 2, across both AZs) and tunes scaling for assessment-window peaks.",
    ])
    lab_note(doc, "the LMS workload described in §2.2 — 200–300 typical and 500–700 peak concurrent users — would "
                  "in reality need an instance several sizes larger than either of these. An AWS Academy Learner "
                  "Lab caps what will launch, so the two options above are ones that actually run there. Size "
                  "between them on the same reasoning you would use at full scale; the reasoning is what matters "
                  "here, not the number of vCPUs.")
    lab_note(doc, "the design calls for a Windows Server build preserving the existing DOODLE stack. Use whichever "
                  "current Windows Server image your environment offers; the version is not what this build turns on.")
    h3("4.6 Load balancing (ALB)")
    add_bullet_list(doc, [
        "Internet-facing ALB, yat-lms-alb, spanning public-web-a and public-web-b — it will not create with only one zone.",
        "Listener: HTTP on port 80, forwarding to the LMS target group.",
        "Target group yat-lms-tg = the ASG instances, HTTP on port 80; health check HTTP on / "
        "(30 s interval; 2 unhealthy → out of service).",
        "The LMS hostname and its TLS certificate are issued by YAT ICT and are put in front of this load balancer "
        "at cutover, which is outside MTS scope.",
    ])
    lab_note(doc, "the listener is deliberately plain HTTP on port 80 rather than HTTPS. That is a simplification, "
                  "and a real production front door would terminate TLS here. It is done this way so that you can "
                  "paste the load balancer's address into a browser and see the platform you built actually "
                  "serving a page — which is worth more, while you are learning, than a certificate you cannot "
                  "issue in a lab.")
    h3("4.7 Database (RDS)")
    add_bullet_list(doc, [
        "Amazon RDS for MySQL (preserves the existing data/schema); engine version confirmed against DOODLE at build time.",
        "General-purpose burstable instance class — db.t3.micro or db.t3.small, a C2 implementer decision.",
        "Storage: gp3, 20 GB; storage encryption enabled.",
        "Multi-AZ DISABLED for the baseline — no standby is created (it is enabled in the HA design).",
        "Placed via the yat-lms-db-subnet-group, which spans private-data-a and private-data-b; not publicly "
        "accessible; 7-day automated backups.",
        "Schema and data migration are YAT ICT's responsibility, not MTS's — MTS provisions an empty instance.",
    ])
    lab_note(doc, "the same applies here as to the application tier: the ~178 GB footprint and the read-heavy LMS "
                  "workload would in reality call for a database class and a storage allocation well beyond "
                  "either option above. An AWS Academy Learner Lab caps both, so the class and the 20 GB are "
                  "sized to what will actually deploy there.")
    h3("4.8 Storage")
    add_data_table(doc, ["Resource", "Type", "Purpose"],
              [["EC2 root volume", "EBS gp3 30 GB", "OS and LMS application install"],
               ["EC2 data volume", "EBS gp3 8 GB (xvdb)", "LMS application data"],
               ["Database storage", "RDS gp3 20 GB, encrypted", "The MySQL data files"]],
              widths=[4.0, 4.5, 7.5])
    h3("4.9 Security")
    add_data_table(doc, ["Security group", "Inbound", "Outbound"],
              [["yat-lms-alb-sg", "HTTP:80 from 0.0.0.0/0", "default allow-all"],
               ["yat-lms-app-sg", "HTTP:80 from yat-lms-alb-sg", "default allow-all"],
               ["yat-lms-db-sg", "MySQL:3306 from yat-lms-app-sg only", "default allow-all"]],
              widths=[3.6, 7.4, 5.0])
    add_bullet_list(doc, [
        "Each rule names a security group as its source rather than an address range, so the rule follows the "
        "instances however many the Auto Scaling group launches and whatever addresses they get.",
        "Administrative access to instances is by Session Manager: no key pair, no open management port, no public "
        "IP address and no bastion host.",
        "Encryption at rest: EBS and RDS both enabled.",
        "Operates under the AWS Shared Responsibility Model — AWS secures the cloud; YAT/MTS secure the OS, application, IAM, data and access in the cloud.",
    ])
    h3("4.10 Monitoring (baseline)")
    add_body_paragraph(doc, "Standard CloudWatch metrics for EC2, RDS, ALB and Auto Scaling. Two baseline alarms — one "
                 "for the application tier, one for the database tier — establish availability monitoring "
                 "for the platform. The Auto Scaling policy tracks CPU and creates its own alarms, so no "
                 "separate CPU alarm is required. HA-tuned alarms come in the follow-on HA design:")
    add_data_table(doc, ["Alarm", "Metric", "Threshold"],
              [["yat-lms-unhealthy-hosts", "UnHealthyHostCount, per load balancer per target group",
                "Maximum ≥ 1 over 1 minute — any target has failed its health check"],
               ["yat-lms-db-storage-low", "FreeStorageSpace on the database instance",
                "Minimum below 15% of allocated storage, over two 5-minute periods"]],
              widths=[4.6, 5.4, 6.0])
    add_bullet_list(doc, [
        "Both alarms notify an SNS topic, yat-lms-alerts.",
        "The storage alarm's threshold is entered in bytes and is a percentage of the storage allocated at build "
        "time; its description records the allocated size, because the number alone does not tell anyone whether "
        "it is still correct after the storage is resized.",
    ])
    h3("4.11 Naming and tagging conventions")
    add_body_paragraph(doc, "Naming pattern yat-lms-<resource-type> — yat-lms-vpc, yat-lms-alb, yat-lms-tg, "
                 "yat-lms-lt, yat-lms-db. Subnets are named for their tier and zone: public-web-a, "
                 "public-web-b, private-app-a, private-data-a, private-data-b. Mandatory tags:")
    add_data_table(doc, ["Tag", "Value"],
              [["Project", "YAT-LMS-Migration"], ["Environment", "Production"], ["Owner", "YAT-ICT"],
               ["ManagedBy", "MTS-Migration during build → YAT-ICT post-handover"],
               ["CostCentre", "YAT-LMS"], ["DataClassification", "Confidential (PII) or Internal"]],
              widths=[5.0, 11.0])
    h3("4.12 Backup")
    add_data_table(doc, ["Resource", "Mechanism", "Retention"],
              [["RDS database", "Automated daily backups + transaction logs", "7 days"],
               ["EC2 EBS volumes", "None — the servers hold no state worth recovering. They are rebuilt "
                                   "from the launch template, not restored", "n/a"]],
              widths=[4.5, 7.5, 4.0])
    add_bullet_list(doc, ["Cross-Region backup copies are out of scope for the baseline — addressed in the follow-on HA design."])
    h3("4.13 Recovery objectives — baseline state")
    add_body_paragraph(doc, "The baseline's recovery posture is backup-based: RPO is up to the last automated backup "
                 "(daily, within a 7-day retention window) and RTO is restore-based, on a single instance "
                 "with no standby. The HA targets (99.9% availability, RPO ≤ 1 h, RTO ≤ 4 h) are "
                 "deliberately not met by this baseline and are the objective of the follow-on HA design.")
    h3("4.14 Components requiring vertical scaling")
    add_body_paragraph(doc, "RDS instance-class changes require a modify-and-apply (a brief interruption, taken in the "
                 "maintenance window); EBS volumes support online resize with no downtime.")
    h3("4.15 Single points of failure removed")
    na(doc, "this baseline is single-AZ by design; the single RDS instance and the single AZ are known "
            "single points of failure, deliberately deferred to the follow-on HA design.")
    h3("4.16 Configuration decisions left to the implementer")
    add_body_paragraph(doc, "The design is opinionated where it matters and silent where the implementer must show "
                 "judgement. Two sizing decisions are left open. For each, consider at least two "
                 "candidates, choose one, and record the choice and the reason in the build run sheet.")
    add_data_table(doc, ["#", "Decision", "Why left open"],
              [["C1", "Application-tier instance type (general-purpose family)",
                "Size against the LMS concurrent-user load"],
               ["C2", "Database instance class and storage size (general-purpose family)",
                "Size against the database workload and its data footprint"]],
              widths=[1.0, 7.0, 8.0])
    add_body_paragraph(doc, "Everything else is specified by this design: the Auto Scaling policy tracks CPU with a "
                 "70% scale-out target; the MTS-Consultants permission boundary is limited to the "
                 "engagement's tagged resources; administrative access to instances is by Session "
                 "Manager, with no bastion host; the database engine is MySQL; and the LMS hostname "
                 "and its certificate are issued by YAT ICT at cutover.")

    h1("5. Implementation Sequencing")
    na(doc, "this is a greenfield build in a new account, not a change to a running system; build order is "
            "at the implementer's discretion and there is no live service to sequence around or roll back.")

    h1("6. Simulation and Verification Plan")
    na(doc, "failure/resize simulation and availability verification are the subject of the follow-on HA "
            "design; the baseline build is verified functionally by the tests in the build run sheet — "
            "reaching the server, reaching the internet through it, reaching the database from it, "
            "reaching the application through the load balancer, and watching the group scale.")

    h1("7. Out of Scope")
    add_body_paragraph(doc, "Stated explicitly so the implementer knows what not to build (these are the deliberate "
                 "inputs to the follow-on HA design):")
    add_bullet_list(doc, [
        "Multi-AZ database; the private-app-b subnet; ASG capacity ≥ 2 across AZs.",
        "HA-tuned monitoring (cross-AZ latency, RDS replica lag); cross-Region backup copies and DR runbook.",
        "Failure-simulation testing; automated availability reporting against the 99.9% target.",
        "Object storage for course attachments and submissions, and the archiving policy over it.",
        "And, out of MTS scope entirely: LMS application install, data migration, cutover (including the LMS "
        "hostname and its TLS certificate), and change management (YAT ICT).",
    ])

    h1("8. References")
    add_bullet_list(doc, [
        "LMS Application Specification; LMS Cloud Migration Requirements; Engagement Role Brief.",
        "ICT Environment Overview; On-Premises Network Diagram.",
        "Privacy Policy; User Access Policy; Security and Incident Response; Industry Standards Reference (AWS Well-Architected, ACSC Essential Eight).",
    ])

    h1("Document control")
    add_data_table(doc, ["Field", "Value"],
              [["Document version", "v1.0 — Approved for implementation"],
               ["Authored by", "MTS Senior Architecture Team, in consultation with YAT ICT"],
               ["Approved by", "Pat Lin (MTS Senior Consultant) · Sam Walker (YAT ICT Manager)"],
               ["Implemented by", "the AT2 build run sheet (foundation build)"],
               ["Successor document", "YAT LMS Cloud Architecture — HA Design (follow-on phase)"]],
              widths=[5.0, 11.0])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "../diploma-cloud-cyber-website-s1/public/documents/YAT-LMS-Baseline-Solution-Design.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
