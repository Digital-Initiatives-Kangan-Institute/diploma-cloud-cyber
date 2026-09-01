#!/usr/bin/env python3
"""Build the S1-CL2 AT1 STUDENT instrument (.docx) by populating the Kangan template.

The blank half of the pair. Content, criteria and front matter all come from the assessor
builder, so the two cannot drift; this one renders the workbook in `student` mode and omits
the marking benchmark and its UoC traceability entirely.

The criteria table here carries the criterion STATEMENTS without their UoC-traceability lines —
the student is told what they are marked against, not which unit item it maps to.

Usage:  python scripts/s1_cl2/build_s1_cl2_at1_student.py [output.docx]
Default: S1-CL2-Cloud-Disaster-Recovery/assessments/AT1/AT1-Design-DR-Plan-Student.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))  # noqa: E402
import build_s1_cl2_at1_assessor as a  # noqa: E402  (single source for content + criteria)
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # noqa: E402
sys.path.insert(0, str(next(d / "scripts" for d in Path(__file__).resolve().parents
                            if (d / "scripts" / "helpers" / "__init__.py").exists())))  # noqa: E402
from helpers.docx_tables import clear_table_rows, find_instruction_row, set_cell_content  # noqa: E402

from docx import Document  # noqa: E402

TEMPLATE = str(Path(__file__).resolve().parents[2] / "kangan-templates" /
               "Project Assessment - Student.docx")


def build(path):
    doc = Document(TEMPLATE)

    t_details = doc.tables[0]
    set_cell_content(t_details.rows[2].cells[1], a.DETAILS["qualification"])
    set_cell_content(t_details.rows[3].cells[1], a.DETAILS["units"])
    set_cell_content(t_details.rows[4].cells[1], a.DETAILS["task_title"])
    set_cell_content(t_details.rows[5].cells[1], a.DETAILS["task_number"])

    t_instr = doc.tables[1]
    set_cell_content(find_instruction_row(t_instr, "Assessment overview"), a.STUDENT_OVERVIEW)
    set_cell_content(find_instruction_row(t_instr, "Task"), a.TASKS)
    set_cell_content(find_instruction_row(t_instr, "Time allowed"), a.TIME_ALLOWED)
    set_cell_content(find_instruction_row(t_instr, "Resources required"), a.STUDENT_RESOURCES)
    set_cell_content(find_instruction_row(t_instr, "Assessment criteria"), a.CRITERIA)

    # The criteria the student is marked against — statements only, no UoC tags.
    t_mark = doc.tables[2]
    clear_table_rows(t_mark, 2)
    for crit in a.CRITERIA_MAP:
        row = t_mark.add_row()
        set_cell_content(row.cells[0], f"{crit['code']} {crit['text']}")
        set_cell_content(row.cells[1], a.CHECK)

    def h1(t):
        return doc.add_paragraph(t, style="Heading 1")

    def h2(t):
        return a.part_a.R.heading2(doc, t)

    a.part_a.render_front_matter(doc, h1)
    a.part_a.render(doc, h1, h2, mode="student")
    a.part_b.render(doc, h1, h2, mode="student")
    a.part_c.render(doc, h1, h2, mode="student")

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = str(a.CLUSTER / "assessments" / "AT1" / "AT1-Design-DR-Plan-Student.docx")
    build(sys.argv[1] if len(sys.argv) > 1 else default)
