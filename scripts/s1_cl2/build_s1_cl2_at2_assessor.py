#!/usr/bin/env python3
"""Build the S1-CL2 AT2 ASSESSOR instrument (.docx) by populating the Kangan template.

AT2 = Cloud Microservice & IaC Implementation — the build phase of the Website Global Expansion
engagement. ICTCLD503 elements 3–4 (deploy and configure the microservice, monitoring, sign-off)
+ ICTCLD505 in full (infrastructure as code).

ONE DEFINITION, TWO INSTRUMENTS — content in s1_cl2_at2_run_sheet.py, rendered worked here and blank in
build_s1_cl2_at2_student.py. The marking guide's traceability lines and the reverse map are
derived from the workbook's own tags by helpers.workbook_instrument.

NO DEPLOYMENT REPORT TEMPLATE. ICTCLD505's assessment conditions (AC 1–9) are environment
conditions and name no document format, so `[ICTCLD505 PC 4.1]` "create user documentation
including cloud infrastructure as code templates" is met by task 20 of the worksheet, written
for an operator audience with the student's templates attached.

Usage:  python scripts/s1_cl2/build_s1_cl2_at2_assessor.py [output.docx]
"""
import sys
from pathlib import Path

from docx import Document  # noqa: E402

TEMPLATE = str(Path(__file__).resolve().parents[2] / "kangan-templates" /
               "Project Assessment - Assessor.docx")
CLUSTER = (Path(__file__).resolve().parents[2] / "S1-CL2-Cloud-Disaster-Recovery")

sys.path.insert(0, str(Path(__file__).resolve().parent))  # noqa: E402
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # noqa: E402
sys.path.insert(0, str(next(d / "scripts" for d in Path(__file__).resolve().parents
                            if (d / "scripts" / "helpers" / "__init__.py").exists())))  # noqa: E402
from helpers.docx_tables import (clear_table_rows, find_instruction_row,  # noqa: E402
                                 set_cell_content)
from helpers.instrument_layout import render_benchmark  # noqa: E402
from helpers.workbook_instrument import (benchmark_sections, build_reverse_map,  # noqa: E402
                                         collect_elements, load_uoc_text, marking_guide,
                                         unevidenced_items)
import s1_cl2_at2_run_sheet as content  # noqa: E402

CHECK = "☐ Yes  ☐ No"

DETAILS = {
    "qualification": "ICT50220 Diploma of Information Technology",
    "units": [
        "ICTCLD503 Implement web-scale cloud infrastructure",
        "ICTCLD505 Deploy infrastructure as code",
    ],
    "task_title": "AT2 — Cloud Microservice & IaC Implementation",
    "task_number": "2 of 2",
}

TIME_ALLOWED = [
    "A 12-hour build window across sessions, plus the handover session with the assessor.",
    "Deployments fail and get fixed — that is the work, not lost time. Time is indicative; a "
    "student who needs longer continues in the next session.",
]

OVERVIEW = [
    "Students are assessed on implementing the audit-log microservice approved in AT1, and on "
    "provisioning it as infrastructure as code. AT2 is the second of two assessment tasks in the "
    "S1-CL2 Cloud Disaster Recovery cluster and closes the Website Global Expansion engagement.",
    "The assessment is one guided workbook of twenty-two build tasks and three knowledge "
    "questions, submitted as a single document together with the student's own template files. "
    "The student analyses why infrastructure as code suits the engagement, operates a template "
    "written by someone else (including diagnosing why it fails), authors and deploys their own, "
    "tests the microservice end to end, parameterises it for a second environment, adds a resource "
    "by update, configures monitoring, writes the user documentation, tears the environment down, "
    "and obtains build sign-off.",
    "THE PROVIDED TEMPLATE DOES NOT DEPLOY AS SUPPLIED. Its AttributeDefinitions declares an "
    "attribute named `id` while its KeySchema uses `event_id`, so CloudFormation rejects it. This "
    "is deliberate and it is the evidence for [ICTCLD505 PC 2.6] and [ICTCLD505 KE 7]. Do not "
    "correct the supplied file, and do not tell a student what the fault is — task 7 exists so "
    "they diagnose it. A student who fixed it before deploying has skipped the evidence and "
    "should be asked to deploy it as supplied.",
    "WHAT IS BEING MARKED. Every settings table and worked answer contains values we chose so "
    "there is a concrete task. Each element carries two assessor-only lines: 'Evidences', naming "
    "the UoC items, and 'Satisfactory when', naming what has to be true for them to be met. Mark "
    "the second.",
    "WHAT IS SUBMITTED. The completed workbook, and the student's own infrastructure-as-code "
    "template files in whatever form they kept them. Task 20's user documentation is written in "
    "the workbook; the templates it refers to are attached.",
    "This is an open-book assessment. Students may use the YAT intranet, AWS documentation, course "
    "materials and external research (which must be cited). They may not use another student.",
    "Reasonable adjustment may include extending the build window, providing one-on-one verbal "
    "explanation of the supplied files, allowing alternative evidence formats where assistive "
    "technology requires it, or splitting the work across more sittings.",
    "Teacher/assessor support level: the assessor may clarify what a task is asking and explain "
    "the supplied code and contract, but must not diagnose the planted fault, write template "
    "syntax for the student, or debug their template for them.",
    "The assessment will not proceed if for any reason it is not safe to do so. You must advise "
    "the student of the reason for suspending the assessment, and what safety action should be "
    "taken. Advise the student of revised arrangements when it is safe to do so.",
    "There is a zero tolerance for plagiarism, cheating and collusion. Students will be expected "
    "to make a declaration that all work is their own prior to submission. Refer to the Training "
    "and Assessment Policy for further information.",
]

TASKS = [
    "The design and disaster recovery plan for the YAT website global expansion were approved in "
    "AT1. This is the implementation phase: the student builds the audit-log microservice that "
    "design calls for, as code, so YAT can reproduce it.",
    "Tasks 1–4 establish why the work is being done as code, what automation the platform "
    "supplies, what can go wrong with the approach, and which infrastructure-as-code service is "
    "selected.",
    "Tasks 5–10 operate a template the student did not write: read it and determine what it "
    "creates and depends on, deploy it, diagnose and fix the failure it produces, confirm the "
    "resource, update a parameter and observe that the update REPLACES the table rather than "
    "renaming it, and return the store to a known state.",
    "Tasks 11–19 are the student's own build: read the provided handler and contract to determine "
    "what the infrastructure must supply, author a template provisioning the API, queue, function "
    "and permissions, deploy it, confirm the wiring, test a valid event, a duplicate and an "
    "invalid one end to end, record the troubleshooting, parameterise for a second environment, "
    "add a dead-letter queue by update, and configure a metric and alarm.",
    "Tasks 20–22 close the engagement: user documentation written for the YAT operator who "
    "inherits the stack, teardown through the tooling, and build sign-off.",
    "The provided data-store template, the microservice handler and the webhook contract are "
    "supplied inline in the workbook — deliberately not as a lab-pack, because authoring and "
    "operating the deployment templates is the assessed skill.",
    "Lab environment: AWS Academy Learner Lab, deploying to us-east-1. The approved design places "
    "the audit store in India — [scenario: ap-south-1 | deploy: us-east-1].",
]

RESOURCES = [
    "Teacher/assessor supplied resources",
    "AWS Academy Learner Lab access — providing the cloud vendor service provider, the "
    "infrastructure-as-code service, the serverless environment, the managed data store, and "
    "console / CLI / SDK tooling",
    "The two provided files as downloads — datastore.yaml (carrying the deliberate fault) and "
    "handler.py. Both are reproduced in full in the workbook; the downloads save retyping",
    "Access to the YAT scenario site / intranet — the engagement requirements, the residency "
    "requirements, the industry-standards reference and the change-management procedure",
    "A person to role-play Sam Walker, YAT ICT Manager, for the build sign-off — normally the "
    "assessor",
    "The worked workbook later in this document — model answers, per-element UoC mapping, and the "
    "standard each element is marked against",
    "Student supplied resources",
    "Computer with web browser",
    "A code editor or IDE",
    "A screenshot tool",
]

STUDENT_OVERVIEW = [
    OVERVIEW[0], OVERVIEW[1],
    "WHAT YOU SUBMIT. The completed workbook, and your own infrastructure-as-code template files. "
    "The user documentation you write at task 20 goes in the workbook; the templates it refers to "
    "are attached.",
    "Things will fail. That is expected and it is part of what is assessed — task 16 exists for "
    "exactly that. Record what broke and what you did about it rather than quietly fixing it.",
    OVERVIEW[5],
    OVERVIEW[8], OVERVIEW[9],
]

STUDENT_RESOURCES = [
    "AWS Academy Learner Lab access — you build in us-east-1",
    "The two provided files as downloads — datastore.yaml and handler.py. Both are reproduced in "
    "full in this document; the downloads just save you retyping",
    "Access to the YAT scenario site / intranet — every document you need is linked from the task "
    "it belongs to",
    "Your assessor will role-play Sam Walker, YAT ICT Manager, for the build sign-off",
    "Computer with web browser",
    "A code editor or IDE",
    "A screenshot tool",
]

CRITERIA = [
    "To receive a Satisfactory outcome for this assessment the student must:",
    "Achieve Satisfactory on every criterion in the Marking Guide below",
    "Submit the completed workbook (.docx) with every task and question answered and every "
    "evidence box populated",
    "Submit their own infrastructure-as-code template files",
    "Obtain build sign-off at task 22",
]

CONDITIONS = [
    "These are conditions the assessor verifies as present before marking begins. They are not "
    "student-performance criteria — they are the conditions under which the assessment can validly "
    "be conducted.",
    "C1 Lab environment is accessible to the student throughout the assessment — AWS Academy "
    "Learner Lab — providing cloud vendor service provider access, a cloud managed database "
    "service, a serverless environment, an infrastructure-as-code service, console / CLI / SDK "
    "tooling, an SSH or RDP client, and internet and web browser access",
    "C2 The two provided files have been made available to the student — datastore.yaml with its "
    "deliberate fault intact, and handler.py",
    "C3 The YAT scenario site / intranet is accessible to the student throughout the assessment — "
    "supplying the engagement requirements, the residency requirements, the organisational "
    "procedures and the industry-standards reference",
    "C4 A code editor or IDE is available to the student for authoring templates",
    "C5 A person is available to role-play Sam Walker, YAT ICT Manager, for the build sign-off",
]

CRITERIA_MAP = [
    dict(code="M1", tasks=["1-4"],
         text="Infrastructure as code assessed and selected (tasks 1–4) — the student ties the "
              "benefits to this engagement's stated needs rather than listing them generically, "
              "identifies what the platform automates beyond provisioning, assesses what can go "
              "wrong including the destructive-update risk, and selects a service against real "
              "alternatives"),
    dict(code="M2", tasks=["5", "6"],
         text="Provided template reviewed and deployed (tasks 5–6) — the student determines what "
              "the template creates and what depends on it, including the export their own stack "
              "will consume, then deploys it with the service's own tooling and captures the "
              "failure rather than pre-emptively fixing it"),
    dict(code="M3", tasks=["7"],
         text="Fault diagnosed and fixed (task 7) — the student reads the error, identifies the "
              "mismatch between the attribute definition and the key schema, and records the "
              "diagnosis rather than a sequence of guesses"),
    dict(code="M4", tasks=["8", "9", "10"],
         text="Deployment confirmed, updated and returned to a known state (tasks 8–10) — the "
              "student inspects the created resource itself rather than the stack message, "
              "redeploys with a changed parameter and observes what the update did to the existing "
              "resource, and verifies the state the rest of the build needs"),
    dict(code="M5", tasks=["11"],
         text="Microservice code and contract reviewed (task 11) — the student extracts the "
              "infrastructure requirements from the provided code: the queue trigger, the "
              "environment variable, the write permission and the runtime"),
    dict(code="M6", tasks=["12", "13", "14"],
         text="Own template authored, deployed and confirmed (tasks 12–14) — the template "
              "provisions a set of related resources and connects to the provided store through "
              "its export rather than a hard-coded name; it deploys successfully; and the student "
              "confirms the wiring including the environment variable"),
    dict(code="M7", tasks=["15"],
         text="Microservice tested end to end (task 15) — a valid event is posted and the record "
              "shown in the store, a duplicate is shown not to create a second record, and an "
              "invalid event is shown rejected and logged"),
    dict(code="M8", tasks=["16"],
         text="Troubleshooting recorded (task 16) — at least two genuine problems with the "
              "symptom, the diagnosis and the fix, showing a method rather than trial and error"),
    dict(code="M9", tasks=["17", "18"],
         text="Parameterised and updated (tasks 17–18) — a second environment deployed from the "
              "same template by changing configuration only, and a new resource added by updating "
              "the existing stack rather than replacing it"),
    dict(code="M10", tasks=["19"],
         text="Monitoring configured (task 19) — a metric and a working alarm, with the student "
              "able to say what condition it detects and why that condition matters to YAT"),
    dict(code="M11", tasks=["20"],
         text="User documentation created (task 20) — documentation written for the operator who "
              "inherits the stack: what it is, deployment order, every parameter and what changing "
              "it does, how to update and remove, what the alarm means, and the templates "
              "themselves included"),
    dict(code="M12", tasks=["21"],
         text="Environment removed (task 21) — all stacks removed through the infrastructure-as-"
              "code tooling rather than by hand, and confirmed gone"),
    dict(code="M13", tasks=["22"],
         text="Build confirmed and signed off (task 22) — the student confirms the build against "
              "the approved design, seeks feedback, responds to it, and obtains recorded sign-off"),
    dict(code="M14", tasks=["Q1", "Q2", "Q3"],
         text="Knowledge (questions 1–3) — industry standards and standard products including "
              "storage technology, the testing and debugging techniques actually used, and how the "
              "templates would be managed and measured over their life"),
]

AC_CONDITIONS = {
    "ICTCLD503 AC 1": "C1", "ICTCLD503 AC 2": "C1", "ICTCLD503 AC 3": "C1",
    "ICTCLD503 AC 4": "C2", "ICTCLD503 AC 5": "C3", "ICTCLD503 AC 6": "C4",
    "ICTCLD503 AC 7": "C3", "ICTCLD503 AC 8": "C1", "ICTCLD503 AC 9": "C3",
    "ICTCLD505 AC 1": "C1", "ICTCLD505 AC 2": "C1", "ICTCLD505 AC 3": "C3",
    "ICTCLD505 AC 4": "C3", "ICTCLD505 AC 5": "C4", "ICTCLD505 AC 6": "C1",
    "ICTCLD505 AC 7": "C1", "ICTCLD505 AC 8": "C1",
}

EXPECTED = (
    [f"ICTCLD503 PC {n}" for n in "3.1 3.2 3.3 3.4 4.1 4.2 4.3".split()]
    + [f"ICTCLD503 PE {n}" for n in "3 4".split()]
    + [f"ICTCLD503 KE {n}" for n in "1 2 5".split()]
    + ["ICTCLD503 FS Writing"]
    + [f"ICTCLD505 PC {n}" for n in
       "1.1 1.2 1.3 1.4 2.1 2.2 2.3 2.4 2.5 2.6 3.1 3.2 3.3 3.4 3.5 3.6 3.7 4.1 4.2".split()]
    + [f"ICTCLD505 PE {n}" for n in "1 2 3 4".split()]
    + [f"ICTCLD505 KE {n}" for n in range(1, 12)]
    + [f"ICTCLD505 FS {s}" for s in
       ["Oral communication", "Problem solving", "Reading", "Self-management", "Writing"]]
)

STYLE = {"h1": "Heading 1", "h2": "Heading 2", "p": "Normal"}


def _elements():
    return collect_elements(content.TASKS, (content.QUESTIONS, "Q"))


# The mapping engine reads this: the same per-criterion tag lists the marking guide
# uses, inverted per unit to produce the Assessment Mapping documents.
BENCHMARK = benchmark_sections(CRITERIA_MAP, _elements(), "AT2 — Cloud Microservice & IaC Implementation")


def _assessor_body(elements):
    uoc = load_uoc_text(CLUSTER / "consolidated_uoc.md")
    return [
        ("h1", "Marking Benchmark — UoC traceability (reverse map)"),
        ("p", "Every UoC requirement AT2 claims to evidence, with the marking criterion that "
              "evidences it. Each workbook element also carries its own 'Evidences' line, where "
              "the work is done. Both this table and the criteria lines are generated from those "
              "element tags, so they cannot drift from the workbook."),
        ("p", "ICTCLD503 — Implement web-scale cloud infrastructure (AT2-evidenced items; "
              "elements 1–2 are evidenced in AT1)"),
        ("tbl", build_reverse_map(CRITERIA_MAP, elements, uoc, "ICTCLD503", AC_CONDITIONS)),
        ("p", "ICTCLD505 — Deploy infrastructure as code"),
        ("tbl", build_reverse_map(CRITERIA_MAP, elements, uoc, "ICTCLD505", AC_CONDITIONS)),
    ]


def render_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    for r, cells in enumerate(rows):
        for c, val in enumerate(cells):
            set_cell_content(t.rows[r].cells[c], val)
    doc.add_paragraph()


def delete_body_paragraph(doc, text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            p._element.getparent().remove(p._element)
            return


def fill_front_matter(doc, overview, resources, student=False):
    t_details = doc.tables[0]
    off = 1 if student else 0
    set_cell_content(t_details.rows[1 + off].cells[1], DETAILS["qualification"])
    set_cell_content(t_details.rows[2 + off].cells[1], DETAILS["units"])
    set_cell_content(t_details.rows[3 + off].cells[1], DETAILS["task_title"])
    set_cell_content(t_details.rows[4 + off].cells[1], DETAILS["task_number"])

    t_instr = doc.tables[1]
    set_cell_content(find_instruction_row(t_instr, "Assessment overview"), overview)
    set_cell_content(find_instruction_row(t_instr, "Task"), TASKS)
    set_cell_content(find_instruction_row(t_instr, "Time allowed"), TIME_ALLOWED)
    set_cell_content(find_instruction_row(t_instr, "Resources required"), resources)
    set_cell_content(find_instruction_row(t_instr, "Assessment criteria"), CRITERIA)
    if not student:
        set_cell_content(find_instruction_row(t_instr, "Location"), "")
        cond_row = t_instr.add_row()
        set_cell_content(cond_row.cells[0], "Assessment Conditions & Setup Requirements")
        for r in cond_row.cells[0].paragraphs[0].runs:
            r.bold = True
        set_cell_content(cond_row.cells[1], CONDITIONS)


def build(path, mode="assessor"):
    elements = _elements()
    gaps = unevidenced_items(CRITERIA_MAP, elements, None, EXPECTED)
    if gaps:
        raise SystemExit("No criterion evidences: " + ", ".join(gaps))

    tmpl = TEMPLATE if mode == "assessor" else TEMPLATE.replace("Assessor", "Student")
    doc = Document(tmpl)
    fill_front_matter(doc,
                      OVERVIEW if mode == "assessor" else STUDENT_OVERVIEW,
                      RESOURCES if mode == "assessor" else STUDENT_RESOURCES,
                      student=(mode != "assessor"))

    t_mark = doc.tables[2]
    clear_table_rows(t_mark, 2)
    rows = (marking_guide(CRITERIA_MAP, elements) if mode == "assessor"
            else [[f"{c['code']} {c['text']}"] for c in CRITERIA_MAP])
    for lines in rows:
        row = t_mark.add_row()
        set_cell_content(row.cells[0], lines)
        set_cell_content(row.cells[1], CHECK)

    delete_body_paragraph(doc, "Add or delete rows as required")
    delete_body_paragraph(doc, "If questioning or observation is incorporated into this assessment "
                               "task, you can incorporate a Practical Observation Checklist.")

    def h1(t):
        return doc.add_paragraph(t, style="Heading 1")

    def h2(t):
        return content.R.heading2(doc, t)

    content.render_front_matter(doc, h1)
    content.render_supplied(doc, h1, h2)
    content.render(doc, h1, h2, mode=mode)

    if mode == "assessor":
        render_benchmark(doc, _assessor_body(elements), render_table, STYLE)

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}  ({mode})")


if __name__ == "__main__":
    args = list(sys.argv[1:])
    mode = args.pop(0) if args and args[0] in ("assessor", "student") else "assessor"
    name = ("AT2-Microservice-IaC-Implementation-"
            + ("Assessor" if mode == "assessor" else "Student") + ".docx")
    default = str(CLUSTER / "assessments" / "AT2" / name)
    build(args[0] if args else default, mode)
