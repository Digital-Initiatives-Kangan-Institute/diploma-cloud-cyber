#!/usr/bin/env python3
"""Build the S1-CL2 AT1 ASSESSOR instrument (.docx) by populating the Kangan template.

An institutional compliance document, not a YAT-branded artefact: it loads the official Kangan
'Project Assessment - Assessor' template and fills it in, then renders the workbook worked.

AT1 = Cloud Expansion: Design & DR Plan — three parts, one continuous worksheet:
  Part A  Design      (ICTCLD503 elements 1–2 — web-scale + microservice)
  Part B  DR Plan     (ICTCLD501 elements 1–4)
  Part C  Approval    (ICTCLD501 element 5 — the design-approval gate)

ONE DEFINITION, TWO INSTRUMENTS. The content lives in at1_part_{a,b,c}_run_sheet.py and is
rendered worked here and blank in build_s1_cl2_at1_student.py. Nothing about the task text
exists twice.

THE MARKING APPARATUS IS DERIVED. The criteria below declare which tasks each one groups; the
UoC-traceability line on each criterion and the whole reverse-map table are computed from the
workbook's own `uoc` tags by helpers.workbook_instrument. Adding a tag to a task carries it into
the marking guide and the reverse map automatically, and a criterion that groups untagged tasks
raises at build time rather than shipping as a free-floating criterion.

WHY PART A AND PART C CARRY NO SEPARATE TEMPLATE. ICTCLD503's assessment conditions name no
document format, so `[ICTCLD503 PC 1.7]` / `[PC 2.4]` "document and justify" are met by the
worksheet, and `[ICTCLD501 PC 5.1]` requires a VERBAL walkthrough, not a deck. Part B is the
exception: `[ICTCLD501 AC 3]` names "reporting standards for documenting and communicating
disaster recovery plan", so the DR Plan is assembled into the YAT template at task 37 and both
the worksheet and the plan are submitted.

Usage:  python scripts/s1_cl2/build_s1_cl2_at1_assessor.py [output.docx]
Default: S1-CL2-Cloud-Disaster-Recovery/assessments/AT1/AT1-Design-DR-Plan-Assessor.docx
"""
import sys
from pathlib import Path

from docx import Document  # noqa: E402

TEMPLATE = str(Path(__file__).resolve().parents[2] / "kangan-templates" /
               "Project Assessment - Assessor.docx")
CLUSTER = (Path(__file__).resolve().parents[2] / "S1-CL2-Cloud-Disaster-Recovery")

sys.path.insert(0, str(Path(__file__).resolve().parent))  # noqa: E402
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # noqa: E402
sys.path.insert(0, str(next(d / "scripts" for d in Path(__file__).resolve().parents
                            if (d / "scripts" / "helpers" / "__init__.py").exists())))  # noqa: E402
from helpers.docx_tables import (clear_table_rows, find_instruction_row,  # noqa: E402
                                 set_cell_content)
from helpers.instrument_layout import render_benchmark  # noqa: E402
from helpers.workbook_instrument import (benchmark_sections, build_reverse_map,  # noqa: E402
                                         collect_elements, load_uoc_text, marking_guide,
                                         unevidenced_items)
import at1_part_a_run_sheet as part_a  # noqa: E402
import at1_part_b_run_sheet as part_b  # noqa: E402
import at1_part_c_run_sheet as part_c  # noqa: E402

CHECK = "☐ Yes  ☐ No"

# ---------------------------------------------------------------- Kangan front matter

DETAILS = {
    "qualification": "ICT50220 Diploma of Information Technology",
    "units": [
        "ICTCLD501 Develop cloud disaster recovery plans",
        "ICTCLD503 Implement web-scale cloud infrastructure",
    ],
    "task_title": "AT1 — Cloud Expansion: Design & DR Plan",
    "task_number": "1 of 2",
}

TIME_ALLOWED = [
    "Part A — Design: 6 hours",
    "Part B — Disaster Recovery Plan: 6 hours, plus the time to assemble the plan document",
    "Part C — Presentation and approval: a 20-minute session, scheduled with the assessor",
    "One continuous worksheet. Time is indicative — a student who needs longer continues in the "
    "next session rather than submitting incomplete work.",
]

OVERVIEW = [
    "Students are assessed on designing the cloud architecture for the YAT website's global "
    "expansion, planning its disaster recovery, and obtaining approval to proceed to "
    "implementation. AT1 is the first of two assessment tasks in the S1-CL2 Cloud Disaster "
    "Recovery cluster.",
    "The assessment is one guided workbook in three parts, submitted as a single document. Part A "
    "leads the student task by task through the web-scale design and the audit-log microservice. "
    "Part B does the same for the disaster recovery analysis and plan, and closes by assembling "
    "that work into the YAT Disaster Recovery Plan template. Part C is a live session in which the "
    "student walks the role-played YAT ICT Manager through both, responds to feedback, lodges the "
    "plan and obtains sign-off.",
    "WHAT IS BEING MARKED. Every settings table and worked answer in this document contains values "
    "we chose so that there is a concrete task to perform — the units' wording is deliberately "
    "general. Each element therefore carries two assessor-only lines: 'Evidences', naming the UoC "
    "items, and 'Satisfactory when', naming what has to be true for those items to be met. Mark "
    "the second. A student whose design differs from ours but who meets the stated standard is "
    "Satisfactory; a student who matches ours but misses the standard is not.",
    "WHAT IS SUBMITTED. The completed workbook, and the completed YAT Disaster Recovery Plan "
    "assembled at task 37. The DR Plan is the artefact marked for [ICTCLD501 PC 4.3]; the "
    "workbook is the evidence that the thinking behind it is the student's own. Where the two "
    "differ materially, ask the student about it before marking.",
    "The design in Part A is the student's own and is NOT corrected before Part B. A disaster "
    "recovery plan is written for a designed system, so Part B plans recovery for whatever the "
    "student designed — including its weaknesses. Correcting the design first would make Part B a "
    "plan for the assessor's architecture rather than the student's.",
    "This is an open-book assessment. Students may use the YAT intranet, AWS documentation, course "
    "materials and external research (which must be cited). They may not use another student.",
    "Reasonable adjustment may include extending the time for either written part, providing "
    "one-on-one verbal explanation of the supplied environment, allowing the Part C session to be "
    "held by video conference, or splitting the work across more sittings.",
    "Teacher/assessor support level: the assessor may clarify what a task is asking, the scenario "
    "context and the supplied environment, but must not identify design options for the student, "
    "supply recovery objectives, or confirm whether a design decision is correct.",
    "The assessment will not proceed if for any reason it is not safe to do so. You must advise "
    "the student of the reason for suspending the assessment, and what safety action should be "
    "taken. Advise the student of revised arrangements when it is safe to do so.",
    "There is a zero tolerance for plagiarism, cheating and collusion. Students will be expected "
    "to make a declaration that all work is their own prior to submission. Refer to the Training "
    "and Assessment Policy for further information.",
]

# The student copy takes the overview and resources that are addressed to the STUDENT. The
# assessor-facing paragraphs — how to mark, support level, reasonable adjustment — and the
# assessor-supplied resource list are omitted, along with any UoC tag. Gate 9 fails the build
# otherwise, which is the check working.
STUDENT_OVERVIEW = [
    OVERVIEW[0], OVERVIEW[1],
    "WHAT YOU SUBMIT. The completed workbook, and the completed YAT Disaster Recovery Plan you "
    "assemble at task 37. The plan is the document your recovery planning is marked on; the "
    "workbook is the evidence that the thinking behind it is yours. Keep the two consistent.",
    "Your Part A design is your own and is not corrected before you start Part B. You plan "
    "recovery for the system you designed, whatever you designed.",
    OVERVIEW[5],
    OVERVIEW[8], OVERVIEW[9],
]

STUDENT_RESOURCES = [
    "Access to the YAT scenario site / intranet — supplying the engagement requirements, the "
    "residency requirements, the infrastructure and application specifications, the deprecated "
    "on-premises DR Plan, the YAT policies and the industry-standards reference. Every document "
    "you need is linked from the task it belongs to",
    "The YAT Disaster Recovery Plan template, from the intranet's Templates section — you need it "
    "at task 37",
    "Your assessor will role-play Sam Walker, YAT ICT Manager, for the Part C session",
    "Computer with web browser",
    "Word-processing software (e.g. Microsoft Word or equivalent)",
]

TASKS = [
    "YAT College has entered an offshore partnership in India, and its public website is now the "
    "enrolment front door for that campus. The website already runs on AWS in Sydney, hardened to "
    "Multi-AZ high availability. It was never built for an audience on the other side of the "
    "Indian Ocean, and the India operation brings log-residency obligations the current "
    "environment cannot meet.",
    "The task has three parts that combine into a single submission:",
    "Part A — Design. Nineteen tasks leading the student through the scaling needs the design must "
    "meet, a review of the current architecture against them, the residency obligation as a design "
    "input, the services required, the design of each layer, the global-delivery and caching "
    "decisions, the checks that the design scales and that availability and security are "
    "maintained, an architecture diagram, and a written justification. Tasks 15–19 do the same for "
    "the audit-log microservice: its data transactions, its supporting services, its architecture, "
    "its interface contract, and a written justification.",
    "Part B — Disaster Recovery Plan. Eighteen tasks covering the recovery requirements, the "
    "existing arrangements, vendor provisions, recovery objectives per component, the data being "
    "protected, at least three major risk events, plan exclusions, recovery options and the "
    "recommended strategy, vendor protections and prioritisation, insurance, the non-technical "
    "recovery components, detection and alerting, the recovery steps with timings, the arithmetic "
    "showing the plan meets its objectives, and the standards the plan reflects. Task 18 assembles "
    "all of it into the YAT Disaster Recovery Plan template.",
    "Part C — Presentation and approval. The student walks the role-played YAT ICT Manager through "
    "both documents, answers questions, responds to feedback, lodges the plan per YAT's Records "
    "Management Policy, and obtains sign-off to proceed to implementation. The first two tasks of "
    "Part C are the student's own preparation and are not marked.",
    "MTS scope: cloud infrastructure only. The website content, the CMS and the application stack "
    "are out of scope, as is legal interpretation of the India obligations — the student designs "
    "to the compliance area's determination.",
]

RESOURCES = [
    "Teacher/assessor supplied resources",
    "Access to the YAT scenario site / intranet — supplying the Website Global Expansion "
    "requirements, the Data Residency & Sovereignty Requirements, the Website Infrastructure "
    "Specifications, the Website Specification, the deprecated on-premises DR Plan, the YAT "
    "policies and the industry-standards reference. Those documents are linked from the "
    "Instructions to Student below",
    "The YAT Disaster Recovery Plan template, available from the intranet's Templates section — "
    "required for Part B task 18",
    "A person to role-play Sam Walker, YAT ICT Manager, for the Part C presentation, feedback and "
    "sign-off — normally the assessor",
    "The worked workbook later in this document — model answers, per-element UoC mapping, and the "
    "standard each element is marked against",
    "Student supplied resources",
    "Computer with web browser",
    "Word-processing software (e.g. Microsoft Word or equivalent)",
]

CRITERIA = [
    "To receive a Satisfactory outcome for this assessment the student must:",
    "Achieve Satisfactory on every criterion in the Marking Guide below",
    "Submit the completed workbook (.docx) with every task and question answered",
    "Submit the completed YAT Disaster Recovery Plan assembled at task 37",
    "Attend the Part C session and obtain sign-off",
]

CONDITIONS = [
    "These are conditions the assessor verifies as present before marking begins. They are not "
    "student-performance criteria — they are the conditions under which the assessment can validly "
    "be conducted.",
    "C1 The YAT scenario site / intranet is accessible to the student throughout the assessment — "
    "supplying the engagement requirements, the residency determination, the infrastructure and "
    "application specifications, the organisational policies and the industry-standards reference",
    "C2 Cloud platform reference access is available for the student to research services and "
    "their capabilities — a cloud vendor service provider, its managed database and serverless "
    "documentation, an internet connection and a web browser. Part A and Part B are design and "
    "planning tasks: nothing is deployed, so no lab session is required",
    "C3 The YAT Disaster Recovery Plan template has been made available to the student before "
    "Part B task 18",
    "C4 A person is available to role-play Sam Walker, YAT ICT Manager, for the Part C "
    "presentation, feedback and sign-off",
]

# ---------------------------------------------------------------- the marking guide
# Each criterion declares the workbook tasks it groups. The UoC line and the reverse map are
# derived from those tasks' own tags — see helpers/workbook_instrument.py.

CRITERIA_MAP = [
    dict(code="A1", tasks=["1-3"],
         text="Scaling needs, current-state review and the residency input (tasks 1–3) — "
              "the student establishes what the design is held to from the supplied documents, "
              "reviews the current architecture against those needs rather than against general "
              "good practice, and reads the residency determination correctly, separating what "
              "must be held in India from what may remain in Australia"),
    dict(code="A2", tasks=["4"],
         text="Services identified (task 4) — the services named address the gaps the "
              "review found, each with a stated purpose"),
    dict(code="A3", tasks=["5-7"],
         text="Layer-by-layer design (tasks 5–7) — the student designs the network entry "
              "point, the compute tier and the data tier as CHANGES to a working system, keeping "
              "what already meets the need and saying so"),
    dict(code="A4", tasks=["8-9"],
         text="Global delivery and caching (tasks 8–9) — the student distinguishes what is "
              "cached from what must reach the origin, sets a policy per content type, addresses "
              "search discoverability, and chooses between edge and in-memory caching with reasons"),
    dict(code="A5", tasks=["10-12"],
         text="Design checks and review (tasks 10–12) — the student checks the design "
              "scales as utilisation increases and identifies where it stops, checks availability "
              "and security are maintained including the new dependencies introduced, and reviews "
              "the whole design back against the needs from task 1"),
    dict(code="A6", tasks=["13"],
         text="Web-scale architecture (task 13) — a diagram showing a multi-tier web "
              "application whose networking, compute and storage all scale, consistent with the "
              "design tables"),
    dict(code="A7", tasks=["14"],
         text="Web-scale justification (task 14) — a written justification tying each "
              "significant choice to a recorded requirement and naming the alternatives rejected"),
    dict(code="A8", tasks=["15-16"],
         text="Microservice identified and its services chosen (tasks 15–16) — the student "
              "identifies the access-event transaction and where it comes to rest, and selects "
              "services for receiving, decoupling, processing and storing, explaining the queue"),
    dict(code="A9", tasks=["17-18"],
         text="Microservice architecture and contract (tasks 17–18) — a diagram showing a "
              "decoupled flow across the region boundary, and an integration contract complete "
              "enough to build either side from, including duplicate handling"),
    dict(code="A10", tasks=["19"],
         text="Microservice justification (task 19) — a written justification of the "
              "separation, the components and the residency decisions, including an honest account "
              "of what an outage of the service would cost"),
    dict(code="A11", tasks=["Q1", "Q2", "Q3", "Q4"],
         text="Design knowledge (questions 1–4) — the student explains the four component "
              "choices, cohesion and coupling, the web-scaling principles applied, and how the "
              "design keeps the residency option open, all with reference to their own design"),
    dict(code="A12", tasks=["20-22"],
         text="Recovery requirements and current position (tasks 20–22) — the student "
              "identifies the recovery requirements with the business need behind each, determines "
              "what recovery arrangements already exist and what they do not cover, and identifies "
              "the vendor's commitments and where YAT's responsibility begins"),
    dict(code="A13", tasks=["23-27"],
         text="Impact analysis (tasks 23–27) — recovery objectives per component, the data "
              "being protected with its volume and sensitivity, at least three major risk events "
              "rated by a stated method, justified plan exclusions, and the analysis recorded "
              "according to YAT's policies"),
    dict(code="A14", tasks=["28-32"],
         text="Recovery strategy (tasks 28–32) — a genuine range of recovery options with "
              "realistic recovery times and costs, a recommendation aligned to the business "
              "requirement rather than to maximum protection, vendor protections mapped against "
              "the risks and prioritised, a reasoned position on insurance, and the non-technical "
              "components the plan needs to work"),
    dict(code="A15", tasks=["33-35"],
         text="The plan itself (tasks 33–35) — detection that does not depend on the failed "
              "region, sequenced recovery steps with owners and timings, and the arithmetic "
              "showing the plan achieves the objectives set in task 4"),
    dict(code="A16", tasks=["36"],
         text="Standards applied (task 36) — the information-security and continuity "
              "standards named and located in the student's own plan"),
    dict(code="A17", tasks=["37"],
         text="The Disaster Recovery Plan (task 37) — the completed plan document, "
              "containing at least three major risk events, assembled to the YAT reporting "
              "standard the template sets. THIS CRITERION IS MARKED ON THE SUBMITTED PLAN, not on "
              "the worksheet"),
    dict(code="A18", tasks=["Q5", "Q6", "Q7"],
         text="Recovery knowledge (questions 5–7) — the student explains what is "
              "distinctive about a public cloud-hosted website's risk environment, the method "
              "behind their risk ratings, and the recovery techniques available and why theirs "
              "fits"),
    dict(code="A19", tasks=["40"],
         text="Verbal walkthrough (task 40, observed) — the student conducts a verbal "
              "walkthrough of the design and the plan with the required person, explaining their "
              "own reasoning in appropriate industry language and answering questions on it"),
    dict(code="A20", tasks=["41"],
         text="Feedback sought and responded to (task 41) — the student actively sought "
              "feedback and responded to each item with a decision and a reason; a reasoned "
              "disagreement is a satisfactory response"),
    dict(code="A21", tasks=["42"],
         text="Lodgement (task 42) — the approved plan is lodged according to YAT's Records "
              "Management Policy, with location, classification and retrievability recorded"),
    dict(code="A22", tasks=["43"],
         text="Sign-off obtained (task 43) — sign-off recorded with a decision, a name and a "
              "date; approval with conditions is a satisfactory outcome where the conditions are "
              "recorded"),
]

# Assessment conditions are verified before marking rather than evidenced by a criterion.
AC_CONDITIONS = {
    "ICTCLD501 AC 1": "C1", "ICTCLD501 AC 2": "C1", "ICTCLD501 AC 3": "C3",
    "ICTCLD503 AC 1": "C2", "ICTCLD503 AC 2": "C2", "ICTCLD503 AC 3": "C2",
    "ICTCLD503 AC 5": "C1", "ICTCLD503 AC 6": "C2", "ICTCLD503 AC 7": "C1",
    "ICTCLD503 AC 8": "C2", "ICTCLD503 AC 9": "C1",
}

# What the assessment plan says AT1 must evidence — checked at build time.
EXPECTED = (
    [f"ICTCLD501 PC {n}" for n in
     "1.1 1.2 1.3 2.1 2.2 2.3 2.4 2.5 3.1 3.2 3.3 3.4 4.1 4.2 4.3 5.1 5.2 5.3 5.4".split()]
    + [f"ICTCLD501 PE {n}" for n in "1 2 3".split()]
    + [f"ICTCLD501 KE {n}" for n in "1 2 3 4 5 6".split()]
    + [f"ICTCLD501 FS {s}" for s in
       ["Oral communication", "Planning and organising", "Problem solving", "Reading",
        "Self-management"]]
    + [f"ICTCLD503 PC {n}" for n in "1.1 1.2 1.3 1.4 1.5 1.6 1.7 2.1 2.2 2.3 2.4".split()]
    + [f"ICTCLD503 PE {n}" for n in "1 2 5".split()]
    + [f"ICTCLD503 KE {n}" for n in "3 4 6".split()]
    + [f"ICTCLD503 FS {s}" for s in
       ["Problem solving", "Reading", "Self-management", "Writing"]]
)

STYLE = {"h1": "Heading 1", "h2": "Heading 2", "p": "Normal"}


def _elements():
    # Task numbering is continuous across all three parts (1–43, questions Q1–Q7), so the
    # marking guide can name tasks the way the student sees them and the traceability validator
    # can resolve "task 27" to exactly one element.
    return collect_elements(part_a.DESIGN, (part_a.QUESTIONS, "Q"),
                            part_b.PLAN, (part_b.QUESTIONS, "Q"),
                            part_c.APPROVAL)


# The mapping engine reads this: the same per-criterion tag lists the marking guide
# uses, inverted per unit to produce the Assessment Mapping documents.
BENCHMARK = benchmark_sections(CRITERIA_MAP, _elements(), "AT1 — Cloud Expansion: Design & DR Plan")


def _assessor_body(elements):
    uoc = load_uoc_text(CLUSTER / "consolidated_uoc.md")
    body = [
        ("h1", "Marking Benchmark — UoC traceability (reverse map)"),
        ("p", "This table closes the loop on bidirectional traceability: every UoC requirement "
              "AT1 claims to evidence is named below with the marking criterion that evidences "
              "it. Each workbook element also carries its own 'Evidences' line, where the work is "
              "actually done. Both this table and the criteria lines are generated from those "
              "element tags, so they cannot drift from the workbook."),
        ("p", "ICTCLD501 — Develop cloud disaster recovery plans"),
        ("tbl", build_reverse_map(CRITERIA_MAP, elements, uoc, "ICTCLD501", AC_CONDITIONS)),
        ("p", "ICTCLD503 — Implement web-scale cloud infrastructure (AT1-evidenced items; "
              "elements 3–4 are evidenced in AT2)"),
        ("tbl", build_reverse_map(CRITERIA_MAP, elements, uoc, "ICTCLD503", AC_CONDITIONS)),
    ]
    return body


def render_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    for r, cells in enumerate(rows):
        for c, val in enumerate(cells):
            set_cell_content(t.rows[r].cells[c], val)
    doc.add_paragraph()


def delete_body_paragraph(doc, text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            p._element.getparent().remove(p._element)
            return


def build(path):
    elements = _elements()
    gaps = unevidenced_items(CRITERIA_MAP, elements, None, EXPECTED)
    if gaps:
        raise SystemExit("No criterion evidences: " + ", ".join(gaps))

    doc = Document(TEMPLATE)

    t_details = doc.tables[0]
    set_cell_content(t_details.rows[1].cells[1], DETAILS["qualification"])
    set_cell_content(t_details.rows[2].cells[1], DETAILS["units"])
    set_cell_content(t_details.rows[3].cells[1], DETAILS["task_title"])
    set_cell_content(t_details.rows[4].cells[1], DETAILS["task_number"])

    t_instr = doc.tables[1]
    set_cell_content(find_instruction_row(t_instr, "Assessment overview"), OVERVIEW)
    set_cell_content(find_instruction_row(t_instr, "Task"), TASKS)
    set_cell_content(find_instruction_row(t_instr, "Time allowed"), TIME_ALLOWED)
    set_cell_content(find_instruction_row(t_instr, "Location"), "")
    set_cell_content(find_instruction_row(t_instr, "Resources required"), RESOURCES)
    set_cell_content(find_instruction_row(t_instr, "Assessment criteria"), CRITERIA)
    cond_row = t_instr.add_row()
    set_cell_content(cond_row.cells[0], "Assessment Conditions & Setup Requirements")
    for r in cond_row.cells[0].paragraphs[0].runs:
        r.bold = True
    set_cell_content(cond_row.cells[1], CONDITIONS)

    t_mark = doc.tables[2]
    clear_table_rows(t_mark, 2)
    for lines in marking_guide(CRITERIA_MAP, elements):
        row = t_mark.add_row()
        set_cell_content(row.cells[0], lines)
        set_cell_content(row.cells[1], CHECK)

    delete_body_paragraph(doc, "Add or delete rows as required")
    delete_body_paragraph(doc, "If questioning or observation is incorporated into this assessment "
                               "task, you can incorporate a Practical Observation Checklist.")

    def h1(t):
        return doc.add_paragraph(t, style="Heading 1")

    def h2(t):
        return part_a.R.heading2(doc, t)

    part_a.render_front_matter(doc, h1)
    part_a.render(doc, h1, h2, mode="assessor")
    part_b.render(doc, h1, h2, mode="assessor")
    part_c.render(doc, h1, h2, mode="assessor")

    render_benchmark(doc, _assessor_body(elements), render_table, STYLE)

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = str(CLUSTER / "assessments" / "AT1" / "AT1-Design-DR-Plan-Assessor.docx")
    build(sys.argv[1] if len(sys.argv) > 1 else default)
