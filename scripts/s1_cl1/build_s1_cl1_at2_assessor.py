#!/usr/bin/env python3
"""Build the S1-CL1 AT2 ASSESSOR instrument (.docx) by populating the Kangan template.

This is an institutional compliance document, NOT a YAT-branded artefact: it loads the
official Kangan 'Project Assessment - Assessor' template and fills it in, preserving the
Kangan structure and styles (Details, Teacher/Assessor instructions, Marking Guide,
Instructions to Student, Benchmark). It retro-fits a generator to the approved, hand-authored
CL1 AT2 instrument, reproducing it exactly at the content level. Mirrors the CL1 AT1 and
CL2 AT2 generators' mechanics.

EVERGREEN: unlike the committed source, this generator emits NO scenario-site URL. Where the
authored text named the intranet URL, only the URL (and its immediate " (…)"/" at …" wrapper)
is dropped; every other word is verbatim. There is no brand.WEBSITE_URL import.

AT2 = Cloud Foundation Build: YAT LMS Migration (ICTICT517 + ICTCLD401 + ICTCLD502) — a single
written Deployment Report (with a single evidence appendix) documenting the implementation of a
supplied AWS foundation architecture in the AWS Academy lab. There is no presentation/observation
event. The assessor instrument carries the task instructions, the Marking Guide, the Deployment
Report Benchmark (per-section guidance + KE model answers), and the UoC coverage reverse-map. The
student instrument (build_s1_cl1_at2_student) carries only the student-facing content.

Usage:  python scripts/s1_cl1/build_s1_cl1_at2_assessor.py [output.docx]
Default: S1-CL1-Cloud-Design-Build/assessments/AT2/AT2-Deployment-Assessor.docx
"""
import sys
from pathlib import Path

from docx import Document  # noqa: E402

TEMPLATE = str(Path(__file__).resolve().parents[2] / "kangan-templates" / "Project Assessment - Assessor.docx")

# Exemplar captures from a worked run of the assessment, placed into the evidence boxes so the
# assessor copy regenerates worked rather than blank. Committed alongside the generator: the
# document is a pure function of its sources, and pasting screenshots in by hand after every
# rebuild is exactly the manual step this removes.
EVIDENCE_DIR = Path(__file__).resolve().parents[2] / \
    "S1-CL1-Cloud-Design-Build" / "assessments" / "AT2" / "exemplar-evidence"

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # content-repo scripts/ (brand + registry)  # noqa: E402
sys.path.insert(0, str(next(d / "scripts" for d in Path(__file__).resolve().parents if (d / "scripts" / "helpers" / "__init__.py").exists())))  # umbrella scripts/ (engine)  # noqa: E402
from helpers.instrument_layout import render_benchmark, render_prose, set_cell_rich  # noqa: E402
import at2_run_sheet  # noqa: E402  (the run sheet — single source for both instruments)
from helpers.docx_tables import add_section_row, clear_table_rows, find_instruction_row, set_cell_content  # noqa: E402


# ---------- content ----------

# Base URL of the scenario site — the single place it is declared; each link appends its path.
SITE = "https://yat.timbaird.com"

CHECK = "☐ Yes  ☐ No"  # marking-guide Satisfactory? cell (matches the Kangan template)

DETAILS = {
    "qualification": "ICT50220 Diploma of Information Technology",
    "units": [
        "ICTICT517 Match ICT needs with the strategic direction of the organisation",
        "ICTCLD401 Configure cloud services",
        "ICTCLD502 Design and implement highly-available cloud infrastructure",
    ],
    "task_title": "AT2 — Cloud Foundation Build: YAT LMS Migration",
    "task_number": "2 of 3",
}

# NOTE (evergreen): the committed doc named the intranet URL here — " (https://www.placeholder.com.au)"
# after "YAT intranet". This generator omits the URL and its wrapper; all other words verbatim.
OVERVIEW = [
    'Students are assessed on the implementation of a supplied AWS cloud architecture for the YAT LMS migration, and the production of a Deployment Report documenting the build.',
    'The assessment is a single-task project delivered as one written artefact (the Deployment Report) with one populated appendix. Implementation occurs in the AWS Academy lab environment authorised for this cluster. There is no presentation or observation event — all assessment evidence is captured in the Deployment Report and its appendices.',
    'This is an open-book practical assessment. Students may use the YAT intranet, the AWS Pricing Calculator, AWS Academy lab environments, AWS documentation, course reference materials, and external research (which must be cited) throughout.',
    'AT2 is the second of three assessment tasks in the S1-CL1 Cloud Design and Build cluster. It builds on AT1 (Business Case + presentation event) and feeds into AT3 (HA hardening + project closure). The student remains in the same MTS consultant role across all three.',
    'Reasonable adjustment for this assessment may include extending the time allowed, varying the lab access arrangement (e.g. extended lab hours), allowing alternative screenshot-evidence formats where assistive technology requires it, or providing one-on-one verbal explanation of the supplied design where needed.',
    'Teacher/assessor support level: the teacher/assessor may clarify task requirements, scenario context, or the supplied design but must not guide students to specific configuration decisions or correct knowledge-evidence answers. The two Configuration Decisions left to the implementer (per the supplied design §4.16) must be made by the student.',
    'Submission: the completed Deployment Report (.docx) with the evidence appendix populated is submitted via the LMS.',
    'The assessment will not proceed if for any reason it is not safe to do so. You must advise the student of the reason for suspending the assessment, and what safety action should be taken. Advise the student of revised arrangements for the assessment when it is safe to do so.',
    'There is a zero tolerance for plagiarism, cheating and collusion. Students will be expected to make a declaration that all work is their own prior to submission. Refer to the Training and Assessment Policy for further information.',
]

# NOTE (evergreen): the committed doc named the intranet URL in the last line — " at https://www.placeholder.com.au".
TASKS = [
    "Following the YAT board's approval of the action plan in the student's AT1 Business Case engagement, the student took a period of planned annual leave (in-scenario narrative).",
    'During that time MTS Senior Architecture worked with YAT IT to translate the approved direction into a detailed technical design — the YAT LMS Cloud Architecture — Baseline Design — which is now the student\'s build specification.',
    'The student has returned to MTS to lead the foundation-build phase of the engagement. The task has two parts that combine into a single deliverable:',
    'Implement the supplied AWS architecture for the YAT LMS migration foundation build, in the AWS Academy lab environment authorised for this engagement.',
    "Produce a Deployment Report documenting the build, using the YAT-provided Deployment Report template. The report covers what was built, configuration decisions made where the supplied design left them open, testing and validation outcomes, an operational handover for YAT IT, written responses to six contextual Knowledge Evidence questions about the student's own build, and one appendix of evidence.",
    "The Deployment Report is the single submitted deliverable. All build evidence (screenshots, test results) is captured in the report's appendices — there is no separate portfolio submission and no presentation event.",
    'The architecture being implemented is intentionally non-HA at this stage. HA hardening is the next phase (AT3) and is explicitly out of scope for AT2.',
    "MTS scope: cloud infrastructure provisioning only. Per the LMS Migration Role Brief on the YAT intranet (§ Scope of the MTS consulting engagement), students must not perform LMS application installation, MySQL data migration, cutover activities, or organisational change management as part of AT2. Those are YAT in-house IT's responsibility in-scenario. The AT2 deliverable stops at infrastructure ready for application deployment. Assessors should not award credit for application-deployment work that falls outside this scope; equally, students who include such work in their report should be redirected to focus their evidence on the in-scope infrastructure deliverables.",
    'Note:',
    'All scenario materials, organisational policies, supplied design, report template, and previous-project examples for YAT College are available on the YAT intranet — students sign in to the intranet at the start of the engagement and refer to it throughout.',
]

# NOTE (evergreen): the committed doc named the intranet URL in the first item — " (https://www.placeholder.com.au)".
RESOURCES = [
    'Teacher/assessor supplied resources',
    'Access to the YAT intranet — supplying the LMS Application Specification and the Records Management Policy',
    'Both documents are linked in the Instructions to Student below',
    'AWS Academy lab access — Cloud Foundations [104469] + Cloud Architecting [172221]',
    'The worked run sheet later in this document — screenshot descriptions, model answers and per-task UoC mapping',
    'Student supplied resources',
    'Computer with web browser',
    'Word-processing software (e.g. Microsoft Word or equivalent)',
    'A screenshot tool',
]

CRITERIA = [
    'To receive a Satisfactory outcome for this assessment the student must:',
    'Achieve Satisfactory on every criterion in the Marking Guide below',
    'Submit a completed Deployment Report (.docx) using the YAT-supplied template, with every section you are asked to write and every appendix populated',
]

# NOTE (evergreen): C2 named the intranet URL — " at https://www.placeholder.com.au".
CONDITIONS = [
    'These are conditions the assessor verifies as present before marking begins. They are not student-performance criteria — they are the conditions under which the assessment can validly be conducted.',
    'C1 Lab environment is accessible to the student throughout the assessment — AWS Academy Cloud Foundations [104469] + AWS Academy Cloud Architecting [172221] — providing cloud vendor service provider access, cloud managed database service (RDS), IDE / console / CLI / SSH-RDP tooling, and internet/web browser access',
    '[ICTCLD401 AC 1] · [ICTCLD401 AC 2] · [ICTCLD401 AC 3] · [ICTCLD502 AC 1] · [ICTCLD502 AC 2] · [ICTCLD502 AC 4] · [ICTCLD502 AC 6] · [ICTCLD502 AC 7]',
    'C2 The YAT scenario site / intranet is accessible to the student throughout the assessment — supplying the LMS Application Specification and the organisational policies',
    '[ICTCLD401 AC 4] · [ICTCLD502 AC 3] · [ICTCLD502 AC 5] · [ICTCLD502 AC 8]',
]

# Marking Guide — single-part (A1-A13); each criterion is [statement, UoC-traceability line].
MARKING = [
    ['A1 Network foundation (tasks 2-7) — student builds the VPC; the five subnets, with the three workload subnets in one zone and the two second-zone subnets the load balancer and the database subnet group each require; the internet and NAT gateways; two route tables, the data subnets left on the VPC default which carries no internet route; and the three security groups, with the database reachable only from the application tier', '[ICTCLD401 PC 1.2] \u00b7 [ICTCLD401 PC 1.7] \u00b7 [ICTCLD401 PC 2.2] \u00b7 [ICTCLD401 PE 1] \u00b7 [ICTCLD502 PC 1.3]'],
    ['A2 Identity and access (tasks 1, 8-9) — student accesses the platform, creates the four IAM groups, enables MFA on an administrator, and creates the instance role so no credentials are stored on a server', '[ICTCLD401 PC 1.4] · [ICTCLD401 PC 1.5] · [ICTCLD401 PC 1.6] · [ICTCLD401 PC 2.1]'],
    ['A3 Compute and load balancing (tasks 10-13) — student creates the launch template with the instance role and both volumes, the target group and internet-facing load balancer, and the Auto Scaling group attached to that target group with a scaling policy', '[ICTCLD401 PC 2.3] · [ICTCLD401 PC 3.1] · [ICTCLD502 PC 4.1] · [ICTCLD401 PE 2]'],
    ['A4 Database (tasks 14-15) — student creates the subnet group spanning two zones and deploys the managed database: single-AZ, encrypted, not publicly accessible, and in the database security group', '[ICTCLD401 PC 2.4] · [ICTCLD401 PC 2.5] · [ICTCLD401 PE 2]'],
    ['A5 Monitoring (task 16) — student creates both baseline alarms and evidences their state', '[ICTCLD502 PC 4.3]'],
    ['A6 Configuration decisions (tasks 10 and 15) — for each of the two decisions the run sheet leaves open, student names the option not chosen, states the choice made, and justifies it against the YAT LMS workload. Naming the gap between what the environment permits and what the workload needs is a satisfactory answer', '[ICTCLD401 PC 1.1] · [ICTCLD401 PC 1.3]'],
    ['A7 Testing (T1-T5) — student connects to the application server, reaches the internet, the database and the load balancer from it, and demonstrates the Auto Scaling group scaling out and back in on its own. Each screenshot shows what its box asks for', '[ICTCLD401 PC 2.6] · [ICTCLD401 PC 3.2] · [ICTCLD401 PE 3] · [ICTCLD502 PC 4.2]'],
    ['A8 Knowledge questions (Q1-Q6) — student answers all six with reference to their own build, not generically, in clear written English', '[ICTCLD401 KE 5] · [ICTCLD401 KE 6] · [ICTCLD401 KE 7] · [ICTCLD401 KE 8] · [ICTCLD401 KE 9] · [ICTCLD401 KE 10] · [ICTCLD401 FS Reading] · [ICTCLD401 FS Writing]'],
    ['A9 Handover — student states where the completed run sheet was filed and which YAT policy required that location', '[ICTCLD401 PC 4.3]'],
]

# ---- Shared 'Instructions to Student' prose (single-sourced; the student builder imports these) ----
# The intro + Part 2/Tips blocks are IDENTICAL across the assessor and student copies; only the
# Part 1 block differs (student adds detail), so Part 1 is defined per-builder.
ASSESSOR_BODY = [
    ('h1', 'Marking Benchmark — UoC traceability (reverse map)'),
    ('p', 'This table closes the loop on bidirectional traceability: every UoC requirement AT2 claims to evidence is named below with the marking criterion that evidences it, and each run-sheet task carries its own Evidences line where the work is done. No UoC requirement claimed by AT2 is left without a criterion.'),
    ('p', 'ICTCLD401 — Configure cloud services (AT2-evidenced items)'),
    ('tbl', [
        ['UoC item', 'Evidenced by criterion(ia)'],
        ['[ICTCLD401 PC 1.1] Discuss and compare different cloud computing solutions, models and services', 'A6'],
        ['[ICTCLD401 PC 1.2] Identify impact of shared security responsibility models', 'A1'],
        ['[ICTCLD401 PC 1.3] Select best cloud computing solution and service', 'A6'],
        ['[ICTCLD401 PC 1.4] Access account on cloud platform', 'A2'],
        ['[ICTCLD401 PC 1.5] Identify user access protocols and policies', 'A2'],
        ['[ICTCLD401 PC 1.6] Configure access functions within cloud environment', 'A2'],
        ['[ICTCLD401 PC 1.7] Identify and assign security responsibilities', 'A1'],
        ['[ICTCLD401 PC 2.1] Create users and groups', 'A2'],
        ['[ICTCLD401 PC 2.2] Create virtual multi-tier network', 'A1'],
        ['[ICTCLD401 PC 2.3] Create virtual machine', 'A3'],
        ['[ICTCLD401 PC 2.4] Define, add and expand storage', 'A4'],
        ['[ICTCLD401 PC 2.5] Deploy a managed database', 'A4'],
        ['[ICTCLD401 PC 2.6] Test external network access', 'A7'],
        ['[ICTCLD401 PC 3.1] Configure and apply autoscaling', 'A3'],
        ['[ICTCLD401 PC 3.2] Test automatic scaling', 'A7'],
        ['[ICTCLD401 PC 4.3] Save and store user documentation per organisational policies', 'A9'],
        ['[ICTCLD401 PE 1] Build at least one simple virtual network capable of supporting a workload', 'A1'],
        ['[ICTCLD401 PE 2] Configure compute, storage, database and autoscaling resources within virtual network', 'A3 · A4'],
        ['[ICTCLD401 PE 3] Conduct simple tests to confirm access to resources', 'A7'],
        ['[ICTCLD401 KE 5] VM/networking/scaling features (VM sizing, load balancing, autoscaling, monitoring, storage backups, virtual networks/traffic routing)', 'A8 (Q1)'],
        ['[ICTCLD401 KE 6] Vertical vs horizontal scaling; VM vs physical; RDBMS/DW/NoSQL; self-hosted vs managed vs cloud-native DB; storage options (block/object/archive/network filesystems)', 'A8 (Q2)'],
        ['[ICTCLD401 KE 7] User, business and vendor responsibilities according to shared security responsibility models', 'A8 (Q3)'],
        ['[ICTCLD401 KE 8] User access protocols and policies according to organisation hierarchy and job function', 'A8 (Q4)'],
        ['[ICTCLD401 KE 9] Security policies, protocols and mechanisms as they relate to cloud (network traffic limits + security responsibilities per work function/user access)', 'A8 (Q5)'],
        ['[ICTCLD401 KE 10] Purpose of DNS for connecting remote servers when web browsing', 'A8 (Q6)'],
        ['[ICTCLD401 FS Reading]', 'A8'],
        ['[ICTCLD401 FS Writing]', 'A8'],
        ['[ICTCLD401 AC 1] Cloud vendor service provider', 'C1 (pre-condition)'],
        ['[ICTCLD401 AC 2] Cloud managed database service', 'C1 (pre-condition)'],
        ['[ICTCLD401 AC 3] Internet and web browser', 'C1 (pre-condition)'],
        ['[ICTCLD401 AC 4] Data to gather information from to determine output and user requirements', 'C2 (pre-condition)'],
    ]),
    ('p', '(Note: 401 PCs not listed above (PC 1.8, PC 4.1, PC 4.2) are evidenced in AT1. 401 KEs not listed (KE 1, 2, 3, 4, 11) are evidenced in AT1 Appendix 2.)'),
    ('p', 'ICTCLD502 — Design and implement highly-available cloud infrastructure (AT2-evidenced items only)'),
    ('tbl', [
        ['UoC item', 'Evidenced by criterion(ia)'],
        ['[ICTCLD502 PC 1.3] Identify level of shared security responsibility models according to business needs', 'A1'],
        ['[ICTCLD502 PC 4.1] Implement architecture design in cloud environment', 'A3'],
        ['[ICTCLD502 PC 4.2] Demonstrate connectivity between resources at all tiers', 'A7'],
        ['[ICTCLD502 PC 4.3] Monitor and measure availability of resources', 'A5'],
        ['[ICTCLD502 FS Reading]', 'A1'],
        ['[ICTCLD502 AC 1] Cloud vendor service provider', 'C1 (pre-condition)'],
        ['[ICTCLD502 AC 2] Cloud managed database service', 'C1 (pre-condition)'],
        ['[ICTCLD502 AC 3] Information and data sources required to design and implement cloud infrastructure', 'C2 (pre-condition)'],
        ['[ICTCLD502 AC 4] Integrated development environment (IDE)', 'C1 (pre-condition — AWS Console + CLI counts as IDE)'],
        ['[ICTCLD502 AC 5] Specific requirements + industry standards + organisational procedures + legislative requirements + business and functionality requirements', 'C2 (pre-condition — supplied via YAT intranet)'],
        ['[ICTCLD502 AC 6] Internet and web browser', 'C1 (pre-condition)'],
        ['[ICTCLD502 AC 7] Secure shell (SSH) or remote desktop protocol (RDP) client', 'C1 (pre-condition — RDP for Windows EC2 instances)'],
        ['[ICTCLD502 AC 8] Data to gather information from to determine output and user requirements', 'C2 (pre-condition)'],
    ]),
    ('p', '(Note: 502 PCs not listed above (PC 1.1, 1.2, 2.1–2.5, 3.1–3.5, 4.4–4.6, 5.1, 5.2, 5.3) are evidenced in AT1 (PC 1.2, 5.2) or AT3 (the rest). 502 KEs not listed (KE 1, 2, 3) are evidenced in AT1 Appendix 2; KEs 4–9 are evidenced in AT3.)'),
]

# ---------- build helpers ----------

STYLE = {"h1": "Heading 1", "h2": "Heading 2", "p": "Normal"}


def add_marking_row(table, lines):
    """A marking-guide criterion row: criterion (+ UoC line) in col 0, the check box in col 1."""
    row = table.add_row()
    set_cell_content(row.cells[0], lines)
    set_cell_content(row.cells[1], CHECK)


def delete_body_paragraph(doc, text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            p._element.getparent().remove(p._element)
            return


def render_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    for r, cells in enumerate(rows):
        for c, val in enumerate(cells):
            set_cell_content(t.rows[r].cells[c], val)
    doc.add_paragraph()


def build(path):
    doc = Document(TEMPLATE)

    # ---- Table 0: Details ----
    t_details = doc.tables[0]
    set_cell_content(t_details.rows[1].cells[1], DETAILS["qualification"])
    set_cell_content(t_details.rows[2].cells[1], DETAILS["units"])
    set_cell_content(t_details.rows[3].cells[1], DETAILS["task_title"])
    set_cell_content(t_details.rows[4].cells[1], DETAILS["task_number"])

    # ---- Table 1: Teacher/Assessor instructions ----
    # (Second attempt + Assessment retention rows keep the template's standard text.)
    t_instr = doc.tables[1]
    set_cell_content(find_instruction_row(t_instr, "Assessment overview"), OVERVIEW)
    set_cell_content(find_instruction_row(t_instr, "Task"), TASKS)
    set_cell_content(find_instruction_row(t_instr, "Time allowed"), "")
    set_cell_content(find_instruction_row(t_instr, "Location"), "")
    set_cell_content(find_instruction_row(t_instr, "Resources required"), RESOURCES)
    set_cell_content(find_instruction_row(t_instr, "Assessment criteria"), CRITERIA)
    # add a Conditions row at the end (matching the table style), as CL1/CL2 do
    cond_row = t_instr.add_row()
    set_cell_content(cond_row.cells[0], "Assessment Conditions & Setup Requirements")
    for r in cond_row.cells[0].paragraphs[0].runs:
        r.bold = True
    set_cell_content(cond_row.cells[1], CONDITIONS)

    # ---- Table 2: Marking Guide (single-part A1-A13; the 'Project ' heading + intro line stay) ----
    t_mark = doc.tables[2]
    clear_table_rows(t_mark, 2)  # keep 'Assessment criteria' + 'Criteria | Satisfactory?' header rows
    for lines in MARKING:
        add_marking_row(t_mark, lines)

    # ---- drop the template's trailing marking-guide boilerplate ----
    delete_body_paragraph(doc, "Add or delete rows as required")
    delete_body_paragraph(doc, "If questioning or observation is incorporated into this assessment task, "
                                "you can incorporate a Practical Observation Checklist.")

    # ---- Instructions to Student (shared prose; assessor Part 1 variant) ----
    at2_run_sheet.render_front_matter(doc, lambda t: doc.add_paragraph(t, style="Heading 1"))

    # ---- The run sheet, worked: exemplar screenshots, model answers, per-task UoC tags ----
    at2_run_sheet.render_run_sheet(
        doc,
        lambda t: doc.add_paragraph(t, style="Heading 1"),
        lambda t: doc.add_paragraph(t, style="Heading 2"),
        mode="assessor",
        evidence_dir=EVIDENCE_DIR)

    # ---- Reverse map ----
    render_benchmark(doc, ASSESSOR_BODY, render_table, STYLE)

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")
    at2_run_sheet.report_evidence(EVIDENCE_DIR)


if __name__ == "__main__":
    default = "S1-CL1-Cloud-Design-Build/assessments/AT2/AT2-Deployment-Assessor.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
