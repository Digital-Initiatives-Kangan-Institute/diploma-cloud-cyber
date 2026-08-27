#!/usr/bin/env python3
"""Build the S1-CL1 AT2 STUDENT instrument (.docx) by populating the Kangan template.

The student-facing half of AT2. Two sources feed it:

  the Kangan wrapper   this module's own second-person wording — OVERVIEW, TASKS, RESOURCES,
                       CRITERIA, RESULTS, SUBMIT and the A1-A9 MARKING criteria — poured into
                       the official 'Project Assessment - Student.docx' template. DETAILS and
                       the CHECK box are imported from build_s1_cl1_at2_assessor so the two
                       instruments cannot disagree about what task this is.
  the run sheet        at2_run_sheet.render_front_matter + render_run_sheet(mode="student"),
                       the same content module the assessor copy renders. Student mode emits
                       empty screenshot and response boxes, and no UoC tags.

Assessor-only material is not included: the marking benchmark, the UoC reverse-map, the
per-task 'Evidences:' lines, the terracotta screenshot descriptions and the teal model answers.
validate_student_instrument.py is the gate that proves none of it leaked.

Scenario-site links: the run sheet emits real intranet URLs, built from at2_run_sheet.SITE —
the required-resources list up front, and the 'Related resources' block under the knowledge
questions that carry one. One constant to change if the site ever moves. Note that this
module's own RESOURCES block names the same two documents WITHOUT a URL ('on the YAT
intranet'), so the wrapper and the run sheet currently follow different conventions.

Usage:  python scripts/s1_cl1/build_s1_cl1_at2_student.py [output.docx]
Default: S1-CL1-Cloud-Design-Build/assessments/AT2/AT2-Deployment-Student.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_s1_cl1_at2_assessor as a  # noqa: E402  (shared content — single source of truth)
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # content-repo scripts/ (brand + registry)  # noqa: E402
sys.path.insert(0, str(next(d / "scripts" for d in Path(__file__).resolve().parents if (d / "scripts" / "helpers" / "__init__.py").exists())))  # umbrella scripts/ (engine)  # noqa: E402
from helpers.instrument_layout import set_cell_rich  # noqa: E402
import at2_run_sheet  # noqa: E402  (the run sheet — single source for both instruments)
from helpers.docx_tables import clear_table_rows, find_instruction_row, set_cell_content  # noqa: E402

from docx import Document  # noqa: E402

TEMPLATE = str(Path(__file__).resolve().parents[2] / "kangan-templates" / "Project Assessment - Student.docx")


# ---------- student-facing content (second person; no UoC traceability) ----------

OVERVIEW = [
    ('You are assessed on building a cloud platform for the YAT LMS migration, following the build '
     'run sheet in the second half of this document, and on the evidence and answers you record in '
     'it as you go.', 'p'),
    ('This is an open-book practical assessment. You may use the YAT intranet, the AWS Academy lab '
     'environment, AWS documentation and course materials throughout.', 'p'),
    ('AT2 is the second of three assessment tasks in the S1-CL1 Cloud Design and Build cluster. It '
     'follows AT1 (the Business Case engagement) and feeds into AT3 (high-availability hardening). '
     'You continue in the same MTS consultant role across all three.', 'p'),
    ('Submission: this document, completed — every screenshot box filled and every question '
     'answered — submitted via the LMS.', 'p'),
    ('The assessment will not proceed if for any reason it is not safe to do so. The assessor must '
     'advise you of the reason for suspending the assessment, what safety action should be taken, '
     'and of revised arrangements when it is safe to resume.', 'p'),
    ('There is zero tolerance for plagiarism, cheating and collusion. You will be expected to make '
     'a declaration that all work is your own prior to submission. Refer to the Training and '
     'Assessment Policy for further information.', 'p'),
]

TASKS = [
    ("Following the YAT board's approval of your action plan in AT1, MTS Senior Architecture worked "
     "with YAT IT to turn that direction into a cloud architecture design, and an MTS "
     "implementation lead turned the design into a build run sheet for the engagement.", 'p'),
    ('You are the implementer. Your work has three parts, all recorded in this one document:', 'p'),
    ('Work through the build run sheet, creating each piece of the platform to the settings given, '
     'and paste the screenshot each task asks for into the box provided.', 'b'),
    ('Run the five tests and paste in what each one returned.', 'b'),
    ('Answer the six knowledge questions about the platform you built, then record where you filed '
     'the completed document.', 'b'),
    ('This phase is the foundation build only. The platform is deliberately single-zone at this '
     'stage — high-availability hardening is the next phase and is out of scope here.', 'p'),
]

RESOURCES = [
    ('Provided to you', 'h'),
    ('This document, containing the build run sheet you work through', 'b'),
    ('LMS Application Specification (on the YAT intranet) — the workload your two sizing decisions '
     'are justified against', 'b'),
    ('Records Management Policy (on the YAT intranet) — where a completed engagement document has '
     'to be filed', 'b'),
    ('Provided externally', 'h'),
    ('AWS Academy lab access', 'b'),
    ('You supply', 'h'),
    ('Computer with a web browser', 'b'),
    ('Word-processing software (Microsoft Word or equivalent)', 'b'),
    ('A screenshot tool', 'b'),
]

CRITERIA = [
    ('To receive a Satisfactory result for AT2 you must:', 'p'),
    ('Achieve Satisfactory on every criterion in the Assessment Criteria table (below)', 'b'),
    ('Complete every task in the run sheet, with the screenshot each one asks for', 'b'),
    ('Run all five tests and record what each returned', 'b'),
    ('Answer all six knowledge questions with reference to your own build', 'b'),
]

RESULTS = [
    ('If you are deemed not satisfactory for this assessment, you will be given one (1) more '
     'attempt at this assessment (or part thereof), or your teacher/assessor will negotiate a '
     'further assessment with you. The second attempt must be completed within 10 working days '
     'from the date your feedback is given.', 'p'),
]

SUBMIT = [
    ('Submit this completed document to the LMS by the due date. Everything is recorded in it — '
     'there are no separate files to submit.', 'p'),
]

# Marking criteria, second person, no UoC traceability.
MARKING = [
    'A1 - Network foundation (tasks 2-7) — you built the VPC, the five subnets across two zones, '
    'the internet and NAT gateways, the route tables, and the three security groups with the '
    'database reachable only from the application tier',
    'A2 - Identity and access (tasks 1, 8-9) — you accessed the platform, worked through the group '
    'and user creation and captured the result, and reviewed the role your servers use',
    'A3 - Compute and load balancing (tasks 10-13) — you created the launch template, the target '
    'group, the load balancer and the Auto Scaling group, and the group is attached to the target '
    'group',
    'A4 - Database (tasks 14-15) — you created the subnet group across two zones and deployed the '
    'database: single-zone, encrypted, not publicly accessible, and in the database security group',
    'A5 - Monitoring (task 16) — you created both alarms with the metrics and thresholds given',
    'A6 - Configuration decisions (tasks 10 and 15) — for each of the two decisions the run sheet '
    'leaves open, you named the option you did not choose, stated the one you did, and justified it '
    'against the YAT LMS workload',
    'A7 - Testing (tests 1-5) — you connected to the application server, reached the internet, the '
    'database and the load balancer, and demonstrated the Auto Scaling group scaling out and back '
    'in on its own. Each screenshot shows what its box asks for',
    'A8 - Knowledge questions (Q1-Q6) — you answered all six with reference to your own build, not '
    'in general terms, in clear written English',
    'A9 - Handover — you stated where you filed the completed document and which YAT policy '
    'required that location',
]

def build(path):
    doc = Document(TEMPLATE)

    # ---- Table 0: Details (student name / ID / assessor name+email left blank to fill) ----
    t_details = doc.tables[0]
    set_cell_content(t_details.rows[2].cells[1], a.DETAILS["qualification"])
    set_cell_content(t_details.rows[3].cells[1], a.DETAILS["units"])
    set_cell_content(t_details.rows[4].cells[1], a.DETAILS["task_title"])
    set_cell_content(t_details.rows[5].cells[1], a.DETAILS["task_number"])

    # ---- Table 1: Student instructions ----
    t_instr = doc.tables[1]
    set_cell_rich(find_instruction_row(t_instr, "Assessment overview"), OVERVIEW)
    set_cell_rich(find_instruction_row(t_instr, "Task"), TASKS)
    set_cell_content(find_instruction_row(t_instr, "Time allowed"), "")
    set_cell_content(find_instruction_row(t_instr, "Location"), "")
    set_cell_rich(find_instruction_row(t_instr, "Resources required"), RESOURCES)
    set_cell_rich(find_instruction_row(t_instr, "Assessment criteria"), CRITERIA)
    set_cell_rich(find_instruction_row(t_instr, "Results"), RESULTS)
    # 'Important information' is left as the template's standard text.
    # add a 'How To Submit' row at the end (matching the table style)
    submit_row = t_instr.add_row()
    set_cell_content(submit_row.cells[0], "How To Submit")
    for r in submit_row.cells[0].paragraphs[0].runs:
        r.bold = True
    set_cell_rich(submit_row.cells[1], SUBMIT)

    # ---- drop the template's marking-guide intro line (the authored student copy omits it) ----
    for p in doc.paragraphs:
        if p.text.strip() == "List the criteria the student will be assessed against for the project assessment.":
            p._element.getparent().remove(p._element)
            break

    # ---- Table 2: Assessment criteria (single-part A1-A13; no UoC/benchmark) ----
    t_mark = doc.tables[2]
    clear_table_rows(t_mark, 2)  # keep the two header rows
    for line in MARKING:
        row = t_mark.add_row()
        set_cell_content(row.cells[0], line)
        set_cell_content(row.cells[1], a.CHECK)

    # ---- drop the template's trailing marking-guide boilerplate ----
    for text in ("Add or delete rows as required",
                 "If questioning or observation is incorporated into this assessment task, "
                 "you can incorporate a Practical Observation Checklist."):
        for p in doc.paragraphs:
            if p.text.strip() == text:
                p._element.getparent().remove(p._element)
                break

    # ---- The run sheet, blank: no UoC tags, no model answers, empty evidence boxes ----
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h2 = lambda t: doc.add_paragraph(t, style="Heading 2")
    at2_run_sheet.render_front_matter(doc, h1)
    at2_run_sheet.render_run_sheet(doc, h1, h2, mode="student")

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL1-Cloud-Design-Build/assessments/AT2/AT2-Deployment-Student.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
