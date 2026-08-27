#!/usr/bin/env python3
"""Build the S1-CL1 AT3 PRACTICE HA run sheet (.docx).

A practice exercise, not an assessment: a branded YAT/MTS document with no Kangan wrapper, no
marking criteria and no UoC tags. Same shape as the AT3 assessment workbook, on the Ledgerline
accounting system rather than the LMS — and with the teaching the assessment withholds:
"Things to consider" leading questions on every Part A design task, and click-by-click steps
on every Part B build task.

It renders through the assessment's own renderer (at3_run_sheet.render) with its content lists
passed in, so the practice and the assessment cannot drift structurally even though every
value in them differs.

Usage:  python scripts/s1_cl1/build_s1_cl1_at3_practice.py [output.docx]
Default: S1-CL1-Cloud-Design-Build/delivery/practice/AT3-Practice-HA-Run-Sheet.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))  # noqa: E402
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # noqa: E402
sys.path.insert(0, str(next(d / "scripts" for d in Path(__file__).resolve().parents if (d / "scripts" / "helpers" / "__init__.py").exists())))  # noqa: E402

import at3_practice_run_sheet as content  # noqa: E402
import at3_run_sheet  # noqa: E402
from helpers.docx_styling import paragraph_bottom_rule  # noqa: E402
from helpers.scenario_document import build_header_footer, configure_styles, wordmark  # noqa: E402
from brand import ADDRESS, GREY, TEAL  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402

R = at3_run_sheet.R


def build(path):
    doc = Document()
    configure_styles(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    build_header_footer(sec)

    # ---- cover ----
    wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(GREY)
    paragraph_bottom_rule(doc.add_paragraph(), TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Practice — High Availability")
    sub = doc.add_paragraph().add_run("Ledgerline — design and build run sheet")
    sub.font.size = Pt(15); sub.italic = True; sub.font.color.rgb = RGBColor.from_string(GREY)

    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h2 = lambda t: R.heading2(doc, t)

    h1("The engagement")
    for para in content.SCENARIO:
        R.p(doc, para, after=8)
    h1("How to work through this")
    for para in content.INSTRUCTIONS:
        R.p(doc, para, after=8)

    at3_run_sheet.render(
        doc, h1, h2, mode="student",
        design=content.DESIGN, build=content.BUILD, tests=content.TESTS,
        closeout=content.CLOSEOUT, questions=[], reflections=[],
        current_arch=content.CURRENT_ARCH, notes=True)

    h1("Cleaning up your environment")
    for para in content.CLEANUP_INTRO:
        R.p(doc, para, after=8)
    for i, (title, detail) in enumerate(content.CLEANUP, 1):
        par = doc.add_paragraph()
        par.paragraph_format.space_after = Pt(2)
        lead = par.add_run(f"{i}.  "); lead.bold = True; lead.font.size = Pt(R.BODY_PT)
        head = par.add_run(title); head.bold = True; head.font.size = Pt(R.BODY_PT)
        R.p(doc, detail, size=9.5, indent=0.6, after=8)

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL1-Cloud-Design-Build/delivery/practice/AT3-Practice-HA-Run-Sheet.docx"
    build(sys.argv[1] if len(sys.argv) > 1 else default)
