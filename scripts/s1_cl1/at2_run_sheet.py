#!/usr/bin/env python3
"""The S1-CL1 AT2 build run sheet — content, and the renderer that puts it in a document.

ONE definition of the run sheet, rendered two ways:

  student   the blank sheet the student works through and fills in — the job, the settings,
            empty screenshot boxes, empty response boxes. No UoC tags.
  assessor  the worked copy: every screenshot box replaced by a terracotta description of
            what should be in it, every response box filled with a teal model answer, and
            each task tagged with the UoC items it evidences.

Source of truth for the task list, its settings and its coverage:
  S1-CL1-Cloud-Design-Build/assessments/AT2/at2-build-tasks.md

Every value specified here is one the lab environment will accept. Sizes are the smallest
that do the job — the previous design named m6i.large / db.m6i.large, which students could
not launch.
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # noqa: E402
sys.path.insert(0, str(next(d / "scripts" for d in Path(__file__).resolve().parents if (d / "scripts" / "helpers" / "__init__.py").exists())))  # noqa: E402
from helpers.docx_styling import set_cell_borders, shade_cell  # noqa: E402
from helpers.instrument_layout import add_hyperlink  # noqa: E402
from brand import CREAM, GREY, TEAL, TERRACOTTA  # noqa: E402

from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402

BODY_PT = 10.5
IMAGE_CM = 16.0         # exemplar captures, sized to sit inside the 16.6 cm evidence box
UOC = "6B6660"          # muted — the assessor-only traceability line
CAPTURE = TERRACOTTA    # exemplar: what the screenshot should have shown
MODEL = TEAL            # exemplar: the model written answer

REGION_NOTE = ("This engagement's production region is ap-southeast-2 (Sydney). The build "
               "environment for this work is us-east-1 — build there, and treat it as standing in "
               "for Sydney throughout.")

# ---------------------------------------------------------------- front matter
# Rendered at body size in the document's own body style — deliberately NOT the shared
# 'Assessor text' style, which is dark blue and would make the top of the instrument look
# like a different document from the run sheet below it.

SITE = "https://yat.timbaird.com"

SCENARIO = [
    "YAT College is migrating their Learning Management System from on-premises to AWS. You are an "
    "MTS Consultant on this engagement, reporting to Pat Lin (MTS Senior Consultant). Sam Walker "
    "(YAT IT Manager) is your primary YAT-side stakeholder.",
    "At the end of the AT1 engagement the YAT board approved your action plan. MTS Senior "
    "Architecture then worked with YAT IT to turn that direction into a cloud architecture design, "
    "and an MTS implementation lead turned the design into the build run sheet that forms the "
    "second half of this document.",
    "Your job is to work through that run sheet and stand up the infrastructure. You are the "
    "implementer: the design decisions have been made, and the run sheet tells you what to build "
    "and what to build it to.",
]

RESOURCES = [
    ("LMS Application Specification — the workload your two sizing decisions are justified against",
     f"{SITE}/intranet/s1-cl1-at2/ict/lms-application-spec"),
    ("Records Management Policy — where a completed engagement document has to be filed",
     f"{SITE}/intranet/s1-cl1-at2/policies/records-management"),
]

INSTRUCTIONS = [
    "Work the tasks in order. Each one tells you what to build and the settings to build it to; "
    "finding your way around the console is left to you.",
    "Take each screenshot as you finish the task, not at the end. Recreating a screen after you "
    "have moved on is painful and sometimes impossible — this is the single most common way "
    "students lose marks here.",
    "Two tasks ask you to choose between two options. Name the one you did not pick as well as the "
    "one you did — a choice with no alternative beside it is not a comparison.",
    "Run the tests yourself and screenshot what actually happened. If a test fails, fix it, run it "
    "again, and say what you changed.",
    "Answer the knowledge questions about your own build. Generic answers about cloud computing "
    "will not pass.",
]

# ---------------------------------------------------------------- build tasks
# job      — what to do, in one sentence
# settings — (label, value) pairs: what it must be built to
# capture  — what the screenshot must show
# uoc      — items this task evidences
# decision — optional: (prompt, model answer) rendered as an extra response box

TASKS = [
    dict(n=1, title="Sign in and confirm your working region",
         job="Sign in to the AWS console for this engagement and set your region.",
         settings=[("Region", "us-east-1")],
         capture="the console with the region selector visible, showing the correct region.",
         uoc=["ICTCLD401 PC 1.4"]),

    dict(n=2, title="Create the VPC",
         job="Create the virtual private cloud all of this engagement's infrastructure sits in.",
         settings=[("Name", "yat-lms-vpc"), ("IPv4 CIDR", "10.0.0.0/16"),
                   ("DNS hostnames", "Enabled"), ("DNS resolution", "Enabled")],
         evidence_note="No screenshot required. Creating the VPC is not an assessable action on "
                       "its own — it is the prerequisite for everything that follows, all of "
                       "which is built inside it.",
         uoc=[]),

    dict(n=3, title="Create the subnets",
         job="Create five subnets, using the names, address ranges and zones below. Set the zone "
             "yourself on every one — the console picks a default, and if two subnets that are "
             "meant to be in different zones end up in the same one, the load balancer and the "
             "database will both refuse to create later.",
         settings=[("public-web-a", "10.0.1.0/24  \u2014  zone  us-east-1a"),
                   ("public-web-b", "10.0.2.0/24  \u2014  zone  us-east-1b"),
                   ("private-app-a", "10.0.11.0/24  \u2014  zone  us-east-1a"),
                   ("private-data-a", "10.0.21.0/24  \u2014  zone  us-east-1a"),
                   ("private-data-b", "10.0.22.0/24  \u2014  zone  us-east-1b"),
                   ("Why five", "three carry your workload, all in us-east-1a. The two in "
                                "us-east-1b exist only because the load balancer and the database "
                                "subnet group each refuse to be created unless you give them "
                                "subnets in two zones. Nothing is placed in those two")],
         capture="the subnet list filtered to your VPC, showing all five with their address ranges "
                 "and, in particular, the Availability Zone column.",
         uoc=["ICTCLD401 PC 2.2", "ICTCLD401 PE 1"]),

    dict(n=4, title="Create the internet gateway",
         job="Create the gateway that gives the public subnet a path to the internet, and attach it "
             "to your VPC.",
         settings=[("Name", "yat-lms-igw"),
                   ("Attach to", "yat-lms-vpc")],
         capture="the internet gateway showing State: Attached, and the VPC it is attached to.",
         uoc=["ICTCLD401 PC 2.2", "ICTCLD401 PE 1"]),

    dict(n=5, title="Create the NAT gateway",
         job="Create a NAT gateway, using the settings below. "
             "It lets the application servers reach out to the internet for updates and external services without anything on the internet being able to reach them. It sits in the public subnet and takes a few minutes to become available.",
         settings=[("Name", "yat-lms-nat"),
                   ("Subnet", "public-web-a — the NAT gateway goes in the public subnet, not the "
                              "one it serves"),
                   ("Connectivity type", "Public"),
                   ("Elastic IP", "allocate a new one here"),
                   ("Before moving on", "wait until its state reads Available — you cannot route to "
                                        "it until then")],
         capture="the NAT gateway showing State: Available, its subnet and its Elastic IP.",
         uoc=["ICTCLD401 PC 2.2", "ICTCLD401 PE 1"]),

    dict(n=6, title="Create the route tables",
         job="Create two route tables, using the settings below. "
             "The data subnets are deliberately left on the default route table the VPC already has, which carries no internet route at all — which is exactly what a database tier should have.",
         settings=[("public-rt", "associated with public-web-a AND public-web-b; 0.0.0.0/0 → the "
                                 "internet gateway"),
                   ("private-app-rt", "associated with private-app-a; 0.0.0.0/0 → the NAT gateway"),
                   ("The data subnets", "leave them alone — no route table of your own")],
         capture="both route tables, showing their routes and their subnet associations.",
         uoc=["ICTCLD401 PC 2.2", "ICTCLD401 PE 1"]),

    dict(n=7, title="Create the three security groups",
         job="Create three security groups, using the names and rules below. "
             "Create all three as empty shells first, then go back and add the rules — each one names the others as a source, which you cannot do until they exist. AWS will not accept a name starting with sg- : that prefix is reserved for the IDs it generates.",
         settings=[("yat-lms-alb-sg", "inbound HTTP 80 from 0.0.0.0/0"),
                   ("yat-lms-app-sg", "inbound HTTP 80 from yat-lms-alb-sg"),
                   ("yat-lms-db-sg", "inbound MySQL 3306 from yat-lms-app-sg only"),
                   ("Outbound", "leave the default allow-all on all three")],
         capture="all three security groups with their inbound rules expanded.",
         uoc=["ICTCLD401 PC 1.7", "ICTCLD502 PC 1.3"]),

    dict(n=8, title="Configure the access model",
         job="Create the group that will operate the platform after handover, and a user in it. "
             "Complete each form with the values below, submit it, and capture the result.",
         note="the lab environment does not permit groups or users to be created. This is a known "
              "constraint of the AWS educational environment, not a fault in your work — you will "
              "be refused, and the refusal is what you capture.",
         settings=[("Group name", "YAT-ICT-Admins"),
                   ("Policy on the group", "ReadOnlyAccess  \u2014 filter for it and pick the policy named "
                                           "exactly that, not one of the service-specific ones"),
                   ("User name", "swalker  (Sam Walker, YAT IT Manager)"),
                   ("Policy on the user", "CloudWatchFullAccess"),
                   ("Expected result", "both submissions are refused with a permissions error")],
         captures=["the completed group form — YAT-ICT-Admins with ReadOnlyAccess attached — and "
                   "the error returned.",
                   "the completed user form — swalker with CloudWatchFullAccess attached — and the "
                   "error returned."],
         uoc=["ICTCLD401 PC 1.5", "ICTCLD401 PC 2.1"],
         decision=("Why does this group get read access to everything but no ability to change "
                   "identity settings, when it is the group that operates the whole platform?",
                   "Identity is what decides who can grant themselves more access. A group that can "
                   "rewrite its own permissions has no boundary at all — it could give itself "
                   "anything. So YAT ICT can see and operate everything, and identity stays with "
                   "whoever governs the account. The user carries CloudWatch access directly rather "
                   "than through the group, which shows permissions can sit in either place; groups "
                   "are for not repeating yourself across many users, not a requirement.")),

    dict(n=9, title="Review the instance role",
         job="Open the role named below and look at the policies attached to it, then write its name down. "
             "Your servers reach the storage buckets and the logging service through this role, so no credentials are ever stored on a server. You select it in the next task.",
         note="the lab environment supplies this role ready-made and does not permit you to create "
              "your own, so this task is a review rather than a build.",
         settings=[("Go to", "IAM \u2192 Roles"),
                   ("Open", "LabRole  (the role the environment provides)"),
                   ("Look at", "the Permissions tab — the policies attached to it"),
                   ("Answer for yourself", "which services does it let your servers reach?"),
                   ("Write down", "the role name, exactly as shown — task 10 asks for it")],
         capture="the LabRole summary page with the Permissions tab open, showing the policies "
                 "attached to it.",
         uoc=["ICTCLD401 PC 1.6", "ICTCLD401 PC 1.7"]),

    dict(n=10, title="Create the launch template",
         job="Create a launch template, using the settings below. "
             "It defines how every application server is built. The Auto Scaling group creates servers from it, so nothing is configured on a server by hand.",
         settings=[("Name", "yat-lms-lt"),
                   ("AMI", "Windows Server (latest available)"),
                   ("Instance type", "your choice — see the decision below"),
                   ("Key pair", "do not include one"),
                   ("Subnet", "do not include one — the Auto Scaling group supplies it"),
                   ("Security group", "yat-lms-app-sg"),
                   ("Root volume", "gp3, 30 GB"),
                   ("Data volume", "add a second volume: gp3, 8 GB, device name xvdb. The "
                                    "console will not accept the volume without a device name, and "
                                    "the /dev/sd... names are greyed out because this is a Windows "
                                    "image"),
                   ("Advanced details \u2192 IAM instance profile", "LabInstanceProfile"),
                   ("Advanced details \u2192 User data", "the script below, exactly as written")],
         code=("User data", [
             "<powershell>",
             "Install-WindowsFeature -Name Web-Server -IncludeManagementTools",
             "Set-Content -Path \"C:\\inetpub\\wwwroot\\index.html\" -Value \"<html><body>",
             "<h1>YAT LMS</h1><p>Infrastructure ready - awaiting application",
             "deployment.</p></body></html>\"",
             "</powershell>",
         ]),
         capture="the launch template summary, showing the AMI, instance type, instance profile, "
                 "security group and both volumes.",
         uoc=["ICTCLD401 PC 2.3", "ICTCLD401 PC 2.4", "ICTCLD401 PC 1.1", "ICTCLD401 PC 1.3"],
         decision=("Choose the instance type for the application tier: t3.micro or t3.small. "
                   "Name the one you did not choose, say which you chose, and explain why against "
                   "the LMS workload described in the LMS Application Specification.",
                   "t3.micro has 1 GiB of memory and t3.small has 2 GiB. Windows Server alone uses "
                   "most of 1 GiB before the application starts, so t3.micro would spend its time "
                   "paging and the health check would be slow and unreliable. I chose t3.small. "
                   "Both are well below what the specification's 200–300 typical concurrent users "
                   "would need — in production I would size to a general-purpose instance several "
                   "times this — but this environment caps what I can launch, and t3.small is the "
                   "better of the two available.")),

    dict(n=11, title="Create the target group",
         job="Create a target group, using the settings below. It is the list of servers the load "
             "balancer will send traffic to, and the health check that decides which of them are "
             "fit to receive it.",
         settings=[("Name", "yat-lms-tg"),
                   ("Target type", "Instances"),
                   ("Protocol and port", "HTTP, port 80"),
                   ("VPC", "yat-lms-vpc"),
                   ("Health check protocol", "HTTP"),
                   ("Health check path", "/"),
                   ("Interval", "30 seconds"),
                   ("Unhealthy threshold", "2"),
                   ("Register targets", "skip this step and create the group with no targets. The "
                                        "Auto Scaling group registers them for you in task 13")],
         evidence_note="No screenshot required. Creating the target group is not an assessable "
                       "action on its own — it is the prerequisite for the load balancer, which "
                       "sends traffic to it, and the Auto Scaling group, which fills it with "
                       "servers.",
         uoc=[]),

    dict(n=12, title="Create the load balancer",
         job="Create an Application Load Balancer, using the settings below. It sits in front of "
             "the application tier and sends traffic only to servers the target group reports as "
             "healthy.",
         settings=[("Type", "Application Load Balancer"),
                   ("Name", "yat-lms-alb"),
                   ("Scheme", "Internet-facing"),
                   ("VPC", "yat-lms-vpc"),
                   ("Subnets", "public-web-a and public-web-b. It requires two zones and will not "
                               "create with only one"),
                   ("Security group", "yat-lms-alb-sg. Remove the default group if the console adds "
                                      "it"),
                   ("Listener", "HTTP on port 80"),
                   ("Listener forwards to", "yat-lms-tg")],
         capture="the load balancer summary, showing its scheme, DNS name, subnets and security "
                 "group.",
         uoc=["ICTCLD401 PC 2.2", "ICTCLD502 PC 4.1"]),

    dict(n=13, title="Create the Auto Scaling group",
         job="Create an Auto Scaling group, using the settings below. "
             "It adds and removes application servers on its own, without anyone touching the instance count.",
         settings=[("Step 1 \u2014 Launch template", "yat-lms-lt"),
                   ("Step 2 \u2014 VPC", "yat-lms-vpc"),
                   ("Step 2 \u2014 Subnet", "private-app-a only"),
                   ("Step 3 \u2014 Load balancing", "choose Attach to an existing load balancer, "
                                               "then Choose from your load balancer target groups, "
                                               "and select yat-lms-tg. This step defaults to No "
                                               "load balancer, so it is easily skipped"),
                   ("Step 3 \u2014 Health check type", "tick Turn on Elastic Load Balancing health "
                                                  "checks; grace period 300 seconds"),
                   ("Step 4 \u2014 Group size", "desired 1, minimum 1, maximum 2"),
                   ("Step 4 \u2014 Scaling policy", "Target tracking, Average CPU utilization, "
                                               "target value 70"),
                   ("Step 4 \u2014 Warm-up", "60 seconds")],
         capture="the Auto Scaling group details, showing group size, the attached target group and "
                 "the scaling policy.",
         uoc=["ICTCLD401 PC 3.1", "ICTCLD401 PE 2"]),

    dict(n=14, title="Create the database subnet group",
         job="Create a database subnet group, using the settings below. It tells the database "
             "service which subnets it is allowed to place the database in.",
         note="create this in the RDS console, not the VPC console \u2014 RDS \u2192 Subnet groups \u2192 "
              "Create DB subnet group. It is easy to go looking for it where the subnets "
              "themselves live.",
         settings=[("Name", "yat-lms-db-subnet-group"),
                   ("VPC", "yat-lms-vpc"),
                   ("Availability Zones", "us-east-1a and us-east-1b"),
                   ("Subnets", "private-data-a and private-data-b"),
                   ("Why two", "the service requires subnets in two zones even though this database "
                               "is single-AZ. The second one stays empty until the next phase")],
         evidence_note="No screenshot required. Creating the subnet group is not an assessable "
                       "action on its own — it is the prerequisite for the database, which "
                       "cannot be created without one.",
         uoc=[]),

    dict(n=15, title="Deploy the database",
         job="Create the database, using the settings below. Create it empty \u2014 loading the schema "
             "and the data is YAT ICT job, not yours.",
         note="use Standard create, not Easy create. Easy create chooses the network, the subnet "
              "group, the security group and the public-access setting for you and does not let "
              "you change them, which is four of the settings this task specifies.",
         settings=[("Creation method", "Standard create"),
                   ("Engine", "MySQL"),
                   ("Template", "Free tier, if offered"),
                   ("Instance class", "your choice \u2014 see the decision below"),
                   ("Storage", "gp3, 20 GB"),
                   ("Multi-AZ", "Do not create a standby. This build is deliberately single-AZ"),
                   ("VPC", "yat-lms-vpc"),
                   ("Subnet group", "yat-lms-db-subnet-group"),
                   ("Public access", "No"),
                   ("Security group", "yat-lms-db-sg. Remove the default group if the console adds "
                                      "it"),
                   ("Encryption", "enabled"),
                   ("Backup retention", "7 days")],
         capture="the database summary showing Available, its class, encryption enabled and public "
                 "access set to No.",
         uoc=["ICTCLD401 PC 2.5", "ICTCLD401 PC 1.1", "ICTCLD401 PC 1.3", "ICTCLD401 PE 2"],
         decision=("Choose the database instance class: db.t3.micro or db.t3.small. Name the one you "
                   "did not choose, say which you chose, and explain why against the workload.",
                   "db.t3.micro has 1 GiB of memory, db.t3.small has 2 GiB. An LMS database is "
                   "read-heavy and benefits from having enough memory to keep frequently-read pages "
                   "cached, so I chose db.t3.small. As with the application tier, both are well "
                   "below what the real workload would need; the environment caps the choice and "
                   "db.t3.small is the better of the two.")),

    dict(n=16, title="Create the two monitoring alarms",
         job="Create two CloudWatch alarms, using the settings below. They are the baseline "
             "monitoring that tells the platform when something is wrong \u2014 one watching whether "
             "the application tier is serving, one watching whether the database is running out of "
             "room.",
         settings=[("Notification topic", "create an SNS topic named yat-lms-alerts with your own "
                                          "email address, and use it for both alarms. You do not "
                                          "need to confirm the subscription email for the alarm to "
                                          "be created"),
                   ("\u2500 Alarm 1 \u2500", "the application tier is not serving"),
                   ("Metric", "Application ELB \u2192 Per AppELB, per TG Metrics \u2192 "
                              "UnHealthyHostCount, filtered to yat-lms-alb and yat-lms-tg"),
                   ("Statistic", "Maximum"),
                   ("Period", "1 minute"),
                   ("Threshold type", "Static"),
                   ("Condition", "Whenever UnHealthyHostCount is Greater/Equal than 1"),
                   ("Datapoints to alarm", "1 out of 1"),
                   ("Missing data treatment", "Treat missing data as missing"),
                   ("Notification", "In alarm \u2192 yat-lms-alerts"),
                   ("Name", "yat-lms-unhealthy-hosts"),
                   ("Description", "any target behind the load balancer has failed its health check"),
                   ("\u2500 Alarm 2 \u2500", "the database is running out of room"),
                   ("Metric", "RDS \u2192 DBInstanceIdentifier, then filter for FreeStorageSpace "
                              "and tick the row for your database"),
                   ("Statistic", "Minimum"),
                   ("Period", "5 minutes"),
                   ("Threshold type", "Static"),
                   ("Condition", "Whenever FreeStorageSpace is Lower than 3221225472"),
                   ("About that number", "it is bytes \u2014 15% of the 20 GiB you allocated. Type the "
                                         "raw digits, no commas and no units. The graph above the "
                                         "field renders in GB, but the field itself is bytes"),
                   ("Datapoints to alarm", "2 out of 2"),
                   ("Missing data treatment", "Treat missing data as missing"),
                   ("Notification", "In alarm \u2192 yat-lms-alerts"),
                   ("Name", "yat-lms-db-storage-low"),
                   ("Description", "record the allocated size the threshold is based on \u2014 15% of "
                                   "20 GiB allocated. The number alone does not tell anyone whether "
                                   "it is still correct after the storage is resized")],
         capture="the alarm list showing both alarms with their metrics, thresholds and current "
                 "state.",
         uoc=["ICTCLD502 PC 4.3"]),
]

# ---------------------------------------------------------------- tests
# Same shape as a task, but the student runs steps rather than building to settings.

TESTS = [
    dict(n="T1", title="Connect to the application server",
         job="Connect to your application server using Session Manager, then run the command "
             "below. This shows you can reach the server you built.",
         steps=["Open EC2 \u2192 Instances and select the running instance.",
                "Choose Connect, then the Session Manager tab, then Connect.",
                "A PowerShell session opens in the browser. Run the command below."],
         code=("Run this", ["hostname"]),
         note="Session Manager needs no key pair, no open port and no public IP address. That is "
              "why your server can sit in a private subnet with nothing exposed to the internet "
              "and still be reachable by you.",
         capture="the session window, showing the connection open and the hostname it returned.",
         uoc=["ICTCLD401 PE 3"]),

    dict(n="T2", title="Reach the internet from the application server",
         job="Run the command below in the session you opened in test 1. This shows the server can "
             "reach the internet through the NAT gateway.",
         steps=["Run the command below.",
                "Confirm it returns StatusCode 200."],
         code=("Run this", ["Invoke-WebRequest -Uri https://aws.amazon.com -UseBasicParsing |",
                            "    Select-Object StatusCode"]),
         capture="the command and the status code it returned.",
         uoc=["ICTCLD401 PC 2.6"]),

    dict(n="T3", title="Reach the database from the application server",
         job="Run the command below in the same session. This shows the application tier can reach "
             "the database tier privately, over the database port.",
         steps=["Open RDS \u2192 Databases \u2192 your database \u2192 Connectivity & security, and copy "
                "the Endpoint. It looks like yat-lms-db.something.us-east-1.rds.amazonaws.com",
                "Run the command below, replacing <YOUR-DB-ENDPOINT> with the address you just "
                "copied. Do not paste the command as it is written \u2014 it will fail to resolve.",
                "Confirm the result reads TcpTestSucceeded : True."],
         code=("Run this", ["Test-NetConnection <YOUR-DB-ENDPOINT> -Port 3306"]),
         capture="the command and its output showing TcpTestSucceeded : True.",
         uoc=["ICTCLD401 PC 2.6", "ICTCLD502 PC 4.2"]),

    dict(n="T4", title="Reach the application through the load balancer",
         job="Open your load balancer address in a web browser on your own computer. This shows "
             "the whole path works \u2014 the internet reaches the load balancer, the load balancer "
             "reaches your private server, and the server serves the page.",
         steps=["Open EC2 \u2192 Load Balancers and copy the DNS name of yat-lms-alb.",
                "Paste it into a browser tab on your own machine, with http:// in front of it.",
                "Confirm the placeholder page loads."],
         note="your server has no public address, so what you are reaching is the load balancer. It "
              "is the only thing allowed to talk to the server, which is what the security groups "
              "enforce.",
         capture="the browser showing the page, with the load balancer address visible in the "
                 "address bar.",
         uoc=["ICTCLD401 PC 2.6", "ICTCLD502 PC 4.2"]),

    dict(n="T5", title="Automatic scaling",
         job="Change the scaling policy as described below and watch what happens. This shows the "
             "Auto Scaling group adds and removes servers on its own, without you changing the "
             "instance count.",
         steps=["Open EC2 \u2192 Instances, select your running instance and open the Monitoring "
                "tab. Read the CPU utilization figure \u2014 on an idle server it will be low, often "
                "only one or two percent.",
                "Open your Auto Scaling group and go to the Automatic scaling tab.",
                "Edit the target tracking policy and set the target value BELOW the figure you "
                "just read, then save. The policy keeps average CPU at the target, so a target "
                "above current usage tells it to scale in rather than out. If your reading is very "
                "low, use 1.",
                "Watch the Activity tab until a second instance launches and enters service.",
                "Edit the policy again, raise the target value to 90, and save.",
                "Watch the Activity tab until the group scales back to one instance.",
                "Set the target value back to 70."],
         note="scaling in is slower than scaling out \u2014 give it several minutes before deciding it "
              "has not worked.",
         assessor_note="proven in the lab: an idle Windows t3.small read about 6% CPU, and setting "
                       "the target value to 1 triggered a scale-out within a few minutes. Expect "
                       "students to report readings of a few percent and targets of 1. A student "
                       "who reports a higher reading and a correspondingly higher target has done "
                       "the same thing correctly. Scale-in takes noticeably longer than scale-out, "
                       "so a submission showing a long gap between the two Activity entries is "
                       "normal, not a fault.",
         capture="the Activity tab, showing both the scale-out and the scale-in entries with their "
                 "times.",
         uoc=["ICTCLD401 PC 3.2"]),
]

# ---------------------------------------------------------------- handover

HANDOVER = dict(
    title="File the completed run sheet",
    job="You have now built the platform, tested it and answered the questions. Hand the build "
        "over by filing this completed run sheet where YAT records procedures require, so YAT ICT "
        "can find it after the engagement ends. Do this last.",
    prompt="State where you filed it and which YAT policy required that location.",
    points=["Names the actual location the report was filed in, not a general intention.",
            "Names the YAT policy that governs it — the Records Management Policy, or the Backup "
            "and Retention Policy.",
            "A student who says only \"I would file it appropriately\" has not evidenced this."],
    uoc=["ICTCLD401 PC 4.3"])

# ---------------------------------------------------------------- knowledge questions
# resources — optional: (label, url) pairs pointing at the YAT intranet page a student would
#             reasonably consult to answer this question in YAT's terms rather than generically.
#             Only where a page genuinely carries the material. Q3 (shared responsibility) and
#             Q6 (DNS) have no page in this state, so they carry none rather than a token link.

QUESTIONS = [
    dict(n="Q1", uoc=["ICTCLD401 KE 5"],
         q="Name the compute, networking and scaling services you deployed. For each, explain the "
           "feature it provides and how it supports the YAT LMS specifically.",
         resources=[
             ("LMS Application Specification — the user population, concurrent load and "
              "service-level expectations your services have to support",
              f"{SITE}/intranet/s1-cl1-at2/ict/lms-application-spec"),
             ("LMS Cloud Architecture — Baseline Design — the design your build implements",
              f"{SITE}/intranet/s1-cl1-at2/projects/lms-cloud-infrastructure/"
              "cloud-architecture-baseline"),
         ],
         points=["Names the three: the instance, the load balancer, the Auto Scaling group.",
                 "Instance: the compute the LMS runs on, sized against the concurrent-user load.",
                 "Load balancer: distributes traffic and removes unhealthy servers from rotation.",
                 "Auto Scaling: adds capacity during assessment-window peaks without anyone acting.",
                 "Each tied to YAT, not explained generically."]),

    dict(n="Q2", uoc=["ICTCLD401 KE 6"],
         q="In your build: (a) why a managed database service rather than running the database on "
           "the server yourself? (b) your build uses block storage in two places — the volumes on "
           "the server and the storage behind the database. Explain how block storage differs from "
           "object storage, and when each is the right choice for an LMS. (c) why scale "
           "horizontally rather than vertically, and what does that trade off?",
         resources=[
             ("LMS Cloud Architecture — Baseline Design — why the design chose a managed database "
              "and a scaling group rather than a single larger server",
              f"{SITE}/intranet/s1-cl1-at2/projects/lms-cloud-infrastructure/"
              "cloud-architecture-baseline"),
             ("LMS Application Specification — what the LMS stores, which is what the storage "
              "choices have to suit",
              f"{SITE}/intranet/s1-cl1-at2/ict/lms-application-spec"),
         ],
         points=["(a) Managed: AWS handles patching, backups and monitoring; less for YAT ICT to "
                 "operate; and multi-AZ is one setting away for the next phase.",
                 "(b) Block storage is a disk attached to one server, addressed in blocks, and it "
                 "is what an operating system and a database need. Object storage holds whole "
                 "files retrieved by name over the network, is effectively unlimited, is cheaper "
                 "at scale, and can age files off to archive. For an LMS: block for the system and "
                 "the database, object for course attachments and student submissions.",
                 "(c) Horizontal adds servers with no downtime, which suits an LMS that must stay "
                 "up; vertical needs a restart. Trade-off: the application must tolerate requests "
                 "landing on different servers."]),

    dict(n="Q3", uoc=["ICTCLD401 KE 7"],
         q="Using your own environment, name two responsibilities that remain YAT's and one that "
           "shifted to AWS under the shared-responsibility model, and say why each falls where it "
           "does.",
         points=["Two YAT-side: operating-system patching after handover; managing users and MFA.",
                 "One shifted: hardware, data-centre power and cooling; or backup media handling, "
                 "now the managed database's job.",
                 "Reasons given, not just a list."]),

    dict(n="Q4", uoc=["ICTCLD401 KE 8"],
         q="Pick one IAM group you created. Describe the permissions you gave it, the job function "
           "it serves, and why its permissions differ from another group in your build.",
         resources=[
             ("User Access Policy — YAT's role-based access model, its role groups and "
              "permissions matrix, and its rules on privileged access",
              f"{SITE}/intranet/s1-cl1-at2/policies/user-access"),
         ],
         points=["Names a real group from their own build.",
                 "Ties the permissions to what that job actually needs to do.",
                 "Contrasts it with another group and explains the boundary — e.g. YAT ICT operate "
                 "the platform but do not change IAM, because that is a security boundary."]),

    dict(n="Q5", uoc=["ICTCLD401 KE 9"],
         q="The yat-lms-db-sg security group you created in task 7 permits MySQL traffic on port "
           "3306 from one source only. State the rule exactly as you configured it, explain why "
           "the database is restricted this way rather than being reachable from the internet, "
           "and describe the risk to YAT if that restriction were removed.",
         resources=[
             ("LMS Application Specification § 4, Data stored — what the MySQL database "
              "actually holds",
              f"{SITE}/intranet/s1-cl1-at2/ict/lms-application-spec#4-data-stored"),
         ],
         points=["States the rule as built: inbound MySQL on 3306, source yat-lms-app-sg — the "
                 "application tier's security group, not an address range and not the internet.",
                 "Explains the restriction: only the application tier has any business reaching the "
                 "database. Naming a security group as the source rather than an address range "
                 "means the rule follows the instances, however many the Auto Scaling group "
                 "launches and whatever addresses they get.",
                 "Names a concrete risk grounded in what the database holds — the ~50 GB of "
                 "student personal information in § 4 of the specification (names, dates of "
                 "birth, USIs, enrolment and fee status, assessment results) becomes reachable "
                 "from the internet and the database engine can be attacked directly for "
                 "credentials. YAT carries Privacy Act and APP obligations over that data.",
                 "A student who says only that it would be \"less secure\" has not evidenced this."]),

    dict(n="Q6", uoc=["ICTCLD401 KE 10"],
         q="When a YAT staff member types the LMS address into their browser, DNS resolution has to "
           "happen before anything else can. Explain what DNS does at that moment in your "
           "deployment, and what would fail for YAT if it were misconfigured.",
         points=["A browser connects to an address, not a name — DNS is the translation step.",
                 "In this deployment the name resolves to the load balancer.",
                 "The load balancer's addresses change, which is why the name points at it rather "
                 "than at a server.",
                 "If misconfigured, staff cannot reach the LMS even though every server is running."]),
]

# ---------------------------------------------------------------- rendering


def _p(doc, text, size=BODY_PT, bold=False, italic=False, colour=None, indent=None, after=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if colour:
        r.font.color.rgb = RGBColor.from_string(colour)
    return p


def _settings_table(doc, settings):
    t = doc.add_table(rows=0, cols=2)
    for label, value in settings:
        cells = t.add_row().cells
        set_cell_borders(cells[0]); set_cell_borders(cells[1]); shade_cell(cells[0], CREAM)
        lr = cells[0].paragraphs[0].add_run(label); lr.bold = True; lr.font.size = Pt(9.5)
        cells[1].paragraphs[0].add_run(value).font.size = Pt(9.5)
        cells[0].width = Cm(4.2); cells[1].width = Cm(11.6)
    doc.add_paragraph()


def _box(doc, lines, height_hint=True):
    """A bordered drop-zone. `lines` is a list of (text, colour_or_None, bold, italic)."""
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    set_cell_borders(cell); shade_cell(cell, CREAM); cell.width = Cm(16.6)
    for i, (text, colour, bold, italic) in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if colour is None else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        r.font.size = Pt(9.5 if colour else 10); r.bold = bold; r.italic = italic
        if colour:
            r.font.color.rgb = RGBColor.from_string(colour)
    doc.add_paragraph()


def _image_box(doc, caption, image):
    """A bordered drop-zone holding a real exemplar capture, under its description."""
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    set_cell_borders(cell); shade_cell(cell, CREAM); cell.width = Cm(16.6)
    p = cell.paragraphs[0]
    r = p.add_run(caption)
    r.font.size = Pt(9.5); r.italic = True
    r.font.color.rgb = RGBColor.from_string(CAPTURE)
    pic = cell.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.add_run().add_picture(str(image), width=Cm(IMAGE_CM))
    doc.add_paragraph()


def _screenshot_slot(doc, capture, mode, image=None):
    if mode == "assessor":
        if image is not None:
            _image_box(doc, f"SCREENSHOT — {capture}", image)
        else:
            _box(doc, [(f"SCREENSHOT — {capture}", CAPTURE, False, True)])
    else:
        _box(doc, [("[ PASTE YOUR SCREENSHOT HERE ]", None, True, False),
                   (capture, GREY, False, True)])


def _evidence_images(evidence_dir, key):
    """The exemplar captures filed for one task or test, in order.

    Convention: <key>.png for a single capture, <key>.<n>.png where a task needed several
    (three security groups are three screenshots). A task with more captures on disk than
    evidence boxes in the run sheet is fine — the extras are placed after the last box.
    """
    if evidence_dir is None:
        return []
    return sorted(Path(evidence_dir).glob(f"{key}.*.png")) + \
        sorted(Path(evidence_dir).glob(f"{key}.png"))


def _place_evidence(doc, captures, mode, images):
    """Render each evidence box, pairing it with its capture where one was filed."""
    for i, cap in enumerate(captures):
        _screenshot_slot(doc, cap, mode, image=images[i] if i < len(images) else None)
    for extra in images[len(captures):]:
        _image_box(doc, "SCREENSHOT — continued", extra)


def report_evidence(evidence_dir, tasks=None, tests=None):
    """What the exemplar folder covers, printed at build time so a gap is not silent."""
    tasks = TASKS if tasks is None else tasks
    tests = TESTS if tests is None else tests
    keys = [f"task-{t['n']:02d}" for t in tasks if not t.get("evidence_note")] + \
           [f"test-{i:02d}" for i, _ in enumerate(tests, 1)]
    exempt = [f"task-{t['n']:02d}" for t in tasks if t.get("evidence_note")]
    found = {k: _evidence_images(evidence_dir, k) for k in keys}
    missing = [k for k, v in found.items() if not v]
    placed = sum(len(v) for v in found.values())
    print(f"Exemplar evidence: {placed} capture(s) placed across "
          f"{len(keys) - len(missing)}/{len(keys)} tasks and tests"
          + (f" ({', '.join(exempt)} ask for none)." if exempt else "."))
    if missing:
        print(f"  NO CAPTURE ON FILE for: {', '.join(missing)} — "
              f"these render as the description alone.")
    orphans = sorted(p.name for p in Path(evidence_dir).glob("*.png")
                     if not any(p in v for v in found.values()))
    if orphans:
        print(f"  NOT PLACED (name matches no task or test): {', '.join(orphans)}")


def _response_slot(doc, model, mode, points=None):
    """Blank box for the student; for the assessor either a model answer or key points."""
    if mode != "assessor":
        _box(doc, [("[ WRITE YOUR ANSWER HERE ]", None, True, False)])
        return
    if points:
        lines = [("Key points the answer should touch on:", MODEL, True, False)]
        lines += [(f"\u2022  {pt}", MODEL, False, False) for pt in points]
        _box(doc, lines)
    else:
        _box(doc, [(model, MODEL, False, False)])


def _flag(doc, text):
    """A small teal label above a heading — the thing a student quotes when they are stuck."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text.upper())
    r.bold = True; r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(TEAL)


def _code(doc, lines):
    """A verbatim block — monospace, shaded, for anything the student must copy exactly."""
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    set_cell_borders(cell); shade_cell(cell, CREAM); cell.width = Cm(16.6)
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.name = "Consolas"; r.font.size = Pt(9)
    doc.add_paragraph()


def _note(doc, text):
    """An environment-constraint note — set apart from the task, in the accent colour."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    lead = p.add_run("Environment constraint\u2003")
    lead.bold = True; lead.font.size = Pt(9.5)
    lead.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    r = p.add_run(text)
    r.italic = True; r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(TERRACOTTA)


def _clicks(doc, clicks):
    """Click-by-click detail. The practice sheet has this; the assessment deliberately does not."""
    _p(doc, "How to do it", bold=True, after=3)
    for i, step in enumerate(clicks, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(2)
        lead = p.add_run(f"{i}.  "); lead.bold = True; lead.font.size = Pt(10)
        p.add_run(step).font.size = Pt(10)
    doc.add_paragraph()


def _assessor_note(doc, text, mode):
    """Guidance for the assessor only — never rendered in the student copy."""
    if mode != "assessor":
        return
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    lead = p.add_run("Assessor note\u2003")
    lead.bold = True; lead.font.size = Pt(9.5)
    lead.font.color.rgb = RGBColor.from_string(MODEL)
    r = p.add_run(text)
    r.italic = True; r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(MODEL)


def _steps(doc, steps):
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(3)
        lead = p.add_run(f"{i}.  "); lead.bold = True; lead.font.size = Pt(BODY_PT)
        p.add_run(step).font.size = Pt(BODY_PT)
    doc.add_paragraph()


def _resources(doc, items):
    """Links to the intranet pages that carry the material a question draws on."""
    _p(doc, "Related resources", bold=True, after=3)
    for label, url in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(4)
        p.add_run("•  ").font.size = Pt(BODY_PT)
        add_hyperlink(p, label, url, size_pt=BODY_PT)
    doc.add_paragraph()


def render_front_matter(doc, h1):
    """Scenario, required resources and instructions — all at body size, one visual voice."""
    h1("Scenario")
    for para in SCENARIO:
        _p(doc, para, after=8)
    h1("Required resources (assessment conditions)")
    for label, url in RESOURCES:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(6)
        p.add_run("\u2022  ").font.size = Pt(BODY_PT)
        add_hyperlink(p, label, url, size_pt=BODY_PT)
    h1("Instructions to Student")
    for para in INSTRUCTIONS:
        _p(doc, para, after=8)


def render_run_sheet(doc, h1, h2, mode="student", tasks=None, tests=None,
                     questions=None, handover=None, region_note=None, evidence_dir=None):
    """Render the run sheet into `doc`. mode = student | assessor.

    evidence_dir — a folder of exemplar captures (see _evidence_images for the naming). When
    given AND mode is assessor, each evidence box carries the real screenshot under its
    description instead of the description alone, so a regenerated assessor copy is worked
    rather than blank. Defaults to None: every other caller renders exactly as before, and no
    other run sheet can pick up this one's captures.
    """
    tasks = TASKS if tasks is None else tasks
    tests = TESTS if tests is None else tests
    questions = QUESTIONS if questions is None else questions
    handover = HANDOVER if handover is None else handover
    _p(doc, REGION_NOTE if region_note is None else region_note, after=10)

    for task in tasks:
        _flag(doc, f"Task {task['n']}")
        h2(task["title"])
        # A prerequisite step carries no tags, so it makes no claim to evidence — the assessor
        # sees nothing to mark rather than a criterion with no capture behind it.
        if mode == "assessor" and task["uoc"]:
            _p(doc, "Evidences: " + " · ".join(f"[{u}]" for u in task["uoc"]),
               size=9, italic=True, colour=UOC, after=4)
        _p(doc, task["job"], after=6)
        if task.get("note"):
            _note(doc, task["note"])
        _settings_table(doc, task["settings"])
        if task.get("code"):
            label, lines = task["code"]
            _p(doc, label, bold=True, after=3)
            _code(doc, lines)
        if task.get("clicks"):
            _clicks(doc, task["clicks"])
        if task.get("decision"):
            prompt, model = task["decision"]
            _p(doc, "Your decision", bold=True, after=3)
            _p(doc, prompt, italic=True, size=9.5, colour=GREY, after=6)
            _response_slot(doc, model, mode)
        _p(doc, "Evidence", bold=True, after=3)
        if task.get("evidence_note"):
            _p(doc, task["evidence_note"], italic=True, size=9.5, colour=GREY, after=10)
        else:
            _place_evidence(doc, task.get("captures", [task.get("capture")]), mode,
                            _evidence_images(evidence_dir, f"task-{task['n']:02d}"))

    # ---- tests ----
    h1("Testing")
    _p(doc, "Five tests confirm the build works. Each tells you what it demonstrates and the steps "
            "to run it. Follow the steps, then paste the screenshot into the box. If a test does "
            "not pass, fix the problem, run it again, and note what you changed.",
       italic=True, size=9.5, colour=GREY, after=10)
    for i, test in enumerate(tests, 1):
        _flag(doc, f"Test {i}")
        h2(test["title"])
        if mode == "assessor":
            _p(doc, "Evidences: " + " \u00b7 ".join(f"[{u}]" for u in test["uoc"]),
               size=9, italic=True, colour=UOC, after=4)
        _p(doc, test["job"], after=6)
        _steps(doc, test["steps"])
        if test.get("code"):
            label, lines = test["code"]
            _p(doc, label, bold=True, after=3)
            _code(doc, lines)
        if test.get("note"):
            _note(doc, test["note"])
        if test.get("assessor_note"):
            _assessor_note(doc, test["assessor_note"], mode)
        _p(doc, "Evidence", bold=True, after=3)
        _place_evidence(doc, [test["capture"]], mode,
                        _evidence_images(evidence_dir, f"test-{i:02d}"))

    # ---- knowledge questions ----
    if questions:
        h1("Knowledge questions")
        _p(doc, "Answer each question about your own build, not in general terms. Refer to what "
                "you actually created and the choices you actually made.",
           italic=True, size=9.5, colour=GREY, after=10)
        for i, q in enumerate(questions, 1):
            _flag(doc, f"Question {i}")
            h2(q["q"][:58].rstrip() + ("..." if len(q["q"]) > 58 else ""))
            if mode == "assessor":
                _p(doc, "Evidences: " + " \u00b7 ".join(f"[{u}]" for u in q["uoc"]),
                   size=9, italic=True, colour=UOC, after=4)
            _p(doc, q["q"], after=6)
            if q.get("resources"):
                _resources(doc, q["resources"])
            _response_slot(doc, None, mode, points=q["points"])

    # ---- handover ----
    if handover:
        h1("Handover")
        h2(handover["title"])
        if mode == "assessor":
            _p(doc, "Evidences: " + " \u00b7 ".join(f"[{u}]" for u in handover["uoc"]),
               size=9, italic=True, colour=UOC, after=4)
        _p(doc, handover["job"], after=6)
        _p(doc, handover["prompt"], italic=True, size=9.5, colour=GREY, after=6)
        _response_slot(doc, None, mode, points=handover["points"])
