#!/usr/bin/env python3
"""Build the AT2 Deployment Report EXEMPLAR (.docx) — assessor marking reference.

A worked model answer for the YAT LMS Cloud Migration *foundation build* (S1-CL1 AT2,
ICTCLD401), built on the generic Deployment Report superset template. Demonstrates the
"Not applicable — [reason]" convention by marking the HA-only sections deferred to the
AT3 hardening phase. Evidence is described, not captured: "[SCREENSHOT — should show ...]".

Assessor-facing: retains UoC `Evidences:` tags, §8 Knowledge Evidence responses, and
reflections (the assessment layers a real org template omits). Reuses the docx brand
helpers from the helpers/ package.

Usage:  python scripts/s1_cl1/build_s1_cl1_at2_deployment_exemplar.py [output.docx]
Default: S1-CL1-Cloud-Design-Build/assessments/AT2/AT2-exemplar-deployment-report.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # content-repo scripts/ (brand + registry)  # noqa: E402
sys.path.insert(0, str(next(d / "scripts" for d in Path(__file__).resolve().parents if (d / "scripts" / "helpers" / "__init__.py").exists())))  # umbrella scripts/ (engine)  # noqa: E402
from helpers.docx_body_text import add_body_paragraph, add_bullet_list  # noqa: E402
from helpers.docx_evidence import add_described_evidence, add_not_applicable  # noqa: E402
from helpers.docx_styling import add_field, paragraph_bottom_rule, set_cell_borders, shade_cell  # noqa: E402
from helpers.docx_tables import add_data_table  # noqa: E402
from helpers.uoc_tags import add_uoc_evidence_tag  # noqa: E402
from brand import ADDRESS, CREAM, GREY, TEAL, TERRACOTTA  # noqa: E402
from helpers.scenario_document import build_header_footer, configure_styles, wordmark  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def build(path):
    doc = Document()
    configure_styles(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    build_header_footer(sec)

    # ---- COVER ----
    wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(GREY)
    paragraph_bottom_rule(doc.add_paragraph(), TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Deployment Report")
    sub = doc.add_paragraph().add_run("YAT LMS Cloud Migration — Foundation Build Phase")
    sub.font.size = Pt(15); sub.bold = True; sub.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    note = doc.add_paragraph().add_run("Assessor exemplar — internal marking reference (not for distribution to students)")
    note.italic = True; note.font.size = Pt(10); note.font.color.rgb = RGBColor.from_string(GREY)
    doc.add_paragraph()
    cover = [
        ("Engagement", "YAT LMS Cloud Migration — Foundation Build"),
        ("Engagement reference", "YAT-LMS-MIG-2026"),
        ("Report version", "v1.0"),
        ("Prepared by", "MTS Consultant"),
        ("Date submitted", "[DD/MM/YYYY]"),
        ("Submitted to", "Sam Walker (YAT ICT Manager) · Pat Lin (MTS Senior Consultant)"),
        ("Related documents", "YAT LMS Cloud Architecture — Baseline Design v1.0; YAT LMS Migration Business Case"),
        ("Classification", "Commercial-in-confidence"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    for k, v in cover:
        cells = ct.add_row().cells
        set_cell_borders(cells[0]); set_cell_borders(cells[1]); shade_cell(cells[0], CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        cells[1].paragraphs[0].add_run(v).font.size = Pt(10)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    doc.add_paragraph("Contents", style="Heading 1")
    add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                 placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h2 = lambda t: doc.add_paragraph(t, style="Heading 2")
    h3 = lambda t: doc.add_paragraph(t, style="Heading 3")

    h1("1. Executive Summary")
    add_uoc_evidence_tag(doc, "[ICTCLD401 PC 4.1] — partial (document and communicate the work)")
    add_body_paragraph(doc, "This report documents the foundation build phase of the YAT LMS cloud migration: the "
                 "baseline AWS infrastructure for the LMS. The design specifies the Sydney region (ap-southeast-2); "
                 "in the AWS Academy Learner Lab it is deployed to us-east-1 - [scenario: ap-southeast-2 (Sydney) | "
                 "deploy: us-east-1] - in availability zone us-east-1a. The build delivered identity and access management, the "
                 "network (a VPC with public, private-app and private-data subnets), the compute tier (an EC2 "
                 "Auto Scaling Group behind an Application Load Balancer), a managed MySQL database on RDS, "
                 "S3 and EBS storage, a three-tier security-group model, and a baseline of CloudWatch "
                 "monitoring. All connectivity, autoscaling and smoke tests passed and the infrastructure is "
                 "ready for YAT IT to deploy the DOODLE application onto.")
    add_body_paragraph(doc, "This phase is deliberately single-AZ: high availability (Multi-AZ database, cross-AZ "
                 "compute, failover testing, DR runbook and cross-Region backup) is out of scope here and is "
                 "the subject of the next phase. The engagement is ready to proceed to the HA design phase.")

    h1("2. Engagement Context")
    add_uoc_evidence_tag(doc, "[ICTCLD502 AC 3, AC 5] (scenario data sources used)")
    add_body_paragraph(doc, "This build implements the board-approved direction from the YAT LMS Migration Business "
                 "Case and the supplied YAT LMS Cloud Architecture — Baseline Design (v1.0, approved by Pat "
                 "Lin and Sam Walker). As the MTS consultant I implemented that architecture in the AWS "
                 "Academy environment. Per the MTS scope, the LMS application install, the MySQL data "
                 "migration, and the cutover remain YAT IT's responsibility; HA hardening is the next MTS "
                 "phase, covered in a separate report.")

    h1("3. Scope of Deployment")
    add_body_paragraph(doc, "In scope of this report: the baseline cloud foundation — IAM, network, compute, database, "
                 "autoscaling, storage, and a monitoring baseline. Out of scope and deferred to the HA "
                 "hardening phase: Multi-AZ database, cross-AZ resilience, failover and resize simulation, "
                 "the DR runbook, cross-Region backup copies, and HA-tuned monitoring.")
    h3("3.1 Maintenance-window context")
    add_not_applicable(doc, "this was a greenfield build in an empty account, not a change to a running system, so no "
            "maintenance window applied.")

    h1("4. Build / Change Narrative")
    add_uoc_evidence_tag(doc, "[ICTCLD401 PC 1.1–1.7, 2.1–2.5, 3.1] · [ICTCLD502 PC 1.3]")
    h3("4.1 Identity and access management (IAM)")
    add_body_paragraph(doc, "I created four IAM groups reflecting the engagement's job functions — YAT-ICT-Admins "
                 "(full administration), MTS-Consultants (build/operate within a permission boundary), "
                 "Application-Service (the EC2 instance role), and Read-Only-Auditors. MFA is enforced on all "
                 "human users via an IAM policy that denies actions without MFA. The EC2 instances assume an "
                 "instance profile granting only the S3 and CloudWatch access the application needs.")
    add_described_evidence(doc, "SCREENSHOT", "should show the IAM Groups list with the four groups and member counts, and an "
                          "IAM user with an MFA device enabled; region us-east-1 visible.")
    h3("4.2 Network topology")
    add_body_paragraph(doc, "I deployed a VPC (10.0.0.0/16) with three subnets in us-east-1a: public-web-a "
                 "(10.0.1.0/24) for the ALB and NAT Gateway, private-app-a (10.0.2.0/24) for the EC2 "
                 "instances, and private-data-a (10.0.3.0/24) for RDS. An Internet Gateway serves the public "
                 "subnet; a NAT Gateway gives the private app tier outbound internet (e.g. Windows Update) "
                 "without inbound exposure. Route tables route the public subnet to the IGW and the private "
                 "subnets to the NAT.")
    add_described_evidence(doc, "SCREENSHOT", "should show the VPC subnets list with the three subnets, their CIDRs and AZ, and "
                          "the route tables showing 0.0.0.0/0 → IGW (public) and → NAT (private-app).")
    h3("4.3 Compute (EC2 + Auto Scaling)")
    add_body_paragraph(doc, "A launch template defines the LMS application instance (Windows Server 2016, m6i.large). "
                 "The Auto Scaling Group runs min 1 / desired 1 / max 3 in this single-AZ baseline, with a "
                 "target-tracking policy that scales out when average CPU exceeds 65%.")
    add_described_evidence(doc, "SCREENSHOT", "should show the EC2 Auto Scaling Group with its launch template, min/desired/max, "
                          "and the scaling policy; and the running instance(s) with tags.")
    h3("4.4 Load balancing (ALB)")
    add_body_paragraph(doc, "An internet-facing Application Load Balancer in the public subnet forwards HTTPS to the app "
                 "tier. The target group health-checks the instances on an HTTP path; the HTTPS listener uses "
                 "an ACM certificate.")
    add_described_evidence(doc, "SCREENSHOT", "should show the ALB target group with the EC2 instance(s) reporting Healthy.")
    h3("4.5 Database (RDS)")
    add_body_paragraph(doc, "The LMS database runs on RDS for MySQL 8.0.35, instance class db.m6i.large, 150 GB gp3 "
                 "storage, encryption at rest enabled, and 7-day automated backup retention. This phase is "
                 "Single-AZ; Multi-AZ is enabled in the HA phase.")
    add_described_evidence(doc, "SCREENSHOT", "should show the RDS database in 'available' state, engine MySQL 8.0.35, "
                          "Single-AZ, storage encrypted.")
    h3("4.6 Storage (EBS + S3)")
    add_body_paragraph(doc, "The EC2 instance has an encrypted gp3 root volume (50 GB) and a gp3 data volume (100 GB). "
                 "Two S3 buckets hold LMS attachments and backups; both have Block Public Access enabled and "
                 "default encryption on.")
    add_described_evidence(doc, "SCREENSHOT", "should show the S3 buckets list with Block Public Access enabled, and the EBS "
                          "volumes attached to the instance.")
    h3("4.7 Security (security groups + encryption)")
    add_body_paragraph(doc, "A three-tier security-group model: sg-alb allows 443 from the internet; sg-app allows the "
                 "application port only from sg-alb; sg-db allows 3306 only from sg-app. Traffic is encrypted "
                 "in transit (HTTPS at the ALB, TLS to RDS) and at rest (EBS, RDS, S3).")
    add_described_evidence(doc, "SCREENSHOT", "should show the three security groups with inbound rules expanded, demonstrating "
                          "the tier-to-tier restriction.")
    h3("4.8 Monitoring (baseline)")
    add_body_paragraph(doc, "Two baseline CloudWatch alarms establish availability monitoring: ALB target health "
                 "status (any unhealthy target) for the application tier, and RDS free storage below 15% for "
                 "the database tier. Both notify an SNS topic. The Auto Scaling policy tracks CPU and manages "
                 "its own alarms. HA-tuned monitoring is added in the next phase.")
    add_described_evidence(doc, "SCREENSHOT", "should show the CloudWatch Alarms list with the baseline alarms configured.")
    h3("4.9 Cross-Region backup / replication")
    add_not_applicable(doc, "cross-Region resilience is deferred to the HA hardening phase (AT3).")

    h1("5. Configuration Decisions")
    add_uoc_evidence_tag(doc, "[ICTCLD401 PC 1.1] (options compared) · [ICTCLD401 PC 1.3] (best option selected)")
    add_data_table(doc, ["#", "Decision", "Options considered", "Choice", "Why this one"],
              [["C1", "Application-tier instance type", "t3.medium (2 vCPU, 4 GB) or m6i.large (2 vCPU, 8 GB)",
                "m6i.large",
                "t3.medium is burstable and cheaper, but the LMS runs sustained load through the working day and would exhaust CPU credits. m6i.large gives steady performance for the 200-300 typical concurrent users; the Auto Scaling group absorbs assessment-window peaks."],
               ["C2", "Database instance class and storage size", "db.t3.medium or db.m6i.large; 100 GB or 150 GB",
                "db.m6i.large, 150 GB",
                "Same burstable objection as C1 — the database carries steady read load. 150 GB covers the current ~68 GB with roughly three years of growth at ~25 GB/yr; bulk attachments go to S3, not the database."]],
              widths=[0.9, 3.3, 4.0, 2.6, 4.7])

    h1("6. Testing and Validation")
    add_uoc_evidence_tag(doc, "[ICTCLD401 PC 2.6, 3.2] · [ICTCLD401 PE 3] · [ICTCLD502 PC 4.2]")
    h3("6.1 Connect to the application server")
    add_body_paragraph(doc, "Connected to the application instance and confirmed its hostname.")
    add_described_evidence(doc, "SCREENSHOT", "should show the terminal session on the instance, the connection "
                            "succeeding and the hostname returned.")
    h3("6.2 Reach the internet from the application server")
    add_body_paragraph(doc, "curl -I https://aws.amazon.com returned HTTP/2 200, confirming outbound access through "
                 "the NAT gateway.")
    add_described_evidence(doc, "SCREENSHOT", "should show the command and the HTTP response header.")
    h3("6.3 Reach the database from the application server")
    add_body_paragraph(doc, "nc -zv against the database endpoint on port 3306 succeeded from the application tier, "
                 "over the private subnet. The database is not reachable from outside the VPC.")
    add_described_evidence(doc, "SCREENSHOT", "should show the command and its output confirming port 3306 was reachable.")
    h3("6.4 Reach the load balancer from the application server")
    add_body_paragraph(doc, "curl -I against the load balancer's DNS name returned an HTTP status line, confirming the "
                 "web and application tiers can see each other. With 6.3 this demonstrates connectivity across "
                 "all three tiers.")
    add_described_evidence(doc, "SCREENSHOT", "should show the command and the HTTP response from the load balancer.")
    h3("6.5 Automatic scaling")
    add_body_paragraph(doc, "Lowering the scaling policy's target value below current utilisation caused the Auto "
                 "Scaling group to launch a second instance without intervention; it entered service in about "
                 "four minutes. Raising the target well above current utilisation caused the group to scale "
                 "back to one instance. The target was then returned to the value recorded in §5.")
    add_described_evidence(doc, "SCREENSHOT", "should show the Auto Scaling Activity tab with both the scale-out and "
                            "the scale-in entries and their timestamps.")

    h1("7. Operational Handover")
    add_uoc_evidence_tag(doc, "[ICTCLD401 PC 4.3] (save documentation per organisational policies)")
    h3("7.1 Access")
    add_body_paragraph(doc, "YAT-ICT-Admins retain full administration; MTS-Consultants remain within the permission "
                 "boundary; MFA stays enforced; instance access is via SSM Session Manager only.")
    h3("7.2 Runbook references")
    add_bullet_list(doc, [
        "The supplied Baseline Design document (operational reference).",
        "The naming and tagging conventions, so YAT ICT can identify resources.",
        "Backup arrangements: RDS automated backups (7-day retention) and the EBS snapshot policy.",
        "The CloudWatch baseline alarms list and the SNS notification topic.",
    ])
    h3("7.3 Known limitations and what's next")
    add_body_paragraph(doc, "This foundation is single-AZ and is not highly available: the RDS instance and each AZ are "
                 "single points of failure. The HA hardening phase will enable RDS Multi-AZ, extend the ASG "
                 "and ALB across two AZs, add HA-tuned monitoring and cross-Region backup, and establish a DR "
                 "runbook with failure-simulation testing.")
    h3("7.4 Documentation filing")
    add_data_table(doc, ["Item", "Filed in", "Reference"],
              [["This Deployment Report (v1.0)", "YAT ICT shared documentation", "[ref]"],
               ["Configuration exports (Appendix B)", "YAT ICT shared documentation", "[ref]"],
               ["Test evidence (§6)", "YAT ICT shared documentation", "[ref]"]],
              widths=[6.5, 5.0, 4.0])
    h3("7.5 Feedback record")
    add_not_applicable(doc, "the formal feedback loop and end-of-engagement sign-off are captured at the close of the HA "
            "phase; this handover is an interim phase handover.")
    h3("7.6 Sign-off")
    add_data_table(doc, ["Role", "Name", "Date", "Acceptance"],
              [["Prepared by", "MTS Consultant", "[date]", "Submitted"],
               ["Accepted by", "Sam Walker (YAT ICT Manager)", "[date]",
                "Foundation build accepted; ready for HA design"]],
              widths=[4.5, 5.0, 2.5, 3.5])

    h1("8. Knowledge Evidence Responses")
    add_uoc_evidence_tag(doc, "[ICTCLD401 KE 5, 6, 7, 8, 9, 10]")
    h3("Q1 — Virtual machine, networking and scaling features")
    add_body_paragraph(doc, "EC2 provides the compute capacity that hosts the LMS application; I chose m6i.large to suit "
                 "the typical concurrent load while keeping cost down, leaning on the ASG for peaks. The Auto "
                 "Scaling Group provides automatic scaling — it adds instances when CPU passes 65% (the "
                 "assessment-window peak pattern) and removes them when load falls, so capacity tracks demand. "
                 "The Application Load Balancer distributes traffic across healthy instances and, via its "
                 "health check, stops routing to an unhealthy one — which is also what lets the ASG replace an "
                 "instance without downtime to users.")
    h3("Q2 — Managed services, storage options, and scaling model")
    add_body_paragraph(doc, "RDS is preferred over self-hosting MySQL on EC2 because AWS manages patching, backups and "
                 "(later) Multi-AZ failover — valuable given YAT ICT's thin cloud experience. EBS and S3 are "
                 "used together because they solve different problems: EBS is block storage attached to the "
                 "instance for the OS and working data, while S3 is durable, cheap object storage for the bulk "
                 "LMS attachments and backups that don't need to sit on a volume. The ASG scales horizontally; "
                 "vertical scaling (resizing the instance) was viable but would mean downtime on each resize "
                 "and a hard ceiling, whereas horizontal scaling matches the spiky assessment-window load "
                 "without interruption.")
    h3("Q3 — Shared security responsibility")
    add_body_paragraph(doc, "Two responsibilities that remain YAT's in this environment: configuring the security groups "
                 "and IAM permissions correctly (AWS provides the controls; using them well is the customer's "
                 "job), and managing the guest OS — patching Windows Server and the application stack on the "
                 "EC2 instances. One responsibility that shifted to AWS by moving to RDS: the database host "
                 "OS patching and the underlying hardware, which YAT previously owned on the on-prem server.")
    h3("Q4 — User access protocols and organisational hierarchy")
    add_body_paragraph(doc, "The MTS-Consultants group has build-and-operate permissions within a permission boundary "
                 "limiting it to us-east-1 and LMS-tagged resources, with no IAM or billing rights. It "
                 "serves the external consultant function. Its permissions must differ from YAT-ICT-Admins "
                 "(who hold full administration as the system owner) so that an external party can do the "
                 "delivery work without being able to alter identity, billing, or resources outside the "
                 "engagement.")
    h3("Q5 — Security policies and network traffic limits")
    add_body_paragraph(doc, "sg-db allows inbound 3306 only from sg-app, and no outbound to the internet. Traffic is "
                 "restricted this way so the database is reachable only from the application tier over the "
                 "private network — never from the internet or even the web tier directly. If the restriction "
                 "were removed (e.g. 3306 open to the VPC or the internet), the database would be exposed to "
                 "any compromised host or external attacker, putting the student personal information and "
                 "records at risk.")
    h3("Q6 — Role of DNS in the deployment")
    add_body_paragraph(doc, "A browser cannot connect to a name, only to an address, so DNS does the translation first. "
                 "When staff type lms.yat.edu.au, YAT's DNS resolves it to the load balancer's DNS name, and "
                 "AWS resolves that to one of the load balancer's current addresses. Those addresses change, "
                 "which is precisely why the hostname points at the load balancer rather than at an instance. "
                 "If it were misconfigured, staff would get the wrong address or none at all and could not "
                 "reach the LMS — even though every server behind it is running normally.")

    # ---- APPENDICES ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1("Appendix A — Build evidence (screenshots)")
    add_body_paragraph(doc, "Evidence captures are described below in lieu of live screenshots for this exemplar.")
    for kind, d in [
        ("SCREENSHOT A1", "IAM dashboard / groups — the four groups with member counts; an MFA-enabled user; region us-east-1."),
        ("SCREENSHOT A2", "VPC subnets — public-web-a, private-app-a, private-data-a with CIDRs and AZ; route tables (IGW / NAT)."),
        ("SCREENSHOT A3", "EC2 instances + ASG — running instance(s) with tags; ASG min/desired/max and scaling policy."),
        ("SCREENSHOT A4", "ALB target group — instance(s) reporting Healthy."),
        ("SCREENSHOT A5", "RDS database — available, MySQL 8.0.35, Single-AZ, encryption enabled."),
        ("SCREENSHOT A6", "S3 buckets — attachments + backups with Block Public Access enabled; EBS volumes attached."),
        ("SCREENSHOT A7", "Security groups — sg-alb / sg-app / sg-db with inbound rules expanded."),
        ("SCREENSHOT A8", "CloudWatch Alarms — the baseline alarm set."),
    ]:
        add_described_evidence(doc, kind, d)
    h1("Document control")
    add_data_table(doc, ["Field", "Value"],
              [["Document version", "v1.0 — Initial submission"],
               ["Author", "MTS Consultant"],
               ["Engagement", "YAT LMS Cloud Migration — Foundation Build Phase"],
               ["Distribution", "Sam Walker (YAT ICT Manager), Pat Lin (MTS Senior Consultant)"],
               ["Successor document", "YAT LMS HA Hardening — HA Design + HA Deployment Report (AT3)"]],
              widths=[5.0, 10.5])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL1-Cloud-Design-Build/assessments/AT2/AT2-exemplar-deployment-report.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
