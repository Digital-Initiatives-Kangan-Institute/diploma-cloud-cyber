#!/usr/bin/env python3
"""Shared rendering primitives for the S1-CL1 workbook instruments.

A workbook instrument is ONE definition rendered two ways:

  student   the blank sheet the student works through — the job, empty capture tables,
            empty screenshot boxes, empty response boxes. No UoC tags.
  assessor  the worked copy: every screenshot box replaced by a terracotta description of
            what should be in it, every response box filled with a teal model answer, each
            element tagged with the UoC items it evidences, and each carrying the standard
            it is actually marked against.

These are the pieces both AT2 and AT3 draw on. They were lifted verbatim from
`at2_run_sheet.py`, which keeps its own copies until the session working on it lands; when
it does, it imports from here and the duplicates go. Everything AT3 needed on top —
`design_table`, `diagram_slot`, `standard_line` — is new here.

The marking model these support: a settings table is context, not the standard. The tags on
`Evidences:` say which UoC items an element carries; the `Satisfactory when:` line says what
has to be true for them to be met. An assessor marks the second, never the table.
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # noqa: E402
sys.path.insert(0, str(next(d / "scripts" for d in Path(__file__).resolve().parents if (d / "scripts" / "helpers" / "__init__.py").exists())))  # noqa: E402
from helpers.docx_styling import set_cell_borders, shade_cell  # noqa: E402
from helpers.instrument_layout import add_hyperlink  # noqa: E402
from brand import CREAM, GREY, TEAL, TERRACOTTA  # noqa: E402

from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402

BODY_PT = 10.5
UOC = "6B6660"          # muted — the assessor-only traceability line
CAPTURE = TERRACOTTA    # exemplar: what the screenshot should have shown
MODEL = TEAL            # exemplar: the model written answer
NOTES_BG = "EAF2F8"     # practice only: the student's own notes box
IMAGE_CM = 16.0         # exemplar captures, sized to sit inside the 16.6 cm evidence box

__all__ = [
    "BODY_PT", "UOC", "CAPTURE", "MODEL", "GREY", "TEAL", "CREAM", "IMAGE_CM", "add_hyperlink",
    "p", "heading2", "settings_table", "design_table", "box", "screenshot_slot", "diagram_slot",
    "image_box", "evidence_images", "place_evidence",
    "response_slot", "flag", "code", "note", "clicks", "assessor_note", "steps", "consider",
    "uoc_line", "standard_line", "resources_block", "notes_box",
]


def p(doc, text, size=BODY_PT, bold=False, italic=False, colour=None, indent=None, after=None):
    par = doc.add_paragraph()
    if indent:
        par.paragraph_format.left_indent = Cm(indent)
    if after is not None:
        par.paragraph_format.space_after = Pt(after)
    r = par.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if colour:
        r.font.color.rgb = RGBColor.from_string(colour)
    return par


def settings_table(doc, settings):
    """A given (label, value) table — what the student builds to. Values are supplied.

    `value` is a plain string, or a list of (text, bold) segments where part of the value
    needs to stand out from the rest.
    """
    t = doc.add_table(rows=0, cols=2)
    for label, value in settings:
        cells = t.add_row().cells
        set_cell_borders(cells[0]); set_cell_borders(cells[1]); shade_cell(cells[0], CREAM)
        lr = cells[0].paragraphs[0].add_run(label); lr.bold = True; lr.font.size = Pt(9.5)
        par = cells[1].paragraphs[0]
        for text, bold in ([(value, False)] if isinstance(value, str) else value):
            r = par.add_run(text); r.bold = bold; r.font.size = Pt(9.5)
        cells[0].width = Cm(4.2); cells[1].width = Cm(11.6)
    doc.add_paragraph()


def design_table(doc, columns, rows, mode, blank_rows=3, given=0, exemplar=0):
    """A capture table the STUDENT fills in — the inverse of settings_table.

    `rows` are the model answers, shown teal in the assessor copy.

    `given` is how many LEADING COLUMNS are pre-filled in the student copy too. It is the
    scaffolding dial: pre-fill a column when knowing it is not the evidence (a student does
    not demonstrate PC 2.3 by knowing that a database has a recovery objective), and leave it
    blank when the column IS the finding the item asks for (naming the single points of
    failure is the whole of PC 2.2). Given cells render as ordinary body text, not as model
    answers — to the student they are part of the question.

    `exemplar` is the same dial turned the other way: how many LEADING ROWS are shown filled
    right across, as a worked example of the sort of answer the table wants. PRACTICE ONLY —
    it shows the shape of an answer, so it must never be one of the answers the task is
    actually asking for. Where a task has only one real answer, the exemplar is drawn from
    the CURRENT environment instead (task 7 describes an existing subnet; task 12 restates
    the alarm that already exists) — form without finding. Exemplar rows render teal italic
    and are announced above the table, so they cannot be mistaken for the student's own work.

    `blank_rows` sets how much room an un-given table offers. Set it deliberately: three
    blank rows on a "list every single point of failure" table reads as "there are three".
    Where an exemplar row was ADDED to an otherwise empty table, blank_rows carries a +1 so
    the student keeps the working space they had.
    """
    if exemplar and mode != "assessor":
        p(doc, "The first row is filled in as an example — yours go underneath.",
          size=9, italic=True, colour=GREY, after=3)
    n = max(blank_rows, len(rows)) if mode != "assessor" else len(rows)
    t = doc.add_table(rows=0, cols=len(columns))
    hdr = t.add_row().cells
    for i, col in enumerate(columns):
        set_cell_borders(hdr[i]); shade_cell(hdr[i], CREAM)
        r = hdr[i].paragraphs[0].add_run(col); r.bold = True; r.font.size = Pt(9)

    shown = 0 if mode == "assessor" else exemplar
    if mode == "assessor":
        body = rows
    else:
        body = [list(row) if j < shown
                else [(row[i] if i < given else "") for i in range(len(columns))]
                for j, row in enumerate(rows[:n])]
        body += [[""] * len(columns) for _ in range(n - len(body))]

    for j, row in enumerate(body):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_borders(cells[i])
            r = cells[i].paragraphs[0].add_run(str(val)); r.font.size = Pt(9)
            if mode == "assessor":
                r.font.color.rgb = RGBColor.from_string(MODEL)
            elif j < shown:
                r.font.color.rgb = RGBColor.from_string(MODEL); r.italic = True
    doc.add_paragraph()


def box(doc, lines):
    """A bordered drop-zone. `lines` is a list of (text, colour_or_None, bold, italic)."""
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    set_cell_borders(cell); shade_cell(cell, CREAM); cell.width = Cm(16.6)
    for i, (text, colour, bold, italic) in enumerate(lines):
        par = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER if colour is None else WD_ALIGN_PARAGRAPH.LEFT
        r = par.add_run(text)
        r.font.size = Pt(9.5 if colour else 10); r.bold = bold; r.italic = italic
        if colour:
            r.font.color.rgb = RGBColor.from_string(colour)
    doc.add_paragraph()


def image_box(doc, caption, image):
    """A bordered drop-zone holding a real exemplar capture, under its description."""
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    set_cell_borders(cell); shade_cell(cell, CREAM); cell.width = Cm(16.6)
    r = cell.paragraphs[0].add_run(caption)
    r.font.size = Pt(9.5); r.italic = True
    r.font.color.rgb = RGBColor.from_string(CAPTURE)
    pic = cell.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.add_run().add_picture(str(image), width=Cm(IMAGE_CM))
    doc.add_paragraph()


def screenshot_slot(doc, capture, mode, image=None):
    if mode == "assessor":
        if image is not None:
            image_box(doc, f"SCREENSHOT — {capture}", image)
        else:
            box(doc, [(f"SCREENSHOT — {capture}", CAPTURE, False, True)])
    else:
        box(doc, [("[ PASTE YOUR SCREENSHOT HERE ]", None, True, False),
                  (capture, GREY, False, True)])


def evidence_images(evidence_dir, key):
    """The exemplar captures filed for one task or test, in order.

    Convention: <key>.png for a single capture, <key>.<n>.png where the task needed several
    (task 22 takes two screens to show the group across both zones). A task with more captures
    on disk than evidence boxes in the run sheet is fine — the extras follow the last box.
    """
    if evidence_dir is None:
        return []
    return sorted(Path(evidence_dir).glob(f"{key}.*.png")) + \
        sorted(Path(evidence_dir).glob(f"{key}.png"))


def place_evidence(doc, captures, mode, images):
    """Render each evidence box, pairing it with the capture filed for it where there is one."""
    for i, cap in enumerate(captures):
        screenshot_slot(doc, cap, mode, image=images[i] if i < len(images) else None)
    for extra in images[len(captures):]:
        image_box(doc, "SCREENSHOT — continued", extra)


def diagram_slot(doc, what, mode):
    """A drawing area. Works on paper or on screen — sketch it or paste it."""
    if mode == "assessor":
        box(doc, [(f"DIAGRAM — {what}", CAPTURE, False, True)])
    else:
        box(doc, [("[ SKETCH OR PASTE YOUR DIAGRAM HERE ]", None, True, False),
                  (what, GREY, False, True)])


def response_slot(doc, model, mode, points=None):
    """Blank box for the student; for the assessor either a model answer or key points."""
    if mode != "assessor":
        box(doc, [("[ WRITE YOUR ANSWER HERE ]", None, True, False)])
        return
    if points:
        lines = [("Key points the answer should touch on:", MODEL, True, False)]
        lines += [(f"•  {pt}", MODEL, False, False) for pt in points]
        box(doc, lines)
    else:
        box(doc, [(model, MODEL, False, False)])


def flag(doc, text):
    """A small teal label above a heading — the thing a student quotes when they are stuck."""
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(10)
    par.paragraph_format.space_after = Pt(0)
    r = par.add_run(text.upper())
    r.bold = True; r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(TEAL)


def code(doc, lines):
    """A verbatim block — monospace, shaded, for anything the student must copy exactly."""
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    set_cell_borders(cell); shade_cell(cell, CREAM); cell.width = Cm(16.6)
    for i, line in enumerate(lines):
        par = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        par.paragraph_format.space_after = Pt(0)
        r = par.add_run(line)
        r.font.name = "Consolas"; r.font.size = Pt(9)
    doc.add_paragraph()


def note(doc, text):
    """An environment-constraint note — set apart from the task, in the accent colour."""
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Cm(0.4)
    par.paragraph_format.space_before = Pt(4)
    par.paragraph_format.space_after = Pt(8)
    lead = par.add_run("Environment constraint ")
    lead.bold = True; lead.font.size = Pt(9.5)
    lead.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    r = par.add_run(text)
    r.italic = True; r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(TERRACOTTA)


def clicks(doc, items):
    """Click-by-click detail. The practice sheet has this; the assessment deliberately does not."""
    p(doc, "How to do it", bold=True, after=3)
    for i, step in enumerate(items, 1):
        par = doc.add_paragraph()
        par.paragraph_format.left_indent = Cm(0.6)
        par.paragraph_format.space_after = Pt(2)
        lead = par.add_run(f"{i}.  "); lead.bold = True; lead.font.size = Pt(10)
        par.add_run(step).font.size = Pt(10)
    doc.add_paragraph()


def assessor_note(doc, text, mode):
    """Guidance for the assessor only — never rendered in the student copy."""
    if mode != "assessor":
        return
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Cm(0.4)
    par.paragraph_format.space_before = Pt(4)
    par.paragraph_format.space_after = Pt(8)
    lead = par.add_run("Assessor note ")
    lead.bold = True; lead.font.size = Pt(9.5)
    lead.font.color.rgb = RGBColor.from_string(MODEL)
    r = par.add_run(text)
    r.italic = True; r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(MODEL)


def steps(doc, items):
    """Numbered steps. A step is a plain string, or a list of (text, bold) segments where
    part of it needs to stand out."""
    for i, step in enumerate(items, 1):
        par = doc.add_paragraph()
        par.paragraph_format.left_indent = Cm(0.6)
        par.paragraph_format.space_after = Pt(3)
        lead = par.add_run(f"{i}.  "); lead.bold = True; lead.font.size = Pt(BODY_PT)
        for text, bold in ([(step, False)] if isinstance(step, str) else step):
            r = par.add_run(text); r.bold = bold; r.font.size = Pt(BODY_PT)
    doc.add_paragraph()


def heading2(doc, text):
    """A task heading.

    The Kangan template defines Heading 2 as WHITE text — presumably for a coloured band it
    no longer sits on. Left alone, every task title in this document is invisible. The run
    colour is set explicitly here rather than redefining the institutional style.
    """
    par = doc.add_paragraph(text, style="Heading 2")
    for r in par.runs:
        r.font.color.rgb = RGBColor.from_string(TEAL)
    return par


def notes_box(doc):
    """A place for the student to write their own notes — PRACTICE ONLY.

    Rendered after a 'How to do it' block, which only the practice sheet has, so the
    assessment never gets one. A light blue fill so it reads as the student's space rather
    than ours, and a Word table cell, so it grows as they type. The point is that the run
    sheet becomes partly theirs — the thing they worked out at the console is worth more to
    them later than anything we wrote.
    """
    p(doc, "Personal notes", bold=True, size=9.5, after=3)
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    set_cell_borders(cell); shade_cell(cell, NOTES_BG); cell.width = Cm(16.6)
    run = cell.paragraphs[0].add_run("< type any personal notes about this section here >")
    run.italic = True; run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(GREY)
    doc.add_paragraph()


def resources_block(doc, items):
    """'Related resources' — the scenario documents this element actually needs.

    A task that asks a student to record an availability target without saying where the
    target is written makes them hunt for it. The work is making the linkage — this document
    tells me that figure — not the search. Each label says WHY the resource matters here, so
    the student knows what they are opening it for, and links deep to a section anchor where
    the document is long.
    """
    p(doc, "Related resources", bold=True, size=9.5, after=3)
    for label, url in items:
        par = doc.add_paragraph()
        par.paragraph_format.left_indent = Cm(0.6)
        par.paragraph_format.space_after = Pt(3)
        par.add_run("\u2022  ").font.size = Pt(9.5)
        add_hyperlink(par, label, url, size_pt=9.5)
    doc.add_paragraph()


def consider(doc, questions):
    """'Things to consider' — the leading questions on a PRACTICE design task.

    The practice sheet presents each design task exactly as the assessment does, then adds
    these underneath. They are the teaching half: questions pointed enough that a student
    who is awake cannot miss the answer, without ever stating it. The assessment has no
    equivalent — there, the task stands on its own.
    """
    p(doc, "Things to consider", bold=True, size=9.5, after=3)
    for q in questions:
        par = doc.add_paragraph()
        par.paragraph_format.left_indent = Cm(0.6)
        par.paragraph_format.space_after = Pt(3)
        par.add_run("•  ").font.size = Pt(9.5)
        r = par.add_run(q); r.font.size = Pt(9.5); r.italic = True
        r.font.color.rgb = RGBColor.from_string(TEAL)
    doc.add_paragraph()


def uoc_line(doc, items, mode):
    """The assessor-only traceability line under an element's heading."""
    if mode != "assessor":
        return
    p(doc, "Evidences: " + " · ".join(f"[{u}]" for u in items),
      size=9, italic=True, colour=UOC, after=2)


def standard_line(doc, text, mode):
    """What the element is actually marked against — the UoC intent, not the supplied values.

    Every settings table in this instrument contains values we invented so the student has a
    concrete task. This line names which of them are load-bearing. An assessor marks this,
    never the table.
    """
    if mode != "assessor":
        return
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(6)
    lead = par.add_run("Satisfactory when ")
    lead.bold = True; lead.font.size = Pt(9)
    lead.font.color.rgb = RGBColor.from_string(UOC)
    r = par.add_run(text)
    r.italic = True; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(UOC)
