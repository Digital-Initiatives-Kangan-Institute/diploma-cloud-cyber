#!/usr/bin/env python3
"""Build the S1-CL1 AT3 STUDENT workbook (.docx) by populating the Kangan template.

The student-facing half of AT3, rendered from the SAME definition as the assessor copy
(at3_run_sheet.py) with `mode="student"`. Nothing here re-authors the workbook — a change to
a task, a table or a step is made once, in the run sheet, and appears in both instruments on
the next build. The two cannot drift.

What `mode="student"` strips:
  * the `Evidences:` UoC line under every element
  * the `Satisfactory when` line — the standard the element is marked against
  * every model answer: capture tables lose their teal rows, response boxes become
    "[ WRITE YOUR ANSWER HERE ]", screenshot boxes become "[ PASTE YOUR SCREENSHOT HERE ]"
    with the capture description as a grey hint
  * assessor-only notes

What it KEEPS, deliberately: the scaffolding columns. A capture table's `given` count says
how many leading columns are pre-filled for the student too — the requirement names in task
1, the tier list in task 2, the setting names in task 8. Those are part of the question, not
the answer. Where the first column IS the finding the unit asks for (task 3's single points
of failure), `given` is 0 and the student faces a blank grid.

The front matter differs from the assessor copy only in the institutional tables — the
student's Details / Student instructions / Assessment criteria, without the Marking Guide,
the Teacher instructions or the reverse map.

Usage:  python scripts/s1_cl1/build_s1_cl1_at3_student.py [output.docx]
Default: S1-CL1-Cloud-Design-Build/assessments/AT3/AT3-High-Availability-Student.docx
"""
import sys
from pathlib import Path

from docx import Document  # noqa: E402

TEMPLATE = str(Path(__file__).resolve().parents[2] / "kangan-templates" / "Project Assessment - Student.docx")

sys.path.insert(0, str(Path(__file__).resolve().parent))  # noqa: E402
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # noqa: E402
sys.path.insert(0, str(next(d / "scripts" for d in Path(__file__).resolve().parents if (d / "scripts" / "helpers" / "__init__.py").exists())))  # noqa: E402
from helpers.docx_tables import clear_table_rows, find_instruction_row, set_cell_content  # noqa: E402
import at3_run_sheet  # noqa: E402  (the workbook — single source for both instruments)
import build_s1_cl1_at3_assessor as a  # noqa: E402  (shared DETAILS — single source of truth)

# ---------- student-facing content (second person; no UoC traceability) ----------

CHECK = "☐ Yes  ☐ No"

TIME_ALLOWED = [
    "Part A — Design: 4 hours",
    "Part B — Implementation window: 3.5 hours",
    "Knowledge questions and reflection: 1 hour",
    "One continuous worksheet — you work the parts in order, in one go.",
]

OVERVIEW = [
    "You are assessed on designing high-availability improvements to the cloud environment you built "
    "in the previous phase, and on implementing that design during a maintenance window. This is the "
    "third and final assessment task in the Cloud Design and Build cluster, and it closes out the YAT "
    "LMS Cloud Migration engagement.",
    "The assessment is one workbook in two parts, submitted as a single document. Part A leads you "
    "task by task through designing the HA architecture. Part B is a run sheet: you deploy the "
    "supplied baseline, then build your own Part A answers, run the simulations you planned, measure "
    "availability across the window, and close the engagement with feedback and sign-off. There is no "
    "presentation and no observation event.",
    "Part B builds what you designed in Part A. Each build task names the design task it comes from "
    "and asks you to copy your own answer across before you build it — so the infrastructure you end "
    "up with is yours, not ours.",
    "This is an open-book practical assessment. You may use the YAT intranet, AWS Academy lab "
    "environments, AWS documentation, course reference materials, and external research (which must "
    "be cited) throughout.",
    "Your assessor will provide a baseline lab-pack — a template that builds the environment described "
    "at the start of Part A, so everyone starts the maintenance window from the same place. Deploying "
    "it is the first task of Part B and takes 10–15 minutes.",
    "Submission: this completed workbook (.docx), with every task and question answered and every "
    "evidence box populated, submitted via the LMS.",
    "The assessment will not proceed if for any reason it is not safe to do so. Your assessor must "
    "advise you of the reason for suspending the assessment, and what safety action should be taken, "
    "and of revised arrangements when it is safe to resume.",
    "There is zero tolerance for plagiarism, cheating and collusion. You will be expected to make a "
    "declaration that all work is your own prior to submission. Refer to the Training and Assessment "
    "Policy for further information.",
]

TASKS = [
    "The LMS now runs on AWS. YAT IT installed the application on the foundation you built, migrated "
    "the database and ran the cutover; the system has been live for a term. It was never built to "
    "survive a failure, and YAT's 99.9% availability target has not been met. The board has approved "
    "a final phase: harden the environment for high availability.",
    "The task has two parts that combine into a single submission:",
    "Part A — Design. You are given the current architecture and led through eighteen tasks: the "
    "targets your design must meet, a review of the current environment against them, the single "
    "points of failure it contains, the recovery objectives it achieves today, your HA design across "
    "every layer, the points of failure that design removes, the recovery objectives it achieves, the "
    "order you will apply the changes in, and the simulations that will verify them.",
    "Part B — Implementation. You deploy the supplied baseline lab-pack, then build your own design "
    "during a simulated Saturday late-night maintenance window of about 3.5 hours. You then run the "
    "failure and resize simulations you planned, measure availability across the window, compare what "
    "happened against what you predicted, record any adjustments, get feedback and sign-off from the "
    "YAT ICT Manager, and file the workbook per YAT's records procedures.",
    "The knowledge questions and the reflection come after the window and are not done under time "
    "pressure.",
    "Scope: cloud infrastructure only. Re-validating the LMS application against the hardened "
    "infrastructure, and any application-layer tuning, are YAT in-house IT's responsibility. Your "
    "deliverable stops at HA-hardened infrastructure handed back to YAT IT.",
]

RESOURCES = [
    "Provided to you",
    "The baseline lab-pack — ask your assessor where to download it from. You deploy it as the first "
    "task of Part B",
    "The YAT intranet — the LMS Application Specification, the LMS Cloud Migration Requirements, the "
    "network diagram and the Records Management Policy. All are linked in this document where you "
    "need them",
    "AWS Academy lab access — Cloud Foundations [104469] + Cloud Architecting [172221]",
    "You supply",
    "Computer with web browser",
    "Word-processing software (Microsoft Word or equivalent)",
    "A screenshot tool",
]

CRITERIA = [
    "To receive a Satisfactory result for this assessment you must:",
    "Achieve Satisfactory on every criterion in the Assessment Criteria table",
    "Submit this completed workbook (.docx) with every task and question answered and every evidence "
    "box populated",
]

# Assessment criteria, student voice — the same 19 the assessor marks, without the UoC tags.
CRITERIA_TABLE = [
    "A1 Requirements and review (tasks 1–2) — you established the availability, recovery and "
    "service-level targets from the supplied documents, then reviewed the current environment tier by "
    "tier against those targets",
    "A2 Current-state analysis (tasks 3–5) — you identified the single points of failure, quantified "
    "the recovery objectives the environment achieves today, and identified what can only scale "
    "vertically and what that costs in availability",
    "A3 Review findings documented (task 6) — you summarised the gap between the current environment "
    "and the targets, written for the YAT ICT Manager",
    "A4 HA architecture designed (tasks 7–12) — you designed the HA architecture across every layer: "
    "an application subnet in the second zone, an Auto Scaling group that survives losing a zone, a "
    "database with automatic failover, an outbound path that does not depend on one zone, monitoring "
    "that would reveal a partial loss of availability, and a reasoned decision about the load balancer",
    "A5 Designed-state analysis (tasks 13–15) — you accounted for every point of failure you found in "
    "task 3, quantified the recovery objectives your design achieves, and stated what still scales "
    "vertically",
    "A6 Design complete and coherent (task 16) — your design reads as one consistent document that "
    "addresses YAT's business needs and could be built from",
    "A7 Implementation sequence (task 17) — you sequenced the changes with a duration, service impact, "
    "verification and rollback for each, and reconciled the total against the window",
    "A8 Simulation plan and findings (tasks 18 and 25) — you planned failure and resize simulations "
    "with a stated expected outcome for each, then compared what actually happened against those "
    "predictions",
    "A9 Implementation (tasks 19–23) — you deployed the supplied baseline, then built your own design: "
    "the second-zone application subnet, its outbound path, the Auto Scaling group across both zones, "
    "and the database failover configuration",
    "A10 Monitoring and availability measurement (task 24, T4) — you built the HA monitoring you "
    "designed, and measured and recorded the LMS's availability across the window with your working "
    "shown",
    "A11 Connectivity (T1) — you demonstrated connectivity at every tier and across both zones, "
    "including the negative test that the database is not reachable from the internet",
    "A12 Failure simulation (T2) — you executed a real failure against your own environment and "
    "recorded the outcome with timings",
    "A13 Resize simulation (T3) — you executed a resize and measured its availability impact from "
    "observation",
    "A14 Adjustments (task 26) — you documented what you changed as a result of the simulations, or "
    "justified from evidence that nothing needed changing",
    "A15 Feedback (task 27) — you took the completed work to the YAT ICT Manager and recorded the "
    "feedback and your response to it",
    "A16 Sign-off (task 28) — the YAT ICT Manager has accepted the work and the sign-off block is "
    "completed and signed",
    "A17 Filing (task 29) — you named where you filed the completed workbook and which YAT policy "
    "required that location",
    "A18 Knowledge questions (Q1–Q6) — you answered all six with reference to your own design and your "
    "own build, in clear written English",
    "A19 Reflection (R1–R3) — three honest reflective responses, all specific to your own work",
]

RESULTS = [
    "If you are deemed not satisfactory, you will be given one (1) more attempt at this assessment or "
    "your teacher/assessor will negotiate a further assessment with you. The second attempt must be "
    "completed within 10 working days from the date your feedback is given.",
]


def add_criterion_row(table, text):
    row = table.add_row()
    set_cell_content(row.cells[0], text)
    set_cell_content(row.cells[1], CHECK)


def delete_body_paragraph(doc, text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            p._element.getparent().remove(p._element)
            return


def build(path):
    doc = Document(TEMPLATE)

    t_details = doc.tables[0]
    set_cell_content(t_details.rows[1].cells[1], a.DETAILS["qualification"])
    set_cell_content(t_details.rows[2].cells[1], a.DETAILS["units"])
    set_cell_content(t_details.rows[3].cells[1], a.DETAILS["task_title"])
    set_cell_content(t_details.rows[4].cells[1], a.DETAILS["task_number"])

    t_instr = doc.tables[1]
    set_cell_content(find_instruction_row(t_instr, "Assessment overview"), OVERVIEW)
    set_cell_content(find_instruction_row(t_instr, "Task"), TASKS)
    set_cell_content(find_instruction_row(t_instr, "Time allowed"), TIME_ALLOWED)
    set_cell_content(find_instruction_row(t_instr, "Location"), "")
    set_cell_content(find_instruction_row(t_instr, "Resources required"), RESOURCES)
    set_cell_content(find_instruction_row(t_instr, "Assessment criteria"), CRITERIA)
    row = find_instruction_row(t_instr, "Results")
    if row is not None:
        set_cell_content(row, RESULTS)

    # ---- Assessment criteria table (student voice; no UoC tags) ----
    t_crit = doc.tables[2]
    clear_table_rows(t_crit, 2)
    for text in CRITERIA_TABLE:
        add_criterion_row(t_crit, text)

    delete_body_paragraph(doc, "Add or delete rows as required")

    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h2 = lambda t: at3_run_sheet.R.heading2(doc, t)
    at3_run_sheet.render_front_matter(doc, h1)
    at3_run_sheet.render(doc, h1, h2, mode="student")

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL1-Cloud-Design-Build/assessments/AT3/AT3-High-Availability-Student.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
