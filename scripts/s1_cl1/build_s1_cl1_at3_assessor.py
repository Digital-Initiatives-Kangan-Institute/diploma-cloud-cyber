#!/usr/bin/env python3
"""Build the S1-CL1 AT3 ASSESSOR workbook (.docx) by populating the Kangan template.

AT3 as a guided design-and-build workbook, in the same form as AT2: the institutional Kangan
'Project Assessment - Assessor' template filled in, then the workbook itself rendered worked —
every capture table carrying its model answer, every screenshot box describing what should be
in it, every element tagged with the UoC items it evidences AND the standard it is marked
against.

  Part A  the student designs the HA-equivalent of what AT2 built. The tasks are given;
          findings and design decisions are not. Part A elements carry performance criteria, so
          they are framed as tasks; only the knowledge section asks questions.
  Part B  a run sheet whose tasks copy forward the student's own Part A answers.

THE MARKING MODEL, and the thing that separates this from the report it replaces: the values
in this document are ours, invented so there is a concrete task to perform. The unit's wording
is vague. So each element carries `Evidences:` (which UoC items) and `Satisfactory when` (what
must be true for them to be met). An assessor marks the second. A student who sets a minimum
of three where we wrote two has met the item; one who leaves it at one has not — not because
it differs from ours, but because one instance in one zone is not fault tolerant.

Content lives in at3_run_sheet.py — ONE definition, rendered here worked and in
build_s1_cl1_at3_student.py blank.

Usage:  python scripts/s1_cl1/build_s1_cl1_at3_assessor.py [output.docx]
Default: S1-CL1-Cloud-Design-Build/assessments/AT3/AT3-High-Availability-Assessor.docx
"""
import sys
from pathlib import Path

from docx import Document  # noqa: E402

TEMPLATE = str(Path(__file__).resolve().parents[2] / "kangan-templates" / "Project Assessment - Assessor.docx")

sys.path.insert(0, str(Path(__file__).resolve().parent))  # noqa: E402
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # noqa: E402
sys.path.insert(0, str(next(d / "scripts" for d in Path(__file__).resolve().parents if (d / "scripts" / "helpers" / "__init__.py").exists())))  # noqa: E402
from helpers.instrument_layout import render_benchmark  # noqa: E402
from helpers.docx_tables import clear_table_rows, find_instruction_row, set_cell_content  # noqa: E402
import at3_run_sheet  # noqa: E402  (the workbook — single source for both instruments)

# ---------- content ----------

CHECK = "☐ Yes  ☐ No"

DETAILS = {
    "qualification": "ICT50220 Diploma of Information Technology",
    "units": [
        "ICTICT517 Match ICT needs with the strategic direction of the organisation",
        "ICTCLD401 Configure cloud services",
        "ICTCLD502 Design and implement highly-available cloud infrastructure",
    ],
    "task_title": "AT3 — High Availability: YAT LMS Migration",
    "task_number": "3 of 3",
}

TIME_ALLOWED = [
    "Part A — Design: 4 hours",
    "Part B — Implementation window: 3.5 hours",
    "Knowledge questions and reflection: 1 hour",
    "Part A is marked before Part B begins. A student whose design will not work is told so and "
    "corrects it before building — see Assessment overview.",
]

OVERVIEW = [
    "Students are assessed on designing high-availability improvements to the cloud environment they "
    "built in AT2, and implementing that design during a maintenance window. AT3 is the third and "
    "final assessment task in the S1-CL1 Cloud Design and Build cluster — it closes out the YAT LMS "
    "Cloud Migration engagement.",
    "The assessment is one guided workbook in two parts, submitted as a single document. Part A leads "
    "the student task by task through designing the HA-equivalent architecture. Part B is a "
    "run sheet: the student deploys a supplied baseline, then builds their own Part A answers, runs "
    "failure and resize simulations, measures availability across the window, and closes the "
    "engagement with feedback and sign-off. There is no presentation or observation event.",
    "PART A IS MARKED BEFORE PART B BEGINS. This is not optional sequencing. A student whose design "
    "would not deliver high availability — most commonly an Auto Scaling minimum left at one — is told "
    "so, corrects the design, and then implements the corrected version. Part B builds the student's "
    "own design; without the Part A checkpoint, a student can spend the whole window faithfully "
    "building something that cannot meet the requirement.",
    "WHAT IS BEING MARKED. Every settings table and worked answer in this document contains values we "
    "chose so that there is a concrete task to perform — the unit's wording is deliberately general. "
    "Each element therefore carries two assessor-only lines: 'Evidences', naming the UoC items, and "
    "'Satisfactory when', naming what has to be true for those items to be met. Mark the second. A "
    "student whose values differ from ours but who meets the stated standard is Satisfactory; a "
    "student who matches our values exactly but misses the standard is not.",
    "This is an open-book practical assessment. Students may use the YAT intranet, AWS Academy lab "
    "environments, AWS documentation, course reference materials, and external research (which must be "
    "cited) throughout.",
    "BASELINE STARTING STATE. So that every student starts the maintenance window from the same "
    "architecture regardless of what they personally completed in AT2, the assessor provides a "
    "baseline lab-pack — a CloudFormation template reproducing the end state of the AT2 run sheet. "
    "How it reaches students is the assessor's choice (LMS, Teams, SharePoint); the requirement is "
    "only that it is available to them at the start of Part B. Deploying it is Part B task 19 and takes "
    "10–15 minutes.",
    "Reasonable adjustment for this assessment may include extending either the design time or the "
    "maintenance window, providing one-on-one verbal explanation of the current architecture, allowing "
    "alternative screenshot-evidence formats where assistive technology requires it, or splitting the "
    "work across more sittings.",
    "Teacher/assessor support level: the assessor may clarify what a task is asking, the scenario "
    "context, and the current architecture, but must not identify single points of failure for the "
    "student, supply design values, or confirm whether a design decision is correct before Part A is "
    "marked. The design is the student's own work.",
    "Submission: the completed workbook (.docx) with every task and question answered and every evidence box "
    "populated, submitted via the LMS.",
    "The assessment will not proceed if for any reason it is not safe to do so. You must advise the "
    "student of the reason for suspending the assessment, and what safety action should be taken. "
    "Advise the student of revised arrangements for the assessment when it is safe to do so.",
    "There is a zero tolerance for plagiarism, cheating and collusion. Students will be expected to "
    "make a declaration that all work is their own prior to submission. Refer to the Training and "
    "Assessment Policy for further information.",
]

TASKS = [
    "The LMS now runs on AWS. YAT IT installed the application on the foundation the student built in "
    "AT2, migrated the database and ran the cutover; the system has been live for a term. It was never "
    "built to survive a failure, and YAT's 99.9% availability target has not been met. The board has "
    "approved a final phase: harden the environment for high availability.",
    "The task has two parts that combine into a single submission:",
    "Part A — Design. The student is given the current architecture and led through eighteen tasks: "
    "the targets the design must meet, a review of the current environment against them, the single "
    "points of failure it contains, the recovery objectives it achieves today, the HA-equivalent design "
    "across every layer, the points of failure that design removes, the recovery objectives it "
    "achieves, the order the changes will be applied in, and the simulations that will verify them.",
    "Part B — Implementation. The student deploys the assessor-provided baseline lab-pack, then builds "
    "their own design during a simulated Saturday late-night maintenance window of approximately 3.5 "
    "hours. Each task names the Part A question whose answer it builds and asks the student to copy "
    "that answer across first. The student then runs the failure and resize simulations they planned, "
    "measures availability across the window, compares what happened against what they predicted, "
    "records any adjustments, secures feedback and sign-off from the role-played YAT ICT Manager, and "
    "files the workbook per YAT's records procedures.",
    "Knowledge questions and reflection follow the window and are not performed under time pressure.",
    "MTS scope: cloud infrastructure only. Re-validating the LMS application against the hardened "
    "infrastructure, and any application-layer tuning, remain YAT in-house IT's responsibility and are "
    "out of scope. The deliverable stops at HA-hardened infrastructure handed back to YAT IT.",
]

RESOURCES = [
    "Teacher/assessor supplied resources",
    "The baseline lab-pack — a CloudFormation template reproducing the AT2 end state. THE ASSESSOR MUST "
    "PROVIDE THIS to students at the start of Part B; the delivery channel is the assessor's choice",
    "Access to the YAT intranet — supplying the LMS Application Specification, the LMS Cloud Migration "
    "Requirements and the Records Management Policy",
    "Those documents are linked in the Instructions to Student below",
    "A person to role-play Sam Walker, YAT ICT Manager, for the feedback and sign-off at the end of "
    "Part B — normally the assessor",
    "AWS Academy lab access — Cloud Foundations [104469] + Cloud Architecting [172221]",
    "The worked workbook later in this document — model answers, screenshot descriptions, per-element "
    "UoC mapping, and the standard each element is marked against",
    "Student supplied resources",
    "Computer with web browser",
    "Word-processing software (e.g. Microsoft Word or equivalent)",
    "A screenshot tool",
]

CRITERIA = [
    "To receive a Satisfactory outcome for this assessment the student must:",
    "Achieve Satisfactory on every criterion in the Marking Guide below",
    "Submit the completed workbook (.docx) with every task and question answered and every evidence box populated",
]

CONDITIONS = [
    "These are conditions the assessor verifies as present before marking begins. They are not "
    "student-performance criteria — they are the conditions under which the assessment can validly be "
    "conducted.",
    "C1 Lab environment is accessible to the student throughout the assessment — AWS Academy Cloud "
    "Foundations [104469] + AWS Academy Cloud Architecting [172221] — providing cloud vendor service "
    "provider access, cloud managed database service (RDS), IDE / console / CLI / SSH-RDP tooling, and "
    "internet/web browser access",
    "[ICTCLD401 AC 1] · [ICTCLD401 AC 2] · [ICTCLD401 AC 3] · [ICTCLD502 AC 1] · [ICTCLD502 AC 2] · "
    "[ICTCLD502 AC 4] · [ICTCLD502 AC 6] · [ICTCLD502 AC 7]",
    "C2 The YAT scenario site / intranet is accessible to the student throughout the assessment — "
    "supplying the LMS Application Specification, the LMS Cloud Migration Requirements and the "
    "organisational policies",
    "[ICTCLD401 AC 4] · [ICTCLD502 AC 3] · [ICTCLD502 AC 5] · [ICTCLD502 AC 8]",
    "C3 The baseline lab-pack has been made available to the student before Part B begins",
    "C4 A person is available to role-play Sam Walker, YAT ICT Manager, for the feedback and sign-off "
    "at the close of Part B",
    "[ICTICT517 AC 4]",
]

# Marking Guide — A1-A19. Each criterion is [statement, UoC-traceability line]. Criteria group the
# workbook's elements; each element also carries its own Evidences line where the work is done.
MARKING = [
    ["A1 Requirements and review (tasks 1–2) — student establishes the availability, recovery and "
     "service-level targets from the supplied documents, then reviews the current environment tier by "
     "tier against those targets rather than against general good practice",
     "[ICTCLD502 PC 1.1] · [ICTCLD502 PC 2.1] · [ICTCLD502 FS Reading] · [ICTCLD401 FS Reading]"],

    ["A2 Current-state analysis (tasks 3–5) — student identifies the single points of failure in the "
     "environment, quantifies the recovery objectives it achieves today, and identifies the components "
     "that can only scale vertically together with the availability cost of doing so",
     "[ICTCLD502 PC 2.2] · [ICTCLD502 PC 2.3] · [ICTCLD502 PC 2.4]"],

    ["A3 Review findings documented (task 6) — student summarises the gap between the current environment "
     "and the targets, written for the YAT ICT Manager rather than for another engineer",
     "[ICTCLD502 PC 2.5]"],

    ["A4 HA architecture designed (tasks 7–12) — student designs the HA-equivalent architecture across "
     "every layer: an application subnet in the second zone, an Auto Scaling group spanning both zones "
     "with a minimum that survives losing one, a database with automatic failover, an outbound path "
     "that does not depend on a single zone, monitoring that would reveal a partial loss of "
     "availability, and a reasoned decision about the load balancer",
     "[ICTCLD502 PC 3.1] · [ICTCLD502 PE 1] (design half) · [ICTCLD502 PE 2] (design half) · "
     "[ICTCLD502 PE 5] (design half)"],

    ["A5 Designed-state analysis (tasks 13–15) — student accounts for every point of failure found in Q3, "
     "quantifies the recovery objectives the design achieves, and states what still scales vertically "
     "and at what cost",
     "[ICTCLD502 PC 3.2] · [ICTCLD502 PC 3.3] · [ICTCLD502 PC 3.4]"],

    ["A6 Design complete and coherent (task 16) — the design reads as one internally consistent document "
     "that addresses YAT's stated business needs and could be built from",
     "[ICTCLD502 PC 3.5]"],

    ["A7 Implementation sequence (task 17) — student sequences the changes with a duration, service impact, "
     "verification and rollback for each, and reconciles the total against the maintenance window",
     "[ICTCLD401 FS Planning and organising]"],

    ["A8 Simulation plan and findings (tasks 18 and 25) — student plans failure and resize simulations with a "
     "stated expected outcome for each, then compares what actually happened against those predictions "
     "and explains any divergence",
     "[ICTCLD502 PC 4.6]"],

    ["A9 Implementation (tasks 19–23) — student deploys the supplied baseline, then builds their own "
     "design: the second-zone application subnet, its outbound path, the Auto Scaling group across both "
     "zones, and the database failover configuration. Each task is built to the student's own Part A "
     "answer, copied forward",
     "[ICTCLD502 PC 4.1] · [ICTCLD502 PE 1] (implement half) · [ICTCLD502 PE 2] (deploy half) · "
     "[ICTCLD502 PE 4]"],

    ["A10 Monitoring and availability measurement (task 24, T4) — student builds the HA monitoring they "
     "designed, and defines, measures and records the LMS's availability across the maintenance window "
     "from real data, with the calculation shown",
     "[ICTCLD502 PC 4.3] · [ICTCLD502 PE 5] (measurement half)"],

    ["A11 Connectivity (T1) — student demonstrates connectivity at every tier and across both zones, "
     "including the negative test that the database remains unreachable from the internet",
     "[ICTCLD502 PC 4.2]"],

    ["A12 Failure simulation (T2) — student executes a real failure against their own environment, "
     "records the outcome with timings, and shows whether the service survived it",
     "[ICTCLD502 PC 4.4] · [ICTCLD502 PE 3]"],

    ["A13 Resize simulation (T3) — student executes a resize and measures its availability impact from "
     "observation rather than estimating it",
     "[ICTCLD502 PC 4.5]"],

    ["A14 Adjustments (task 26) — student documents what they changed as a result of the simulations, or "
     "justifies from evidence that nothing needed changing",
     "[ICTCLD502 PC 5.1]"],

    ["A15 Feedback (task 27) — student takes the completed work to the role-played YAT ICT Manager, records "
     "the feedback received and their response to it",
     "[ICTCLD502 PC 5.2]"],

    ["A16 Sign-off (task 28) — the role-played YAT ICT Manager has accepted the work and the sign-off block "
     "is completed and signed",
     "[ICTCLD502 PC 5.3]"],

    ["A17 Filing (task 29) — student names where the completed workbook was filed and which YAT policy "
     "required that location",
     "[ICTCLD401 PC 4.3]"],

    ["A18 Knowledge questions (Q1–Q6) — student answers all six with reference to their own design and "
     "their own build, not generically, in clear written English",
     "[ICTCLD502 KE 4] · [ICTCLD502 KE 5] · [ICTCLD502 KE 6] · [ICTCLD502 KE 7] · [ICTCLD502 KE 8] · "
     "[ICTCLD502 KE 9] · [ICTCLD401 FS Writing]"],

    ["A19 Reflection (R1–R3) — three honest reflective responses: decisions in hindsight, working under "
     "time pressure, and a lesson carried forward — all specific to the student's own work",
     "[ICTCLD401 FS Learning] · [ICTCLD401 FS Self-management skills] · [ICTCLD502 FS Problem solving] · "
     "[ICTCLD502 FS Self-management]"],
]

ASSESSOR_BODY = [
    ("h1", "Marking Benchmark — UoC traceability (reverse map)"),
    ("p", "This table closes the loop on bidirectional traceability: every UoC requirement AT3 claims "
          "to evidence is named below with the marking criterion that evidences it, and each workbook "
          "element carries its own Evidences line where the work is done. No UoC requirement claimed by "
          "AT3 is left without a criterion."),
    ("p", "ICTCLD502 — Design and implement highly-available cloud infrastructure (AT3-evidenced items)"),
    ("tbl", [
        ["UoC item", "Evidenced by criterion(ia)"],
        ["[ICTCLD502 PC 1.1] Determine reliability, recoverability and service levels required for application", "A1"],
        ["[ICTCLD502 PC 2.1] Review architecture of traditional multi-tier web application and identify HA requirements", "A1"],
        ["[ICTCLD502 PC 2.2] Identify any single points of failure", "A2"],
        ["[ICTCLD502 PC 2.3] Estimate recovery objectives for components and overall architecture", "A2"],
        ["[ICTCLD502 PC 2.4] Determine components that must scale vertically and the impact on availability", "A2"],
        ["[ICTCLD502 PC 2.5] Document architecture review findings according to business needs", "A3"],
        ["[ICTCLD502 PC 3.1] Design equivalent architecture for high availability using cloud services", "A4"],
        ["[ICTCLD502 PC 3.2] Identify and remove single points of failure as required", "A5"],
        ["[ICTCLD502 PC 3.3] Estimate recovery objectives for each component and overall architecture (designed)", "A5"],
        ["[ICTCLD502 PC 3.4] Determine components that must scale vertically and the impact on availability (designed)", "A5"],
        ["[ICTCLD502 PC 3.5] Document architecture design according to business needs", "A6"],
        ["[ICTCLD502 PC 4.1] Implement architecture design in cloud environment", "A9"],
        ["[ICTCLD502 PC 4.2] Demonstrate connectivity between resources at all tiers", "A11"],
        ["[ICTCLD502 PC 4.3] Monitor and measure availability of resources", "A10"],
        ["[ICTCLD502 PC 4.4] Simulate failures of component and confirm infrastructure is fault tolerant", "A12"],
        ["[ICTCLD502 PC 4.5] Simulate resizing components and measure availability impact", "A13"],
        ["[ICTCLD502 PC 4.6] Compare and document simulation findings according to documented design", "A8"],
        ["[ICTCLD502 PC 5.1] Adjust and improve availability of architecture according to simulations", "A14"],
        ["[ICTCLD502 PC 5.2] Confirm, seek and respond to feedback with required personnel", "A15"],
        ["[ICTCLD502 PC 5.3] Obtain final sign off from required personnel", "A16"],
        ["[ICTCLD502 PE 1] Design and implement at least one fault tolerant cloud infrastructure", "A4 (design) · A9 (implement)"],
        ["[ICTCLD502 PE 2] Design and deploy automated infrastructure scaling for at least one business need", "A4 (design) · A9 (deploy)"],
        ["[ICTCLD502 PE 3] Simulate failures of at least one component and demonstrate fault tolerance", "A12"],
        ["[ICTCLD502 PE 4] Use cloud management console, software development kits or command line tools", "A9"],
        ["[ICTCLD502 PE 5] Define, monitor and record resource availability in cloud environment", "A4 (design) · A10 (measure)"],
        ["[ICTCLD502 KE 4] HA concepts — fault tolerance, SPOFs, MTTF/MTTR/MTBF, RPO/RTO, SLAs, scalability", "A18 (Q1)"],
        ["[ICTCLD502 KE 5] Testing and debugging techniques, including avoiding single point failures", "A18 (Q2)"],
        ["[ICTCLD502 KE 6] Tools and techniques to measure availability impact", "A18 (Q3)"],
        ["[ICTCLD502 KE 7] Built-in fault tolerance versus infrastructure designed for fault tolerance", "A18 (Q4)"],
        ["[ICTCLD502 KE 8] Purpose and features of load balancing and autoscaling for availability", "A18 (Q5)"],
        ["[ICTCLD502 KE 9] Techniques, methods and metrics used to monitor cloud resource performance", "A18 (Q6)"],
        ["[ICTCLD502 FS Problem solving]", "A19"],
        ["[ICTCLD502 FS Reading]", "A1"],
        ["[ICTCLD502 FS Self-management]", "A19"],
        ["[ICTCLD502 AC 1] Cloud vendor service provider", "C1 (pre-condition)"],
        ["[ICTCLD502 AC 2] Cloud managed database service", "C1 (pre-condition)"],
        ["[ICTCLD502 AC 3] Information and data sources required to design and implement cloud infrastructure", "C2 (pre-condition)"],
        ["[ICTCLD502 AC 4] Integrated development environment (IDE)", "C1 (pre-condition — AWS Console + CLI counts as IDE)"],
        ["[ICTCLD502 AC 5] Specific requirements + industry standards + organisational procedures + legislative requirements", "C2 (pre-condition — supplied via YAT intranet)"],
        ["[ICTCLD502 AC 6] Internet and web browser", "C1 (pre-condition)"],
        ["[ICTCLD502 AC 7] Secure shell (SSH) or remote desktop protocol (RDP) client", "C1 (pre-condition)"],
        ["[ICTCLD502 AC 8] Data to gather information from to determine output and user requirements", "C2 (pre-condition)"],
    ]),
    ("p", "(Note: 502 PC 1.2 and PC 5.2's AT1 half are evidenced in AT1; PC 1.3 and PC 4.1–4.3's "
          "foundation half are evidenced in AT2. 502 KE 1–3 are evidenced in AT1.)"),
    ("p", "ICTCLD401 — Configure cloud services (AT3-evidenced items only)"),
    ("tbl", [
        ["UoC item", "Evidenced by criterion(ia)"],
        ["[ICTCLD401 PC 4.3] Save and store user documentation according to organisational policies", "A17"],
        ["[ICTCLD401 FS Learning]", "A19"],
        ["[ICTCLD401 FS Planning and organising]", "A7"],
        ["[ICTCLD401 FS Reading]", "A1"],
        ["[ICTCLD401 FS Self-management skills]", "A19"],
        ["[ICTCLD401 FS Writing]", "A18"],
        ["[ICTCLD401 AC 1] Cloud vendor service provider", "C1 (pre-condition)"],
        ["[ICTCLD401 AC 2] Cloud managed database service", "C1 (pre-condition)"],
        ["[ICTCLD401 AC 3] Internet and web browser", "C1 (pre-condition)"],
        ["[ICTCLD401 AC 4] Data to gather information from to determine output and user requirements", "C2 (pre-condition)"],
    ]),
    ("p", "(Note: 401 PC 4.1 — document and communicate work to required personnel — is evidenced in "
          "AT1, through the business-case presentation event and its sign-off. AT3 does not claim it.)"),
    ("p", "ICTICT517 — Match ICT needs with the strategic direction of the organisation"),
    ("tbl", [
        ["UoC item", "Evidenced by criterion(ia)"],
        ["—", "AT3 does not evidence any 517 items directly. 517 is AT1. [ICTICT517 AC 4] (a superior "
              "to role-play with) is discharged here as pre-condition C4, for the closure sign-off."],
    ]),
]

# ---------- build ----------

STYLE = {"h1": "Heading 1", "h2": "Heading 2", "p": "Normal"}


def add_marking_row(table, lines):
    row = table.add_row()
    set_cell_content(row.cells[0], lines)
    set_cell_content(row.cells[1], CHECK)


def delete_body_paragraph(doc, text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            p._element.getparent().remove(p._element)
            return


def render_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    for r, cells in enumerate(rows):
        for c, val in enumerate(cells):
            set_cell_content(t.rows[r].cells[c], val)
    doc.add_paragraph()


def build(path):
    doc = Document(TEMPLATE)

    t_details = doc.tables[0]
    set_cell_content(t_details.rows[1].cells[1], DETAILS["qualification"])
    set_cell_content(t_details.rows[2].cells[1], DETAILS["units"])
    set_cell_content(t_details.rows[3].cells[1], DETAILS["task_title"])
    set_cell_content(t_details.rows[4].cells[1], DETAILS["task_number"])

    t_instr = doc.tables[1]
    set_cell_content(find_instruction_row(t_instr, "Assessment overview"), OVERVIEW)
    set_cell_content(find_instruction_row(t_instr, "Task"), TASKS)
    set_cell_content(find_instruction_row(t_instr, "Time allowed"), TIME_ALLOWED)
    set_cell_content(find_instruction_row(t_instr, "Location"), "")
    set_cell_content(find_instruction_row(t_instr, "Resources required"), RESOURCES)
    set_cell_content(find_instruction_row(t_instr, "Assessment criteria"), CRITERIA)
    cond_row = t_instr.add_row()
    set_cell_content(cond_row.cells[0], "Assessment Conditions & Setup Requirements")
    for r in cond_row.cells[0].paragraphs[0].runs:
        r.bold = True
    set_cell_content(cond_row.cells[1], CONDITIONS)

    t_mark = doc.tables[2]
    clear_table_rows(t_mark, 2)
    for lines in MARKING:
        add_marking_row(t_mark, lines)

    delete_body_paragraph(doc, "Add or delete rows as required")
    delete_body_paragraph(doc, "If questioning or observation is incorporated into this assessment task, "
                               "you can incorporate a Practical Observation Checklist.")

    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h2 = lambda t: doc.add_paragraph(t, style="Heading 2")
    at3_run_sheet.render_front_matter(doc, h1)
    at3_run_sheet.render(doc, h1, h2, mode="assessor")

    render_benchmark(doc, ASSESSOR_BODY, render_table, STYLE)

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL1-Cloud-Design-Build/assessments/AT3/AT3-High-Availability-Assessor.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
