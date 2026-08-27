#!/usr/bin/env python3
"""Build the S1-CL1 AT2 PRACTICE build run sheet (.docx).

A practice exercise, not an assessment: a branded YAT/MTS document with no Kangan wrapper,
no marking criteria and no UoC tags. Same shape as the AT2 assessment run sheet, on the
Ledgerline accounting system rather than the LMS, with different addresses, a Linux server
and a PostgreSQL database — and with the click-by-click detail the assessment withholds.

Usage:  python scripts/s1_cl1/build_s1_cl1_at2_practice.py [output.docx]
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # noqa: E402
sys.path.insert(0, str(next(d / "scripts" for d in Path(__file__).resolve().parents if (d / "scripts" / "helpers" / "__init__.py").exists())))  # noqa: E402

import at2_practice_run_sheet as content  # noqa: E402
from at2_run_sheet import render_run_sheet, _p  # noqa: E402
from helpers.docx_styling import paragraph_bottom_rule  # noqa: E402
from helpers.scenario_document import build_header_footer, configure_styles, wordmark  # noqa: E402
from brand import ADDRESS, GREY, TEAL  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def build(path):
    doc = Document()
    configure_styles(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    build_header_footer(sec)

    # ---- cover ----
    wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(GREY)
    paragraph_bottom_rule(doc.add_paragraph(), TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Practice Build")
    sub = doc.add_paragraph().add_run("Ledgerline migration — build run sheet")
    sub.font.size = Pt(15); sub.italic = True; sub.font.color.rgb = RGBColor.from_string(GREY)

    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h2 = lambda t: doc.add_paragraph(t, style="Heading 2")

    h1("The engagement")
    for para in content.SCENARIO:
        _p(doc, para, after=8)
    h1("How to work through this")
    for para in content.INSTRUCTIONS:
        _p(doc, para, after=8)

    render_run_sheet(doc, h1, h2, mode="student",
                     tasks=content.TASKS, tests=content.TESTS,
                     questions=[], handover=False, region_note=content.REGION_NOTE,
                     notes=True)

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL1-Cloud-Design-Build/delivery/practice/AT2-Practice-Build-Run-Sheet.docx"
    build(sys.argv[1] if len(sys.argv) > 1 else default)
