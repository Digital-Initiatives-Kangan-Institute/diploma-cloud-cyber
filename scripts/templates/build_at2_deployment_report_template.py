#!/usr/bin/env python3
"""Build the YAT / MTS Deployment Report template for the LMS foundation build (.docx).

A deployment-report template scoped to ONE deployment — the YAT LMS foundation build — holding
exactly and only the sections that deployment requires. Split from the generic superset template
(build_deployment_report_template.py), which stays in place for other deployments.

What a foundation build does not have, and so is not in here: a maintenance window (§3.1),
cross-Region backup (§4.9), failure/resize simulation, availability measurement and the findings
and adjustments that follow from them (§6.5-§6.9), a stakeholder feedback record, a sign-off
block, an appendix of configuration exports, and an appendix of reflections. The feedback record and
sign-off go because no AT2 marking criterion looks at them; the exports go because the Appendix A
screenshots already evidence the same performance item; the reflections go because AT3's reflections
appendix carries every foundation skill this one did, and the engagement ends there. Nothing carries an "Applicability / mark Not applicable" note — every
section in this template is a section the report needs.

[TBD - needs discussion: the Appendix A and C evidence lists. The AT2 marking criteria name counts
(17 screenshots, 6 test-evidence items) that no artefact enumerates and that no version of this
template has ever carried. The lists are to be derived from the review of what the assessment
actually asks the student to do, once that review is complete; until then the appendices carry the
generic examples inherited from the superset template.

The rule the two lists are built on comes from the UoC wording of what each appendix evidences:

  Appendix A (criterion A9)  -> [ICTCLD401 PE 1] "build ... virtual network"
                                [ICTCLD401 PE 2] "configure compute, storage, database and
                                                  autoscaling resources"
                                = CONFIGURATION STATE. The settings, as set.

  Appendix C (criterion A11) -> [ICTCLD401 PE 3] "conduct simple tests to confirm access"
                                [ICTCLD502 PC 4.2] "demonstrate connectivity between resources
                                                    at all tiers"
                                [ICTCLD502 PC 4.3] "monitor and measure availability"
                                = TEST OUTCOMES. What happened when it was exercised.

Nothing appears in both. The exemplar currently files one screenshot twice - "ALB target group
reporting Healthy" as A4 and again as C1 - which is the failure this rule prevents: a target group
reporting healthy is a health-check RESULT, so it is Appendix C; Appendix A carries the ALB's
listener / target-group / health-check configuration instead.]

Usage:  python scripts/templates/build_at2_deployment_report_template.py
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # content-repo scripts/ (brand + registry)  # noqa: E402
sys.path.insert(0, str(next(d / "scripts" for d in Path(__file__).resolve().parents if (d / "scripts" / "helpers" / "__init__.py").exists())))  # umbrella scripts/ (engine)  # noqa: E402
from helpers.docx_body_text import add_guidance_text, add_response_placeholder  # noqa: E402
from helpers.docx_callouts import add_convention_box  # noqa: E402
from helpers.docx_styling import add_field, paragraph_bottom_rule, set_cell_borders, shade_cell  # noqa: E402
from helpers.docx_tables import add_template_table  # noqa: E402
from brand import ADDRESS, CREAM, GREY, TEAL  # noqa: E402
from helpers.scenario_document import build_header_footer, configure_styles, wordmark  # noqa: E402

from build_deployment_report_template import KE_AT2  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402

# Course-side note styling — the same violet convention the Business Case template uses for
# anything addressed to the student rather than to the reader of the finished report.
NOTE = "6A3EA1"
NOTE_PT = 9.0

SUPPLIED_NOTE = ("Supplied — this section is already written for you. It is part of your report: "
                 "leave the text below as it is. Delete this note before you submit.")

# §2 and §3 are the same for every consultant on this engagement — the engagement, the
# stakeholders, the predecessor work and the scope all come from the approved design. They are
# supplied rather than written, so nothing here is student work.
ENGAGEMENT_CONTEXT = [
    "This deployment is the foundation-build phase of the YAT College Learning Management System "
    "(LMS) migration to AWS, carried out by MTS under the engagement agreed following the LMS "
    "Replacement Business Case. The board approved the action plan set out in that business case, "
    "and this phase implements the first stage of it.",
    "MTS Senior Architecture and YAT IT subsequently produced the YAT LMS Cloud Architecture — "
    "Baseline Design, approved by Pat Lin (MTS Senior Consultant) and Sam Walker (YAT IT Manager). "
    "That design is the specification this deployment implements; where it leaves a decision to the "
    "implementer, the decision and its rationale are recorded in §5.",
    "On completion the infrastructure is handed to YAT in-house IT, who are responsible for "
    "installing the LMS application, migrating the database, and running the cutover. Hardening the "
    "environment for high availability is a separate, later phase.",
]

SCOPE_IN = [
    "Identity and access management — groups, users, MFA enforcement and instance roles",
    "Network topology — the VPC, its subnets, gateways and route tables",
    "Compute — the launch template, EC2 instances and the Auto Scaling group",
    "Load balancing — the application load balancer, its target group and listener",
    "Database — the managed relational database instance",
    "Storage — the EBS volumes and S3 buckets, with encryption and public-access settings",
    "Security — the tiered security-group model and encryption in transit and at rest",
    "Monitoring — the baseline alarm set",
]

SCOPE_DEFERRED = [
    "Multi-availability-zone database deployment",
    "Resilience across availability zones",
    "Failure and resize simulation, and the availability measurement that goes with it",
    "Cross-Region backup and replication",
    "Disaster-recovery runbooks",
    "The high-availability-tuned monitoring set",
]

# §6 tests. Each is (number, title, what it shows, [steps], what the screenshot must show).
# One test per UoC item, no more: connect to the instance [401 PE 3]; reach the internet and reach
# the database [401 PC 2.6]; reach the database and the load balancer [502 PC 4.2, all tiers];
# scale out and back in [401 PC 3.2].
#
# [TBD - the connection method in 6.1 assumes the design review enables SSH access to the
# application instance (security group, key pair, and a route in). The exact commands in 6.2-6.4
# depend on that decision and on the instance operating system, and must be confirmed in a live
# lab run before this template is issued.]
TESTS = [
    ("6.1", "Connect to the application server",
     "you can reach the application server you built.",
     ["In the AWS console, open EC2 and find your application instance in the instance list.",
      "Copy its address from the details panel.",
      "Open a terminal on your own computer and connect to the instance using the key pair and "
      "the connection command given in the engagement's build instructions.",
      "When the prompt appears, run: hostname"],
     "your terminal, showing the connection succeeding and the hostname of the instance."),

    ("6.2", "Reach the internet from the application server",
     "the application server can reach the internet through the NAT gateway, so it can be patched "
     "and can call external services.",
     ["In the session you opened in 6.1, run: curl -I https://aws.amazon.com",
      "Confirm the response begins with an HTTP status line."],
     "the command and the HTTP response header it returned."),

    ("6.3", "Reach the database from the application server",
     "the application tier can reach the database tier privately, over port 3306.",
     ["In the AWS console, open RDS and copy your database's endpoint address.",
      "In the session from 6.1, run: nc -zv <your-database-endpoint> 3306",
      "Confirm the result reports the connection succeeded."],
     "the command and its output confirming the connection to port 3306 succeeded."),

    ("6.4", "Reach the load balancer from the application server",
     "the application tier and the web tier can see each other — connectivity across all three "
     "tiers, taken with 6.3.",
     ["In the AWS console, open EC2 → Load Balancers and copy your load balancer's DNS name.",
      "In the session from 6.1, run: curl -I http://<your-load-balancer-dns-name>",
      "Confirm an HTTP status line comes back."],
     "the command and the HTTP response from the load balancer."),

    ("6.5", "Automatic scaling",
     "the Auto Scaling group adds and removes instances on its own, without you touching the "
     "instance count.",
     ["In the AWS console, open EC2 → Auto Scaling Groups and select your group.",
      "On the Automatic scaling tab, edit your scaling policy and lower the target value far "
      "enough that current usage is above it (for example, to 10).",
      "Wait for the alarm to trigger. On the Activity tab, watch for a new instance to launch.",
      "Once the new instance is in service, edit the policy again and raise the target value well "
      "above current usage (for example, to 90).",
      "Wait for the group to scale back in, then return the target to the value you set in §5."],
     "the Activity tab, showing both the scale-out and the scale-in entries with their timestamps."),
]

SCOPE_EXCLUDED = [
    "LMS application installation on the infrastructure built here",
    "Migration of the existing database content",
    "Cutover from the legacy environment, and the change management around it",
    "Ongoing application support after handover",
]


def add_supplied_note(doc):
    """The violet course-side marker that heads a supplied (pre-written) section."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(SUPPLIED_NOTE)
    r.italic = True
    r.font.size = Pt(NOTE_PT)
    r.font.color.rgb = RGBColor.from_string(NOTE)


def add_supplied_body(doc, text):
    """A paragraph of supplied report content — normal body styling, not guidance styling."""
    p = doc.add_paragraph()
    p.add_run(text).font.size = Pt(10.5)
    return p


def add_supplied_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.size = Pt(10.5)


def add_screenshot_slot(doc, caption):
    """A bordered drop-zone for a screenshot, with a caption saying what must be visible."""
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    set_cell_borders(cell)
    shade_cell(cell, CREAM)
    cell.width = Cm(16.6)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("[ PASTE YOUR SCREENSHOT HERE ]")
    r.bold = True; r.font.size = Pt(10)
    p2 = cell.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(caption)
    r2.italic = True; r2.font.size = Pt(9); r2.font.color.rgb = RGBColor.from_string(GREY)
    doc.add_paragraph()


def add_test(doc, h3, number, title, shows, steps, capture):
    """One instrumented test: what it shows, numbered steps to follow, and a screenshot slot."""
    h3(f"{number} {title}")
    add_guidance_text(doc, f"What this test shows: {shows}")
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(3)
        lead = p.add_run(f"{i}.  ")
        lead.bold = True; lead.font.size = Pt(10.5)
        p.add_run(step).font.size = Pt(10.5)
    doc.add_paragraph()
    add_screenshot_slot(doc, capture)


def build(path):
    doc = Document()
    configure_styles(doc)

    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    build_header_footer(sec)

    # ---- COVER ----
    wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(GREY)
    paragraph_bottom_rule(doc.add_paragraph(), TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Deployment Report")
    sub = doc.add_paragraph().add_run("[ Deployment / engagement name ]")
    sub.font.size = Pt(15); sub.italic = True; sub.font.color.rgb = RGBColor.from_string(GREY)
    doc.add_paragraph()

    cover_rows = [
        ("Engagement", "[ Engagement name ]"),
        ("Engagement reference", "[ Reference ]"),
        ("Report version", "[ e.g. v1.0 ]"),
        ("Prepared by", "[ Consultant name ]"),
        ("Consultant role", "[ Role, e.g. MTS Consultant ]"),
        ("Date submitted", "[ DD/MM/YYYY ]"),
        ("Submitted to", "[ Acceptance authority / sponsor ]"),
        ("Related documents", "[ e.g. the approved design this report implements ]"),
        ("Classification", "Commercial-in-confidence"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    ct.alignment = WD_TABLE_ALIGNMENT.LEFT
    for k, v in cover_rows:
        cells = ct.add_row().cells
        set_cell_borders(cells[0]); set_cell_borders(cells[1]); shade_cell(cells[0], CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        vr = cells[1].paragraphs[0].add_run(v); vr.font.size = Pt(10); vr.italic = True
        vr.font.color.rgb = RGBColor.from_string(GREY)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS + convention ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    doc.add_paragraph("How to use this template", style="Heading 1")
    add_convention_box(doc, [
        ("Complete every section that asks you for a response.", "Every section in this template "
         "is one the report needs — there is nothing here to skip."),
        ("Two sections are supplied.", "§2 and §3 are already written for you and are part of your "
         "report. Leave them as they are, and delete the violet note above each before you submit."),
        ("Cross-reference your evidence.", "The build narrative and testing sections reference the "
         "screenshots and test evidence captured in the appendices."),
    ])
    doc.add_paragraph("Contents", style="Heading 1")
    add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
              placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h3 = lambda t: doc.add_paragraph(t, style="Heading 3")

    h1("2. Engagement Context")
    add_supplied_note(doc)
    for para in ENGAGEMENT_CONTEXT:
        add_supplied_body(doc, para)

    h1("3. Scope of Deployment")
    add_supplied_note(doc)
    add_supplied_body(doc, "In scope of this deployment:")
    add_supplied_bullets(doc, SCOPE_IN)
    add_supplied_body(doc, "Deferred to the follow-on high-availability phase:")
    add_supplied_bullets(doc, SCOPE_DEFERRED)
    add_supplied_body(doc, "Outside the MTS engagement entirely — YAT in-house IT's responsibility:")
    add_supplied_bullets(doc, SCOPE_EXCLUDED)

    h1("4. Build Narrative")
    add_guidance_text(doc, "A layer-by-layer account of what was built. For each layer, write a short narrative of "
                     "what you stood up, and cross-reference the Appendix A screenshots.")
    for n, title, hint in [
        ("4.1", "Identity and access management (IAM)", "account access, groups/users/roles, MFA, instance profiles"),
        ("4.2", "Network topology", "VPC, subnets, gateways, route tables"),
        ("4.3", "Compute (EC2 + Auto Scaling)", "launch template, instance type + why, ASG min/desired/max + scaling policy"),
        ("4.4", "Load balancing (ALB)", "the load balancer, target group, listener, health check"),
        ("4.5", "Database (RDS)", "instance class + why, engine version, storage, encryption, backups"),
        ("4.6", "Storage (EBS + S3)", "EBS volumes, S3 buckets, encryption, public-access settings"),
        ("4.7", "Security (security groups + encryption)", "the tiered security-group model, encryption in transit + at rest"),
        ("4.8", "Monitoring", "the baseline CloudWatch alarms and thresholds"),
    ]:
        h3(f"{n} {title}")
        add_guidance_text(doc, f"Cover: {hint}. Cross-reference the relevant Appendix A screenshots.")
        add_response_placeholder(doc)

    h1("5. Configuration Decisions")
    add_guidance_text(doc, "The approved design leaves two sizing decisions to the implementer (design §4.16). For "
                     "each, name at least two options you considered, state the one you chose, and say why it "
                     "suits the YAT LMS workload described in the LMS Application Specification.")
    add_template_table(doc, ["#", "Decision point", "Options you considered", "Your choice", "Why this one"],
             [["C1", "Application-tier instance type", "[ two candidates ]", "[ … ]",
               "[ concurrent-user load from the application spec ]"],
              ["C2", "Database instance class and storage size", "[ two candidates ]", "[ … ]",
               "[ database workload; current data footprint and its growth ]"]],
             widths=[0.9, 3.6, 3.6, 2.6, 4.8])

    h1("6. Testing and Validation")
    add_guidance_text(doc, "Five tests confirm the deployment works. Each one below tells you what it demonstrates "
                     "and the exact steps to run it. Follow the steps, then paste the screenshot into the box "
                     "provided. If a test does not pass, fix the problem, re-run it, and note what you changed.")
    for n, title, shows, steps, capture in TESTS:
        add_test(doc, h3, n, title, shows, steps, capture)

    h1("7. Operational Handover")
    add_guidance_text(doc, "Hand-over information for the team taking over the infrastructure.")
    h3("7.1 Access")
    add_guidance_text(doc, "Who has what access post-handover, MFA enforcement, any IAM group changes.")
    add_response_placeholder(doc)
    h3("7.2 Runbook references")
    add_guidance_text(doc, "Pointers to the design document, naming/tagging conventions, backup arrangements, and "
                     "the alarms list + notification destinations.")
    add_response_placeholder(doc)
    h3("7.3 Known limitations and what's next")
    add_guidance_text(doc, "Be explicit about what is not covered today and what a later phase would add.")
    add_response_placeholder(doc)
    h3("7.4 Documentation filing")
    add_template_table(doc, ["Item", "Filed in", "Reference"],
             [["This Deployment Report", "[ YAT ICT shared documentation ]", "[ ref ]"],
              ["Test evidence (§6)", "[ … ]", "[ ref ]"]],
             widths=[6.5, 5.0, 4.0])
    h1("8. Knowledge Evidence Responses")
    for text, placeholder in KE_AT2:
        add_guidance_text(doc, text)
        add_response_placeholder(doc, placeholder or "[ Write your response here ]")

    # ---- APPENDICES ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1("Appendix A — Build evidence (screenshots)")
    add_guidance_text(doc, "Capture a console screenshot evidencing each component you built, with the region "
                     "indicator visible. List each below and cross-reference it from §4 / §6. Examples:")
    add_template_table(doc, ["#", "Screenshot", "What must be visible"],
             [["A1", "IAM groups / MFA", "[ created groups; MFA enabled ]"],
              ["A2", "VPC subnets", "[ subnets + AZs ]"],
              ["A3", "EC2 instances + ASG", "[ running instances; ASG min/desired/max ]"],
              ["A4", "ALB target group health", "[ healthy targets ]"],
              ["A5", "RDS database", "[ available; encryption ]"],
              ["A6", "CloudWatch alarms / dashboard", "[ the baseline alarms ]"],
              ["…", "[ add as your deployment requires ]", "[ … ]"]],
             widths=[1.0, 5.5, 9.0])

    h1("Document control")
    add_template_table(doc, ["Field", "Value"],
             [["Document version", "[ v1.0 ]"],
              ["Author", "[ Name, role ]"],
              ["Engagement", "[ Engagement name ]"],
              ["Date submitted", "[ DD/MM/YYYY ]"],
              ["Distribution", "[ … ]"],
              ["Related documents", "[ the design implemented; predecessor/successor reports ]"]],
             widths=[5.0, 10.5])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    out = Path("../diploma-cloud-cyber-website-s1/public/templates/AT2-Deployment-Report-Template.docx")
    build(out)
