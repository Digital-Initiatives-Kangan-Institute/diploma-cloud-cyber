#!/usr/bin/env python3
"""Build the YAT / MTS Deployment Report template for the LMS HA hardening (.docx).

A deployment-report template scoped to ONE deployment — the YAT LMS high-availability hardening
phase — holding exactly and only the sections that deployment requires. Split from the generic
superset template (build_deployment_report_template.py) on the same basis as the AT2 template.

A section is here because an AT3-mapped UoC criterion requires it, or because the report does not
hold together without it (§2, §3 and §3.1, which are supplied rather than written). Nothing carries
an "Applicability / mark Not applicable" note except where the section is still open — see below.

What the HA hardening does not have, and so is not in here: the foundation regression tests the
superset carries at §6.2-§6.4 (autoscaling, database connectivity, smoke test — no AT3 criterion
marks them, and B3's §6.1 connectivity tests already cover the post-change regression check), and
Appendix B configuration exports (no criterion; B13 evidences the build through the Appendix A
screenshots and B14 through Appendix C, the same reasoning the AT2 template used).

Dropping §6.2-§6.4 renumbers §6 to what criteria B3-B8 and the student instructions already say:
6.1 connectivity, 6.2 failure simulation, 6.3 resize simulation, 6.4 availability measurement,
6.5 findings vs the design, 6.6 adjustments. The superset numbering the previous template carried
(failure simulation at §6.5) did not match either.

[TBD - needs discussion: §4.9 Cross-Region backup / replication, §5 Configuration Decisions and
§7.1-§7.3 Access / Runbook references / Known limitations. No AT3 criterion marks any of them; the
student instrument's MARKING_NOTE tells students to complete §5 and §7.1-§7.3 anyway. All four are
carried over from the superset UNCHANGED pending that decision — §4.9 keeps its applicability note
and §5 keeps the superset's decision rows, because editing either would settle the question rather
than leave it open. §5's rows are the foundation-build decisions (EC2 instance type, RDS instance
class, storage sizing, ASG threshold); if §5 stays, they need replacing with the decisions this
phase actually leaves to the implementer.]

Usage:  python scripts/templates/build_at3_deployment_report_template.py
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # content-repo scripts/ (brand + registry)  # noqa: E402
sys.path.insert(0, str(next(d / "scripts" for d in Path(__file__).resolve().parents if (d / "scripts" / "helpers" / "__init__.py").exists())))  # umbrella scripts/ (engine)  # noqa: E402
from helpers.docx_body_text import add_guidance_text, add_response_placeholder  # noqa: E402
from helpers.docx_callouts import add_applicability_note, add_convention_box  # noqa: E402
from helpers.docx_styling import add_field, paragraph_bottom_rule, set_cell_borders, shade_cell  # noqa: E402
from helpers.docx_tables import add_template_table  # noqa: E402
from brand import ADDRESS, CREAM, GREY, TEAL  # noqa: E402
from helpers.scenario_document import build_header_footer, configure_styles, wordmark  # noqa: E402

from build_deployment_report_template import KE_AT3, REFLECT_AT3  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402

# Course-side note styling — the same violet convention the Business Case and AT2 Deployment
# Report templates use for anything addressed to the student rather than to the reader of the
# finished report.
NOTE = "6A3EA1"
NOTE_PT = 9.0

SUPPLIED_NOTE = ("Supplied — this section is already written for you. It is part of your report: "
                 "leave the text below as it is. Delete this note before you submit.")

# §2, §3 and §3.1 are the same for every consultant on this engagement — the engagement history,
# the scope of the HA phase and the maintenance window all come from the approved HA Design and the
# scenario. They are supplied rather than written, so nothing here is student work.
ENGAGEMENT_CONTEXT = [
    "This deployment is the high-availability hardening phase of the YAT College Learning Management "
    "System (LMS) migration to AWS, carried out by MTS. It is the final phase of the engagement that "
    "began with the LMS Replacement Business Case and continued through the foundation build.",
    "At the end of the foundation-build phase MTS handed YAT a single-availability-zone environment "
    "that was ready to operate but did not yet meet YAT's strategic 99.9% availability target. YAT "
    "in-house IT installed the LMS application on that infrastructure, migrated the database and ran "
    "the cutover; the LMS is now in production on AWS. Having reviewed the foundation-build "
    "Deployment Report, the YAT board approved this HA hardening phase.",
    "The HA Design produced for this phase — reviewed and approved by Pat Lin (MTS Senior Consultant) "
    "and Sam Walker (YAT ICT Manager) — is the specification this deployment implements.",
    "On completion the HA-hardened infrastructure is handed back to YAT in-house IT. Re-validating "
    "the LMS application against the hardened infrastructure is their responsibility, and the MTS "
    "engagement closes.",
]

SCOPE_IN = [
    "Identity and access management — any changes arising from the HA work",
    "Network topology — the additional subnets giving coverage across availability zones",
    "Compute — the Auto Scaling group expanded across availability zones",
    "Load balancing — application load balancer targets registered across availability zones",
    "Database — conversion of the managed relational database to a Multi-AZ deployment",
    "Storage — HA-related adjustments to volumes, buckets and the backup baseline",
    "Security — security-group adjustments arising from the HA changes",
    "Monitoring — the HA-tuned alarm set and availability tracking",
    "Verification — failure and resize simulation, and availability measurement across the window",
]

SCOPE_EXCLUDED = [
    "Re-validating the LMS application against the HA-hardened infrastructure",
    "Application-layer tuning, such as session affinity during a database failover",
    "Any further LMS application, data or cutover work — completed in earlier phases",
    "Ongoing application support after handover",
]

MAINTENANCE_WINDOW = [
    "This work is performed on a live production system inside a maintenance window: a Saturday "
    "late-night window of approximately 3.5 hours, chosen as the LMS's lowest-traffic period.",
    "Brief service interruptions within the window are expected and acceptable — the conversion of "
    "the database to Multi-AZ incurs a failover of roughly 30 to 60 seconds, and re-registering load "
    "balancer targets can cause a momentary blip.",
    "At the end of the window the LMS infrastructure must be back online, either with the HA "
    "hardening complete or cleanly rolled back to its pre-window state. Leaving the environment "
    "part-way through the change is not acceptable.",
]


def add_supplied_note(doc):
    """The violet course-side marker that heads a supplied (pre-written) section."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(SUPPLIED_NOTE)
    r.italic = True
    r.font.size = Pt(NOTE_PT)
    r.font.color.rgb = RGBColor.from_string(NOTE)


def add_supplied_body(doc, text):
    """A paragraph of supplied report content — normal body styling, not guidance styling."""
    p = doc.add_paragraph()
    p.add_run(text).font.size = Pt(10.5)
    return p


def add_supplied_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.size = Pt(10.5)


def build(path):
    doc = Document()
    configure_styles(doc)

    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    build_header_footer(sec)

    # ---- COVER ----
    wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(GREY)
    paragraph_bottom_rule(doc.add_paragraph(), TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Deployment Report")
    sub = doc.add_paragraph().add_run("[ Deployment / engagement name ]")
    sub.font.size = Pt(15); sub.italic = True; sub.font.color.rgb = RGBColor.from_string(GREY)
    doc.add_paragraph()

    cover_rows = [
        ("Engagement", "[ Engagement name ]"),
        ("Engagement reference", "[ Reference ]"),
        ("Report version", "[ e.g. v1.0 ]"),
        ("Prepared by", "[ Consultant name ]"),
        ("Consultant role", "[ Role, e.g. MTS Consultant ]"),
        ("Date submitted", "[ DD/MM/YYYY ]"),
        ("Submitted to", "[ Acceptance authority / sponsor ]"),
        ("Related documents", "[ e.g. the approved design this report implements ]"),
        ("Classification", "Commercial-in-confidence"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    ct.alignment = WD_TABLE_ALIGNMENT.LEFT
    for k, v in cover_rows:
        cells = ct.add_row().cells
        set_cell_borders(cells[0]); set_cell_borders(cells[1]); shade_cell(cells[0], CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        vr = cells[1].paragraphs[0].add_run(v); vr.font.size = Pt(10); vr.italic = True
        vr.font.color.rgb = RGBColor.from_string(GREY)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS + convention ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    doc.add_paragraph("How to use this template", style="Heading 1")
    add_convention_box(doc, [
        ("Complete every section that asks you for a response.", "Every section in this template "
         "is one the report needs — there is nothing here to skip."),
        ("Two sections are supplied.", "§2 and §3 — including the maintenance-window context at "
         "§3.1 — are already written for you and are part of your report. Leave them as they are, "
         "and delete the violet note above each before you submit."),
        ("Cross-reference your evidence.", "The build narrative and testing sections reference the "
         "screenshots and simulation evidence captured in the appendices."),
    ])
    doc.add_paragraph("Contents", style="Heading 1")
    add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
              placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h3 = lambda t: doc.add_paragraph(t, style="Heading 3")

    h1("1. Executive Summary")
    add_guidance_text(doc, "Write this last. A ≤ 1-page summary the reader sees first: what was hardened; the "
                     "availability-zone footprint you now have; whether the availability target is met; the "
                     "headline outcomes of your failure and resize simulations; and any limitations or items "
                     "deferred to a later phase. ~250–400 words.")
    add_response_placeholder(doc)

    h1("2. Engagement Context")
    add_supplied_note(doc)
    for para in ENGAGEMENT_CONTEXT:
        add_supplied_body(doc, para)

    h1("3. Scope of Deployment")
    add_supplied_note(doc)
    add_supplied_body(doc, "In scope of this deployment:")
    add_supplied_bullets(doc, SCOPE_IN)
    add_supplied_body(doc, "Outside the MTS engagement — YAT in-house IT's responsibility:")
    add_supplied_bullets(doc, SCOPE_EXCLUDED)
    h3("3.1 Maintenance-window context")
    add_supplied_note(doc)
    for para in MAINTENANCE_WINDOW:
        add_supplied_body(doc, para)

    h1("4. Build / Change Narrative")
    add_guidance_text(doc, "A layer-by-layer account of what changed. For each layer, write a short narrative of "
                     "what you changed from the baseline you started the window with, and cross-reference the "
                     "Appendix A screenshots.")
    for n, title, hint in [
        ("4.1", "Identity and access management (IAM)", "any IAM additions or changes the HA work required — state if none"),
        ("4.2", "Network topology (cross-AZ)", "the subnets, gateways and route-table entries added to give coverage across availability zones"),
        ("4.3", "Compute (cross-AZ Auto Scaling)", "the Auto Scaling group's capacity and availability-zone spread after the change, and the scaling policy"),
        ("4.4", "Load balancing (cross-AZ targets)", "the load balancer's availability-zone coverage, the target group and health check after the change"),
        ("4.5", "Database (Multi-AZ)", "the conversion to a Multi-AZ deployment — how and when you applied it, and the failover behaviour it gives you"),
        ("4.6", "Storage (EBS + S3)", "HA-related adjustments to volumes and buckets, and to the backup baseline"),
        ("4.7", "Security (security groups + encryption)", "the security-group adjustments the HA changes required, and any encryption changes"),
        ("4.8", "Monitoring (HA-tuned)", "the HA-tuned alarm set — per-AZ counts, failover detection, replica lag, and the service-level dashboard"),
    ]:
        h3(f"{n} {title}")
        add_guidance_text(doc, f"Cover: {hint}. Cross-reference the relevant Appendix A screenshots.")
        add_response_placeholder(doc)
    h3("4.9 Cross-Region backup / replication")
    add_applicability_note(doc, "deployments with cross-Region resilience in scope")
    add_guidance_text(doc, "Cross-Region snapshot copies, AWS Backup, or S3 cross-Region replication, if in scope. "
                     "Otherwise mark Not applicable.")
    add_response_placeholder(doc)

    h1("5. Configuration Decisions")
    add_guidance_text(doc, "The approved design leaves specific decisions to the implementer. For each decision you "
                     "made, state your choice and justify it against the workload and requirements. Add rows "
                     "as needed.")
    add_template_table(doc, ["#", "Decision point", "Your decision", "Rationale"],
             [["C1", "EC2 instance type (within the chosen family)", "[ … ]", "[ workload + cost ]"],
              ["C2", "RDS instance class", "[ … ]", "[ … ]"],
              ["C3", "Storage sizing (EBS data volume / RDS storage)", "[ show the calc ]", "[ data footprint + growth ]"],
              ["C4", "ASG scaling threshold", "[ … ]", "[ expected CPU profile ]"],
              ["C5", "RDS Multi-AZ apply timing (HA)", "[ … ]", "[ maintenance-window rationale ]"],
              ["C6", "Post-HA ASG capacity (min ≥ 2 across AZs)", "[ … ]", "[ AZ-failure resilience ]"],
              ["C7", "Cross-Region backup destination (HA)", "[ … ]", "[ … ]"],
              ["…", "[ add further decisions ]", "[ … ]", "[ … ]"]],
             widths=[1.0, 6.0, 4.0, 4.5])

    h1("6. Testing, Simulation and Validation")
    add_guidance_text(doc, "Document the tests and simulations run to verify the HA hardening. For each, state what "
                     "you did, the result, and reference the supporting evidence in Appendix C.")
    h3("6.1 Connectivity tests")
    add_template_table(doc, ["Test", "Outcome (Pass/Fail)", "Notes"],
             [["ALB → EC2 health check, both AZs", "[ ]", "[ ]"],
              ["EC2 → RDS connection (private)", "[ ]", "[ ]"],
              ["EC2 → internet via NAT", "[ ]", "[ ]"],
              ["RDS not publicly reachable (negative test)", "[ ]", "[ ]"]],
             widths=[7.0, 4.0, 4.5])
    h3("6.2 Failure simulation")
    add_guidance_text(doc, "Execute each failure simulation from your HA Design §6.1 (e.g. EC2 termination, RDS "
                     "Multi-AZ failover, AZ partition). Record method, observed outcome, and whether the service "
                     "stayed reachable.")
    add_template_table(doc, ["#", "Simulation", "Method", "Observed outcome", "Reachable throughout?", "Evidence"],
             [["F1", "[ e.g. EC2 termination ]", "[ … ]", "[ … ]", "[ Yes/No ]", "C—"],
              ["F2", "[ e.g. RDS failover ]", "[ … ]", "[ … ]", "[ Yes (brief blip) ]", "C—"]],
             widths=[0.8, 3.0, 3.0, 3.6, 2.6, 1.5])
    h3("6.3 Resize simulation")
    add_guidance_text(doc, "Execute each resize simulation from your HA Design §6.2 (e.g. ASG capacity increase, RDS "
                     "instance class change) and record the availability impact.")
    add_response_placeholder(doc)
    h3("6.4 Availability measurement")
    add_guidance_text(doc, "Describe how you defined and measured availability across the maintenance window (e.g. a "
                     "CloudWatch dashboard, a curl-loop against the ALB), and report the availability you "
                     "recorded across the window.")
    add_response_placeholder(doc)
    h3("6.5 Simulation findings vs the design")
    add_guidance_text(doc, "For each simulation, compare the observed outcome against the expected outcome in your "
                     "HA Design §6; document any divergence and why.")
    add_response_placeholder(doc)
    h3("6.6 Adjustments made per simulation outcomes")
    add_guidance_text(doc, "Any changes made to the architecture, configuration, or monitoring based on what the "
                     "simulations revealed — or an explicit statement that none were needed, with reasoning.")
    add_response_placeholder(doc)

    h1("7. Operational Handover")
    add_guidance_text(doc, "Hand-over information for the team taking over the infrastructure.")
    h3("7.1 Access")
    add_guidance_text(doc, "Who has what access post-handover, MFA enforcement, any IAM group changes.")
    add_response_placeholder(doc)
    h3("7.2 Runbook references")
    add_guidance_text(doc, "Pointers to the design document, naming/tagging conventions, backup arrangements, and "
                     "the alarms list + notification destinations.")
    add_response_placeholder(doc)
    h3("7.3 Known limitations and what's next")
    add_guidance_text(doc, "Be explicit about what is not covered today and what a later phase would add.")
    add_response_placeholder(doc)
    h3("7.4 Documentation filing")
    add_template_table(doc, ["Item", "Filed in", "Reference"],
             [["The HA Design", "[ YAT ICT shared documentation ]", "[ ref ]"],
              ["This Deployment Report", "[ YAT ICT shared documentation ]", "[ ref ]"],
              ["Test / simulation evidence (Appendix C)", "[ … ]", "[ ref ]"]],
             widths=[6.5, 5.0, 4.0])
    h3("7.5 Feedback record")
    add_guidance_text(doc, "The feedback you received on the HA Design and the HA implementation, who it came from, "
                     "your response, and what changed as a result.")
    add_template_table(doc, ["Feedback received", "From", "Your response", "Resulting action"],
             [["[ … ]", "[ … ]", "[ … ]", "[ … ]"]],
             widths=[5.0, 3.0, 4.0, 4.0])
    h3("7.6 Sign-off")
    add_template_table(doc, ["Role", "Name", "Date", "Signature"],
             [["Prepared by", "[ … ]", "", ""],
              ["Reviewed by", "[ … ]", "", ""],
              ["Approved by (acceptance authority)", "[ … ]", "", ""]],
             widths=[5.5, 4.5, 2.5, 3.0])

    h1("8. Knowledge Evidence Responses")
    for text, placeholder in KE_AT3:
        add_guidance_text(doc, text)
        add_response_placeholder(doc, placeholder or "[ Write your response here ]")

    # ---- APPENDICES ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1("Appendix A — Build evidence (screenshots)")
    add_guidance_text(doc, "Capture a console screenshot evidencing each component you changed, with the region "
                     "indicator visible. List each below and cross-reference it from §4 / §6. Examples:")
    add_template_table(doc, ["#", "Screenshot", "What must be visible"],
             [["A1", "Cross-AZ subnets", "[ the subnets and their AZs ]"],
              ["A2", "EC2 instances + ASG", "[ running instances across both AZs; ASG min/desired/max ]"],
              ["A3", "ALB target group health", "[ healthy targets in both AZs ]"],
              ["A4", "RDS database", "[ available; Multi-AZ status; encryption ]"],
              ["A5", "CloudWatch alarms / dashboard", "[ the HA-tuned alarms / service-level dashboard ]"],
              ["…", "[ add as your deployment requires ]", "[ … ]"]],
             widths=[1.0, 5.5, 9.0])
    h1("Appendix C — Test and simulation evidence")
    add_guidance_text(doc, "Attach the evidence supporting the results in §6 — the connectivity test captures, the "
                     "failure and resize simulation captures, the metric graphs, and the availability you "
                     "computed across the window.")
    add_response_placeholder(doc, "[ Test and simulation evidence ]")

    h1("Appendix D — Reflections")
    intro, items = REFLECT_AT3
    add_guidance_text(doc, intro)
    for heading, prompt in items:
        h3(heading)
        add_guidance_text(doc, prompt)
        add_response_placeholder(doc, "[ Write your response here]")

    h1("Document control")
    add_template_table(doc, ["Field", "Value"],
             [["Document version", "[ v1.0 ]"],
              ["Author", "[ Name, role ]"],
              ["Engagement", "[ Engagement name ]"],
              ["Date submitted", "[ DD/MM/YYYY ]"],
              ["Distribution", "[ … ]"],
              ["Related documents", "[ the HA Design implemented; predecessor reports ]"]],
             widths=[5.0, 10.5])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    out = Path("../diploma-cloud-cyber-website-s1/public/templates/AT3-Deployment-Report-Template.docx")
    build(out)
