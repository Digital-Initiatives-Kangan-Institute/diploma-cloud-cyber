#!/usr/bin/env python3
"""Build the YAT / MTS Feedback Record template (.docx) from the brand pack.

ONE generic feedback-record form for presentation / review events where feedback must be
captured as a standalone attachment (rather than in-deliverable — most deliverables carry
their own §7.5-style feedback-record section instead; see docs/document-template-system.md).
Currently referenced by the S1-CL1 AT1 Business Case Appendix 4.

Usage:  python scripts/build_feedback_record_template.py [output.docx]
Default: ../diploma-cloud-cyber-website-s1/public/templates/YAT-Feedback-Record-Template.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # content-repo scripts/ (brand + registry)  # noqa: E402
sys.path.insert(0, str(next(d / "scripts" for d in Path(__file__).resolve().parents if (d / "scripts" / "helpers" / "__init__.py").exists())))  # umbrella scripts/ (engine)  # noqa: E402
from helpers.docx_body_text import add_guidance_text  # noqa: E402
from helpers.docx_callouts import add_convention_box  # noqa: E402
from helpers.docx_styling import set_cell_borders, shade_cell  # noqa: E402
from helpers.docx_tables import add_template_table  # noqa: E402
from brand import ADDRESS, CREAM, GREY, TEAL  # noqa: E402
from helpers.scenario_document import build_header_footer, configure_styles, wordmark  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def build(path):
    doc = Document()
    configure_styles(doc)

    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    build_header_footer(sec)

    # ---- COVER / HEADER BLOCK ----
    wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(GREY)
    doc.add_paragraph("Feedback Record", style="Title")
    sub = doc.add_paragraph().add_run("[ Deliverable presented ]")
    sub.font.size = Pt(15); sub.italic = True; sub.font.color.rgb = RGBColor.from_string(GREY)
    doc.add_paragraph()

    detail_rows = [
        ("Engagement", "[ Engagement name ]"),
        ("Deliverable presented", "[ e.g. Business Case, Solution Design ]"),
        ("Presented by", "[ Student / consultant name ]"),
        ("Presented to", "[ Name(s) and role(s) ]"),
        ("Presentation date", "[ DD/MM/YYYY ]"),
        ("Presentation format", "[ In person / video conference ]"),
    ]
    dt = doc.add_table(rows=0, cols=2)
    dt.alignment = WD_TABLE_ALIGNMENT.LEFT
    for k, v in detail_rows:
        cells = dt.add_row().cells
        set_cell_borders(cells[0]); set_cell_borders(cells[1]); shade_cell(cells[0], CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        vr = cells[1].paragraphs[0].add_run(v); vr.font.size = Pt(10); vr.italic = True
        vr.font.color.rgb = RGBColor.from_string(GREY)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)
    doc.add_paragraph()

    add_convention_box(doc, [
        ("Complete during or immediately after the presentation event.", "Capture feedback as it is "
         "given — do not reconstruct it afterwards from memory."),
        ("One row per distinct piece of feedback.", "Record who raised it, your response at the time "
         "(or on reflection), and the resulting action — or state “No action required” with a reason."),
        ("Attach the completed record to your deliverable.", "As directed by the assessment task "
         "(e.g. as an appendix), alongside any separate sign-off block."),
    ])

    doc.add_paragraph("Feedback received", style="Heading 1")
    add_template_table(doc, ["Feedback received", "From", "Your response", "Resulting action"],
             [["[ … ]", "[ … ]", "[ … ]", "[ … ]"],
              ["[ … ]", "[ … ]", "[ … ]", "[ … ]"],
              ["[ … ]", "[ … ]", "[ … ]", "[ … ]"],
              ["…", "[ add further rows as needed ]", "[ … ]", "[ … ]"]],
             widths=[5.2, 3.0, 4.3, 3.0])

    doc.add_paragraph("Summary", style="Heading 1")
    add_guidance_text(doc, "In a sentence or two, summarise the overall tenor of the feedback and whether "
                     "any follow-up remains outstanding beyond the actions recorded above.")
    p = doc.add_paragraph()
    r = p.add_run("[ Write your summary here ]")
    r.italic = True; r.font.color.rgb = RGBColor.from_string(GREY)

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "../diploma-cloud-cyber-website-s1/public/templates/YAT-Feedback-Record-Template.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
