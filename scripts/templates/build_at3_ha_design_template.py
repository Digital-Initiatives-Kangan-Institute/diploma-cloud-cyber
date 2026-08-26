#!/usr/bin/env python3
"""Build the YAT / MTS HA Design template for the LMS high-availability hardening (.docx).

A solution-design template scoped to ONE design — the YAT LMS HA-equivalent architecture that
supersedes the foundation-build baseline — holding exactly and only the sections that design
requires. Split from the generic superset template (build_solution_design_template.py) on the same
basis as the AT2 and AT3 deployment-report templates.

The superset was itself generalised FROM this design's outline, so there is little AT2-specific
content to strip; what comes out is the generic latitude that generalisation added — the
"Not applicable — reason" convention, the "one design for any cloud solution" framing, the
applicability notes on the sections a greenfield design would skip, and the either-or wording
inside the §4 guidance ("or, for a change, any additions", "(if designed)"). Every section those
notes flagged is a section this design uses in full.

Note that §6 Simulation and Verification Plan is criterion-backed even though no Part A criterion
cites it: Part B's B4, B5 and B7 mark the student's execution of, and comparison against, the plan
authored here.

[TBD - needs discussion: §1 Purpose and Scope, §2.1 Inputs, §5 Implementation Sequencing, §7 Out of
Scope, §8 References and §9 Review and Approval. No AT3 criterion marks any of them, though the
student instrument tells students to produce §5. §9 also overlaps the Deployment Report's §7.5
feedback record and §7.6 sign-off, which B10 and B11 do mark. All six are carried over from the
superset UNCHANGED pending that decision — §5 keeps its applicability note, because removing it
would settle the question rather than leave it open.]

Usage:  python scripts/templates/build_at3_ha_design_template.py
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # content-repo scripts/ (brand + registry)  # noqa: E402
sys.path.insert(0, str(next(d / "scripts" for d in Path(__file__).resolve().parents if (d / "scripts" / "helpers" / "__init__.py").exists())))  # umbrella scripts/ (engine)  # noqa: E402
from helpers.docx_body_text import add_guidance_text, add_response_placeholder  # noqa: E402
from helpers.docx_callouts import add_applicability_note, add_convention_box  # noqa: E402
from helpers.docx_styling import add_field, paragraph_bottom_rule, set_cell_borders, shade_cell  # noqa: E402
from helpers.docx_tables import add_template_table  # noqa: E402
from brand import ADDRESS, CREAM, GREY, TEAL  # noqa: E402
from helpers.scenario_document import build_header_footer, configure_styles, wordmark  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def build(path):
    doc = Document()
    configure_styles(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    build_header_footer(sec)

    # ---- COVER ----
    wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(GREY)
    paragraph_bottom_rule(doc.add_paragraph(), TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("HA Design")
    sub = doc.add_paragraph().add_run("[ Solution / initiative name ]")
    sub.font.size = Pt(15); sub.italic = True; sub.font.color.rgb = RGBColor.from_string(GREY)
    doc.add_paragraph()
    cover = [
        ("Engagement", "[ Engagement name ]"),
        ("Engagement reference", "[ Reference ]"),
        ("Document version", "[ e.g. v1.0 ]"),
        ("Prepared by", "[ Consultant name, role ]"),
        ("Date", "[ DD/MM/YYYY ]"),
        ("Submitted to", "[ Acceptance authority / sponsor ]"),
        ("Related documents", "[ requirements + the baseline design this supersedes ]"),
        ("Classification", "Commercial-in-confidence"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    ct.alignment = WD_TABLE_ALIGNMENT.LEFT
    for k, v in cover:
        cells = ct.add_row().cells
        set_cell_borders(cells[0]); set_cell_borders(cells[1]); shade_cell(cells[0], CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        vr = cells[1].paragraphs[0].add_run(v); vr.font.size = Pt(10); vr.italic = True
        vr.font.color.rgb = RGBColor.from_string(GREY)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS + convention ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    doc.add_paragraph("How to use this template", style="Heading 1")
    add_convention_box(doc, [
        ("Complete every section.", "Every section in this template is one this design needs — "
         "there is nothing here to skip."),
        ("This design hardens an existing environment.", "It supersedes the baseline design of the "
         "environment you are changing, and §3 reviews that baseline before §4 designs its "
         "HA-equivalent."),
        ("This design is implemented by a Deployment Report.", "Business Case (why) → HA Design "
         "(what / how) → Deployment Report (what was built)."),
    ])
    doc.add_paragraph("Contents", style="Heading 1")
    add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                 placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h3 = lambda t: doc.add_paragraph(t, style="Heading 3")

    h1("1. Purpose and Scope")
    add_guidance_text(doc, "What this design covers and what it doesn't (≤ ½ page): the solution being designed, "
                     "the objective it serves, and what is in / out of scope.")
    add_response_placeholder(doc)

    h1("2. Design Inputs and Requirements")
    h3("2.1 Inputs")
    add_guidance_text(doc, "The source documents this design is built from — the approved business case, the "
                     "requirements, the workload/application specification, and (if hardening an existing "
                     "system) the baseline design and the report of what was actually built.")
    add_response_placeholder(doc)
    h3("2.2 Requirements the design must meet")
    add_guidance_text(doc, "State the reliability, recoverability and service-level targets this design is held "
                     "to, drawn from the application specification and the cloud migration requirements. "
                     "Add rows as needed.")
    add_template_table(doc, ["Requirement / metric", "Target"],
             [["Availability", "[ e.g. 99.9% ]"],
              ["RPO (acceptable data loss)", "[ e.g. ≤ 1 hour ]"],
              ["RTO (time to recover)", "[ e.g. ≤ 4 hours ]"],
              ["Support response (severity-1)", "[ e.g. ≤ 1 hour ]"],
              ["Data residency", "[ e.g. Australia ]"],
              ["[ add others ]", "[ … ]"]],
             widths=[9.0, 7.0])

    h1("3. Review of the Baseline Architecture")
    add_guidance_text(doc, "Review the baseline environment — the one this design hardens — against the §2 "
                     "requirements, using the supplied baseline design and the deployment report of what was "
                     "actually built.")
    h3("3.1 Architecture review")
    add_guidance_text(doc, "Review the baseline architecture against the requirements, layer by layer (compute, "
                     "load balancing, database, network, monitoring) — where each currently meets or fails "
                     "the targets.")
    add_response_placeholder(doc)
    h3("3.2 Single points of failure")
    add_template_table(doc, ["Component", "Failure mode", "Consequence"],
             [["[ … ]", "[ … ]", "[ … ]"], ["[ … ]", "[ … ]", "[ … ]"]],
             widths=[5.0, 5.0, 6.0])
    h3("3.3 Recovery objectives — current state")
    add_template_table(doc, ["Component", "Current RPO", "Current RTO", "Meets target?"],
             [["[ … ]", "[ … ]", "[ … ]", "[ Yes/No ]"], ["[ … ]", "[ … ]", "[ … ]", "[ Yes/No ]"]],
             widths=[5.5, 3.5, 3.5, 3.5])
    h3("3.4 Components requiring vertical scaling")
    add_template_table(doc, ["Component", "Vertical scale required for", "Availability impact during scale"],
             [["[ … ]", "[ … ]", "[ … ]"]],
             widths=[5.0, 5.0, 6.0])
    h3("3.5 Review findings summary")
    add_guidance_text(doc, "Summarise the gap between the baseline architecture and the §2 requirements: what's "
                     "met, what isn't, and which components drive the gap (≤ 250 words). Write it for the "
                     "stakeholders who review this design.")
    add_response_placeholder(doc)

    h1("4. Architecture Design")
    add_guidance_text(doc, "The design proper — the HA-equivalent of the baseline. Describe the design relative "
                     "to the baseline reviewed in §3.")
    h3("4.1 Assumptions and constraints")
    add_response_placeholder(doc)
    h3("4.2 AWS account and region")
    add_guidance_text(doc, "Region(s) and account; note any data-residency constraint and any new region (e.g. a "
                     "cross-Region backup destination).")
    add_response_placeholder(doc)
    for n, title, hint in [
        ("4.3", "Identity and access management (IAM)", "any additions or changes to groups/roles/users, MFA and instance profiles the HA work requires — state if none"),
        ("4.4", "Network topology", "VPC, subnets, availability-zone distribution, gateways, route tables (use the subnet table below)"),
        ("4.5", "Compute (EC2 + Auto Scaling)", "instance type, ASG capacity and availability-zone spread, scaling policy, health checks"),
        ("4.6", "Load balancing (ALB)", "the load balancer, target group, listener, availability-zone coverage"),
        ("4.7", "Database (RDS)", "engine + version, instance class, storage, encryption, backups, and the Multi-AZ deployment"),
        ("4.8", "Storage (EBS + S3)", "volumes, buckets, encryption, public-access settings, and any cross-Region copy"),
        ("4.9", "Security", "the tiered security-group model, encryption in transit + at rest, and the adjustments the HA changes require"),
        ("4.10", "Monitoring", "the HA-tuned alarms and the service-level / availability tracking (use the alarm table below)"),
        ("4.11", "Naming and tagging conventions", "the tagging scheme (Project, Environment, Owner, CostCentre, DataClassification, AZ)"),
        ("4.12", "Backup", "the backup baseline and any cross-Region copy / retention design"),
    ]:
        h3(f"{n} {title}")
        add_guidance_text(doc, f"Cover: {hint}.")
        if n == "4.4":
            add_template_table(doc, ["Subnet", "CIDR", "AZ", "Purpose"],
                     [["[ … ]", "[ … ]", "[ … ]", "[ … ]"], ["[ … ]", "[ … ]", "[ … ]", "[ … ]"]],
                     widths=[4.0, 3.5, 4.0, 4.5])
        elif n == "4.10":
            add_template_table(doc, ["Alarm", "Metric", "Threshold", "Triggers"],
                     [["[ … ]", "[ … ]", "[ … ]", "[ … ]"]],
                     widths=[4.0, 4.0, 3.5, 4.5])
        else:
            add_response_placeholder(doc)
    h3("4.13 Recovery objectives — designed state")
    add_template_table(doc, ["Component", "Designed RPO", "Designed RTO", "Notes"],
             [["[ … ]", "[ … ]", "[ … ]", "[ … ]"], ["Overall service", "[ … ]", "[ … ]", "[ meets target ]"]],
             widths=[5.0, 3.5, 3.5, 4.0])
    h3("4.14 Components requiring vertical scaling — designed state")
    add_template_table(doc, ["Component", "Vertical scale required for", "Availability impact (designed)"],
             [["[ … ]", "[ … ]", "[ … ]"]],
             widths=[5.0, 5.0, 6.0])
    h3("4.15 Single points of failure removed")
    add_guidance_text(doc, "Each single point of failure from §3.2, and how this design removes it.")
    add_template_table(doc, ["SPOF (from §3.2)", "Mitigation in this design"],
             [["[ … ]", "[ … ]"], ["[ … ]", "[ … ]"]],
             widths=[7.0, 9.0])

    h1("5. Implementation Sequencing")
    add_applicability_note(doc, "deployments into a running system, where order and rollback matter")
    add_guidance_text(doc, "The order the changes are applied, with per-change duration, expected impact, a "
                     "verification step, and a rollback if it fails. State the total window and buffer.")
    add_template_table(doc, ["#", "Change", "Duration", "Expected impact", "Verification", "Rollback"],
             [["1", "[ … ]", "[ … ]", "[ none ]", "[ … ]", "[ … ]"],
              ["2", "[ … ]", "[ … ]", "[ … ]", "[ … ]", "[ … ]"]],
             widths=[0.8, 4.0, 2.0, 3.0, 3.0, 3.0])

    h1("6. Simulation and Verification Plan")
    add_guidance_text(doc, "The plan the Deployment Report then executes — the simulations that verify this "
                     "design delivers the availability §2 requires.")
    h3("6.1 Failure simulation plan")
    add_template_table(doc, ["#", "Simulation", "Method", "Expected outcome", "Verification"],
             [["F1", "[ … ]", "[ … ]", "[ … ]", "[ … ]"]],
             widths=[0.8, 3.0, 3.5, 4.5, 4.0])
    h3("6.2 Resize simulation plan")
    add_template_table(doc, ["#", "Simulation", "Method", "Expected outcome", "Verification"],
             [["R1", "[ … ]", "[ … ]", "[ … ]", "[ … ]"]],
             widths=[0.8, 3.0, 3.5, 4.5, 4.0])
    h3("6.3 Verification criteria")
    add_guidance_text(doc, "The success criteria — what evidence will demonstrate the design works as intended.")
    add_response_placeholder(doc)

    h1("7. Out of Scope")
    add_guidance_text(doc, "What this design deliberately does not address, and why (e.g. cross-Region "
                     "active-active DR, application-layer HA, items owned by another team).")
    add_response_placeholder(doc)

    h1("8. References")
    add_guidance_text(doc, "The source documents and standards informing this design (requirements, baseline "
                     "design, application specification, reference architectures, industry standards), with "
                     "external sources cited with access dates.")
    add_response_placeholder(doc)

    h1("9. Review and Approval")
    add_guidance_text(doc, "The completed design is submitted to the accepting authority / superior for review. "
                     "The reviewer records their feedback below; the author records how each point was "
                     "addressed; the reviewer then signs off on the design before it is implemented.")
    h3("9.1 Reviewer feedback and author response")
    add_template_table(doc, ["#", "Reviewer feedback / comment", "Author response", "Resulting change"],
             [["1", "[ … ]", "[ … ]", "[ … ]"],
              ["2", "[ … ]", "[ … ]", "[ … ]"],
              ["3", "[ … ]", "[ … ]", "[ … ]"]],
             widths=[0.8, 6.0, 5.2, 4.0])
    h3("9.2 Sign-off")
    add_guidance_text(doc, "The accepting authority records their decision. “Approved with comments” means the "
                     "design proceeds subject to the changes recorded above.")
    add_template_table(doc, ["Role", "Name", "Decision", "Date", "Signature"],
             [["Prepared by (author)", "[ name, role ]", "—", "[ DD/MM/YYYY ]", ""],
              ["Reviewed and approved by (accepting authority)", "[ name, role ]",
               "[ Approved / Approved with comments / Rejected ]", "[ DD/MM/YYYY ]", ""]],
             widths=[4.6, 3.2, 3.6, 2.4, 2.2])

    h1("Document control")
    add_template_table(doc, ["Field", "Value"],
             [["Document version", "[ v1.0 ]"],
              ["Author", "[ Name, role ]"],
              ["Engagement", "[ Engagement name ]"],
              ["Supersedes", "[ the baseline design this replaces ]"],
              ["Implemented by", "[ the Deployment Report that builds this design ]"],
              ["Approval status", "[ Pending / Approved / Rejected with comments ]"]],
             widths=[5.0, 10.5])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    out = Path("../diploma-cloud-cyber-website-s1/public/templates/AT3-HA-Design-Template.docx")
    build(out)
